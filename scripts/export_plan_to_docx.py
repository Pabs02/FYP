import os
from docx import Document
from docx.shared import Pt

ROOT = os.path.dirname(os.path.dirname(__file__))
PLAN_MD = os.path.join(ROOT, "docs", "PLAN.md")
OUTPUT_DOCX = os.path.join(ROOT, "docs", "FYP_Plan.docx")


def md_to_docx(md_path: str, out_path: str) -> None:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.strip() == "":
                doc.add_paragraph("")
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:].strip(), style=None)
            elif line[0:2].isdigit() and line[2:4] == ") ":
                # very naive ordered list support
                doc.add_paragraph(line, style=None)
            else:
                doc.add_paragraph(line)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


if __name__ == '__main__':
    md_to_docx(PLAN_MD, OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")

