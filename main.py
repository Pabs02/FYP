# Reference: Flask Documentation - Quickstart
# https://flask.palletsprojects.com/en/3.0.x/quickstart/
from flask import Flask, render_template, jsonify, request, redirect, url_for, flash, session, g, send_file, abort

# Reference: Flask-Login Documentation - Managing User Sessions
# https://flask-login.readthedocs.io/en/latest/#flask_login.LoginManager
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user

# Reference: Werkzeug Documentation - Password Hashing
# https://werkzeug.palletsprojects.com/en/3.0.x/utils/#werkzeug.security
from werkzeug.security import generate_password_hash, check_password_hash
import os
import time as time_module
import json
import re
import secrets
import requests
from io import BytesIO
from email.message import EmailMessage
import smtplib
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
from datetime import datetime, timedelta, time, timezone

from config import (
	get_flask_config,
	get_supabase_database_url,
	get_openai_api_key,
	get_openai_model_name,
	get_smtp_config,
	get_imap_config,
)
from db_supabase import fetch_all as sb_fetch_all, fetch_one as sb_fetch_one, execute as sb_execute  # type: ignore
from services.chatgpt_client import ChatGPTTaskBreakdownService, ChatGPTClientError, AssignmentReviewResponse
from services.analytics import upcoming_tasks_with_priority, assess_progress, normalise_due_datetime
SUPABASE_URL = get_supabase_database_url()


# Reference: Flask Documentation - Application Setup
# https://flask.palletsprojects.com/en/3.0.x/quickstart/#a-minimal-application
app = Flask(__name__, static_folder="static", template_folder="templates")
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-only-change-me")
app.config["OPENAI_API_KEY"] = get_openai_api_key()
app.config["OPENAI_MODEL_NAME"] = get_openai_model_name()
# During local development, optionally auto-login to speed up testing.
# Keep disabled by default so production logout/auth works as expected.
_SKIP_LOGIN_FOR_TESTING = os.getenv("AUTO_LOGIN_FOR_TESTING", "0") in {"1", "true", "True"}

# Reference: Flask-Login Documentation - Initializing Extension
# https://flask-login.readthedocs.io/en/latest/#flask_login.LoginManager.init_app
# Setup follows Flask-Login quickstart guide
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'


@app.get("/healthz")
def healthz():
	return "ok", 200


@app.before_request
def _activity_log_start():
	# Reference: ChatGPT (OpenAI) - Activity Logging Middleware
	# Date: 2026-02-04
	# Prompt: "I need to log all authenticated user actions with timestamps. It should
	# capture method, path, endpoint, status code, duration, IP, and user agent, and
	# store them in a database table. Can you provide a clean before_request/after_request
	# pattern in Flask?"
	# ChatGPT provided the before_request/after_request pattern to capture timing and
	# persist request metadata for authenticated users.
	"""Capture request start time for activity logging."""
	g.activity_start = time_module.monotonic()


@app.after_request
def _activity_log_after(response):
	# Reference: ChatGPT (OpenAI) - Activity Logging Middleware
	# Date: 2026-02-04
	# Prompt: "I need to log all authenticated user actions with timestamps. It should
	# capture method, path, endpoint, status code, duration, IP, and user agent, and
	# store them in a database table. Can you provide a clean before_request/after_request
	# pattern in Flask?"
	# ChatGPT provided the before_request/after_request pattern to capture timing and
	# persist request metadata for authenticated users.
	"""Log authenticated user activity with timestamps."""
	try:
		if current_user.is_authenticated:
			endpoint = request.endpoint or ""
			path = request.path or ""
			if endpoint != "static" and not path.startswith("/static/") and path != "/favicon.ico":
				start_time = getattr(g, "activity_start", None)
				duration_ms = None
				if start_time is not None:
					duration_ms = int((time_module.monotonic() - start_time) * 1000)
				ip_address = request.headers.get("X-Forwarded-For", "").split(",")[0].strip() or request.remote_addr
				sb_execute(
					"""
					INSERT INTO activity_logs (
						student_id, action_type, method, path, endpoint,
						status_code, duration_ms, ip_address, user_agent, referrer
					) VALUES (
						:student_id, :action_type, :method, :path, :endpoint,
						:status_code, :duration_ms, :ip_address, :user_agent, :referrer
					)
					""",
					{
						"student_id": current_user.id,
						"action_type": "http_request",
						"method": request.method,
						"path": path,
						"endpoint": endpoint,
						"status_code": response.status_code,
						"duration_ms": duration_ms,
						"ip_address": ip_address,
						"user_agent": request.headers.get("User-Agent", "")[:500],
						"referrer": request.referrer or "",
					},
				)
	except Exception:
		pass
	return response


def _format_activity_title(item: Dict[str, Any]) -> str:
	# Reference: ChatGPT (OpenAI) - Activity Log Title Mapping
	# Date: 2026-02-04
	# Prompt: "I want activity logs to show friendly titles instead of raw paths.
	# Can you map endpoints/paths to readable labels and add action verbs by method?"
	# ChatGPT provided the mapping approach with method-based verbs.
	endpoint = (item.get("endpoint") or "").strip()
	path = (item.get("path") or "").strip()
	method = (item.get("method") or "GET").upper()

	endpoint_map = {
		"index": "Dashboard",
		"tasks": "Tasks",
		"task_detail": "Task Detail",
		"calendar_view": "Calendar",
		"analytics": "Analytics",
		"study_planner": "AI Study Planner",
		"ai_workspace": "AI Workspace",
		"course_bot": "AI Course Bot",
		"assignment_review": "Assignment Review",
		"reviews_history": "Review History",
		"subtasks_history": "Microtasks",
		"active_semester": "Semester Overview",
		"lecturer_messages_history": "Sent Messages",
		"lecturer_replies": "Inbox (Replies)",
		"refresh_lecturer_replies": "Check Lecturer Replies",
		"sync_canvas": "Canvas Sync",
		"profile": "Profile",
		"add_data_form": "Add Data",
		"login": "Login",
		"register": "Register",
		"contact_lecturer": "Lecturer Message",
		"voice_parse": "Voice Task Capture",
		"group_workspace": "Group Workspace",
		"group_workspace_invite": "Group Invite",
		"group_workspace_download_file": "Group File Download",
		"toggle_lecture_attendance": "Lecture Attendance",
	}

	base_title = endpoint_map.get(endpoint)
	if not base_title and path.startswith("/tasks/") and path.endswith("/review"):
		base_title = "Assignment Review"
	if not base_title and path.startswith("/tasks/"):
		base_title = "Task Detail"
	if not base_title and path.startswith("/lecturer-replies/"):
		base_title = "Lecturer Replies"

	if not base_title:
		base_title = path or "Unknown Action"

	verb_map = {
		"GET": "Viewed",
		"POST": "Submitted",
		"PUT": "Updated",
		"PATCH": "Updated",
		"DELETE": "Deleted",
	}
	verb = verb_map.get(method, method)
	return f"{verb} {base_title}"


@app.before_request
def auto_login_for_testing() -> None:
	"""Auto-login the first student to simplify local testing when enabled."""
	if not _SKIP_LOGIN_FOR_TESTING or current_user.is_authenticated:
		return
	try:
		user_data = sb_fetch_one(
			"""
			SELECT id, name, email, student_number, canvas_api_token,
			       email_notifications_enabled, email_daily_summary_enabled, last_daily_summary_sent_at
			FROM students
			ORDER BY id
			LIMIT 1
			"""
		)
		if user_data:
			login_user(User(
				id=user_data['id'],
				name=user_data['name'],
				email=user_data['email'],
				student_number=user_data.get('student_number'),
				canvas_api_token=user_data.get('canvas_api_token'),
				email_notifications_enabled=bool(user_data.get('email_notifications_enabled', True)),
				email_daily_summary_enabled=bool(user_data.get('email_daily_summary_enabled', True)),
				last_daily_summary_sent_at=user_data.get('last_daily_summary_sent_at'),
			))
	except Exception:
		# If DB isn't available, fall back to normal auth flow.
		return

_chatgpt_service: Optional[ChatGPTTaskBreakdownService] = None
_AI_ALLOWED_SUFFIXES = {".txt", ".md", ".markdown", ".docx", ".pdf"}
_AI_MAX_UPLOAD_BYTES = 4 * 1024 * 1024
_GROUP_MAX_UPLOAD_BYTES = 8 * 1024 * 1024


def get_chatgpt_service() -> ChatGPTTaskBreakdownService:
	"""Return a singleton ChatGPTTaskBreakdownService configured from Flask settings."""
	global _chatgpt_service
	if _chatgpt_service is not None:
		return _chatgpt_service
	api_key = app.config.get("OPENAI_API_KEY")
	if not api_key:
		raise ChatGPTClientError("OPENAI_API_KEY is not configured.")
	model_name = app.config.get("OPENAI_MODEL_NAME")


	_chatgpt_service = ChatGPTTaskBreakdownService(api_key=api_key, model_name=model_name)
	return _chatgpt_service


# Reference: Flask-Login Documentation - User Class Implementation
# https://flask-login.readthedocs.io/en/latest/#flask_login.UserMixin
# Used UserMixin pattern from Flask-Login docs. Adapted to work with our database structure.
class User(UserMixin):
	"""User model for authentication"""
	def __init__(
		self,
		id: int,
		name: str,
		email: str,
		student_number: Optional[str] = None,
		canvas_api_token: Optional[str] = None,
		email_notifications_enabled: bool = True,
		email_daily_summary_enabled: bool = True,
		last_daily_summary_sent_at: Optional[datetime] = None,
	):
		self.id = id
		self.name = name
		self.email = email
		self.student_number = student_number
		self.canvas_api_token = canvas_api_token
		self.email_notifications_enabled = email_notifications_enabled
		self.email_daily_summary_enabled = email_daily_summary_enabled
		self.last_daily_summary_sent_at = last_daily_summary_sent_at
	
	@staticmethod
	def get(user_id: int) -> Optional['User']:
		"""Get user by ID"""
		user_data = sb_fetch_one(
			"""
			SELECT id, name, email, student_number, canvas_api_token,
			       email_notifications_enabled, email_daily_summary_enabled, last_daily_summary_sent_at
			FROM students
			WHERE id = :id
			""",
			{"id": user_id}
		)
		if user_data:
			return User(
				id=user_data['id'],
				name=user_data['name'],
				email=user_data['email'],
				student_number=user_data.get('student_number'),
				canvas_api_token=user_data.get('canvas_api_token'),
				email_notifications_enabled=bool(user_data.get('email_notifications_enabled', True)),
				email_daily_summary_enabled=bool(user_data.get('email_daily_summary_enabled', True)),
				last_daily_summary_sent_at=user_data.get('last_daily_summary_sent_at'),
			)
		return None
	
	@staticmethod
	def get_by_email(email: str) -> Optional['User']:
		"""Get user by email"""
		user_data = sb_fetch_one(
			"""
			SELECT id, name, email, student_number, canvas_api_token,
			       email_notifications_enabled, email_daily_summary_enabled, last_daily_summary_sent_at
			FROM students
			WHERE email = :email
			""",
			{"email": email}
		)
		if user_data:
			return User(
				id=user_data['id'],
				name=user_data['name'],
				email=user_data['email'],
				student_number=user_data.get('student_number'),
				canvas_api_token=user_data.get('canvas_api_token'),
				email_notifications_enabled=bool(user_data.get('email_notifications_enabled', True)),
				email_daily_summary_enabled=bool(user_data.get('email_daily_summary_enabled', True)),
				last_daily_summary_sent_at=user_data.get('last_daily_summary_sent_at'),
			)
		return None


# Reference: Flask-Login Documentation - User Loader Callback
# https://flask-login.readthedocs.io/en/latest/#flask_login.LoginManager.user_loader
# Required callback function for Flask-Login to load users from session
@login_manager.user_loader
def load_user(user_id):
	"""Load user by ID for Flask-Login"""
	return User.get(int(user_id))


# Reference: Flask Documentation - Custom Template Filters
# https://flask.palletsprojects.com/en/3.0.x/templating/#registering-filters
# Custom filter for Irish date format (DD/MM/YYYY)
@app.template_filter('irish_date')
def format_irish_date(value):
	"""Format date to DD/MM/YYYY"""
	if value is None:
		return ""
	
	# Reference: ChatGPT (OpenAI) - ISO Date Parsing with Timezone Handling
	# Date: 2025-10-12
	# Prompt: "I'm getting ISO date strings from a database that sometimes have 'Z' for UTC timezone. 
	# Python's fromisoformat() doesn't accept 'Z', it needs '+00:00'. Can you show me how to safely 
	# parse these ISO strings and handle the timezone conversion?"
	# ChatGPT provided the pattern for replacing 'Z' with '+00:00' before parsing ISO date strings.
	if isinstance(value, str):
		try:
			from datetime import datetime
			value = datetime.fromisoformat(value.replace('Z', '+00:00'))
		except:
			return value
	
	try:
		return value.strftime('%d/%m/%Y')
	except:
		return str(value)


@app.template_filter('irish_datetime')
def format_irish_datetime(value):
	"""Format datetime to DD/MM/YYYY HH:MM"""
	if value is None:
		return ""
	
	# Reference: ChatGPT (OpenAI) - ISO Date Parsing with Timezone Handling
	# Date: 2025-10-12
	# Prompt: "I'm getting ISO date strings from a database that sometimes have 'Z' for UTC timezone. 
	# Python's fromisoformat() doesn't accept 'Z', it needs '+00:00'. Can you show me how to safely 
	# parse these ISO strings and handle the timezone conversion?"
	# ChatGPT provided the pattern for replacing 'Z' with '+00:00' before parsing ISO date strings.
	if isinstance(value, str):
		try:
			from datetime import datetime
			value = datetime.fromisoformat(value.replace('Z', '+00:00'))
		except:
			return value
	
	try:
		return value.strftime('%d/%m/%Y %H:%M')
	except:
		return str(value)


def test_database_connection():
	"""Check database connection on startup"""
	print("\n" + "="*70)
	print("🚀 STUDENT TASK MANAGEMENT SYSTEM")
	print("="*70)
	
	if SUPABASE_URL:
		try:
			from urllib.parse import urlparse
			parsed = urlparse(SUPABASE_URL)
			safe_url = f"{parsed.scheme}://{parsed.hostname}:{parsed.port}/{parsed.path.strip('/')}"
			print(f"📊 Database: Supabase PostgreSQL")
			print(f"🌍 Host: {parsed.hostname}")
			print(f"🔌 Port: {parsed.port}")
		except:
			print(f"📊 Database: Supabase PostgreSQL (configured)")
		
		try:
			result = sb_fetch_one("SELECT version(), current_database(), current_user")
			print(f"✅ Database Connection: SUCCESS")
			print(f"   └─ Database: {result['current_database']}")
			print(f"   └─ User: {result['current_user']}")
			
			tables = sb_fetch_all("""
				SELECT table_name 
				FROM information_schema.tables 
				WHERE table_schema = 'public' 
				AND table_type = 'BASE TABLE'
				ORDER BY table_name
			""")
			print(f"✅ Tables Found: {len(tables)}")
			for table in tables:
				count_result = sb_fetch_one(f"SELECT COUNT(*) as count FROM {table['table_name']}")
				count = count_result['count'] if count_result else 0
				print(f"   └─ {table['table_name']}: {count} records")
			
			print(f"✅ Frontend: Flask Templates Ready")
			print(f"✅ Backend: Supabase Connected")
			print("="*70 + "\n")
			return True
			
		except Exception as e:
			print(f"❌ Database Connection: FAILED")
			print(f"   └─ Error: {str(e)}")
			print("="*70 + "\n")
			return False
	else:
		print(f"❌ Database URL: NOT CONFIGURED")
		print("="*70 + "\n")
		return False


@app.route("/login", methods=["GET", "POST"])
def login():
	"""User login page and handler"""
	if current_user.is_authenticated:
		return redirect(url_for("index"))
	
	if request.method == "POST":
		email = request.form.get("email", "").strip().lower()
		password = request.form.get("password", "")
		
		if not email or not password:
			flash("Please provide both email and password", "error")
			return render_template("login.html")
		
		user_data = sb_fetch_one(
			"SELECT id, name, email, student_number, password_hash, canvas_api_token FROM students WHERE email = :email",
			{"email": email}
		)
		
		if not user_data:
			flash("Invalid email or password", "error")
			return render_template("login.html")
		
		if not user_data.get('password_hash'):
			flash("Account not activated. Please register first.", "error")
			return render_template("login.html")
		
		# Reference: Werkzeug Documentation - Password Verification
		# https://werkzeug.palletsprojects.com/en/3.0.x/utils/#werkzeug.security.check_password_hash
		# Verifies password against stored hash
		if not check_password_hash(user_data['password_hash'], password):
			flash("Invalid email or password", "error")
			return render_template("login.html")
		
		# Create user object and log in
		user = User(
			id=user_data['id'],
			name=user_data['name'],
			email=user_data['email'],
			student_number=user_data.get('student_number'),
			canvas_api_token=user_data.get('canvas_api_token')
		)
		# Reference: Flask-Login Documentation - Logging Users In
		# https://flask-login.readthedocs.io/en/latest/#flask_login.login_user
		# Creates user session
		login_user(user, remember=True)
		
		sb_execute(
			"UPDATE students SET last_login = NOW() WHERE id = :id",
			{"id": user.id}
		)
		
		flash(f"Welcome back, {user.name}!", "success")
		
		next_page = request.args.get('next')
		return redirect(next_page) if next_page else redirect(url_for('index'))
	
	return render_template("login.html")


@app.route("/register", methods=["GET", "POST"])
def register():
	"""User registration page and handler"""
	if current_user.is_authenticated:
		return redirect(url_for("index"))
	
	if request.method == "POST":
		name = request.form.get("name", "").strip()
		email = request.form.get("email", "").strip().lower()
		password = request.form.get("password", "")
		password_confirm = request.form.get("password_confirm", "")
		student_number = request.form.get("student_number", "").strip() or None
		canvas_api_token = request.form.get("canvas_api_token", "").strip()
		
		if not name or not email or not password:
			flash("All fields are required", "error")
			return render_template("register.html")
		
		if len(password) < 6:
			flash("Password must be at least 6 characters long", "error")
			return render_template("register.html")
		
		if password != password_confirm:
			flash("Passwords do not match", "error")
			return render_template("register.html")
		
		existing = sb_fetch_one("SELECT id FROM students WHERE email = :email", {"email": email})
		if existing:
			flash("Email already registered. Please login instead.", "error")
			return redirect(url_for("login"))
		
		# Reference: Werkzeug Documentation - Password Hashing
		# https://werkzeug.palletsprojects.com/en/3.0.x/utils/#werkzeug.security.generate_password_hash
		# Hashes password before storing in database
		password_hash = generate_password_hash(password)
		
		try:
			sb_execute(
				"""INSERT INTO students (name, email, student_number, password_hash, canvas_api_token, created_at) 
				   VALUES (:name, :email, :student_number, :password_hash, :canvas_api_token, NOW())""",
				{
					"name": name,
					"email": email,
					"student_number": student_number,
					"password_hash": password_hash,
					"canvas_api_token": canvas_api_token if canvas_api_token else None
				}
			)
			
			flash("Registration successful! Please login.", "success")
			return redirect(url_for("login"))
			
		except Exception as e:
			flash(f"Registration failed: {str(e)}", "error")
			return render_template("register.html")
	
	return render_template("register.html")


@app.route("/logout")
@login_required
def logout():
	"""Log out current user"""
	logout_user()
	flash("You have been logged out successfully.", "success")
	return redirect(url_for("login"))


@app.route("/profile")
@login_required
def profile():
	"""User profile page"""
	return render_template("profile.html")


@app.route("/update-profile", methods=["POST"])
@login_required
def update_profile():
	"""Update user profile (Canvas token)"""
	canvas_api_token = request.form.get("canvas_api_token", "").strip()
	student_number = request.form.get("student_number", "").strip() or None
	
	try:
		sb_execute(
			"UPDATE students SET canvas_api_token = :canvas_api_token, student_number = :student_number WHERE id = :id",
			{
				"canvas_api_token": canvas_api_token if canvas_api_token else None,
				"student_number": student_number,
				"id": current_user.id
			}
		)
		
		current_user.canvas_api_token = canvas_api_token if canvas_api_token else None
		current_user.student_number = student_number
		
		if canvas_api_token:
			flash("✅ Canvas API token updated successfully!", "success")
		else:
			flash("Canvas API token removed.", "success")
		
	except Exception as e:
		flash(f"Error updating profile: {str(e)}", "error")
	
	return redirect(url_for("profile"))


@app.route("/update-email-preferences", methods=["POST"])
@login_required
def update_email_preferences():
	email_notifications_enabled = request.form.get("email_notifications_enabled") == "1"
	email_daily_summary_enabled = request.form.get("email_daily_summary_enabled") == "1"
	try:
		sb_execute(
			"""
			UPDATE students
			SET email_notifications_enabled = :email_notifications_enabled,
			    email_daily_summary_enabled = :email_daily_summary_enabled
			WHERE id = :id
			""",
			{
				"email_notifications_enabled": email_notifications_enabled,
				"email_daily_summary_enabled": email_daily_summary_enabled,
				"id": current_user.id,
			}
		)
		current_user.email_notifications_enabled = email_notifications_enabled
		current_user.email_daily_summary_enabled = email_daily_summary_enabled
		flash("Email preferences updated.", "success")
	except Exception as e:
		flash(f"Error updating email preferences: {str(e)}", "error")
	return redirect(url_for("profile"))


@app.route("/email/test", methods=["POST"])
@login_required
def send_test_email():
	error = _send_reminder_email(
		to_email=current_user.email,
		subject="Test Email - Student Planner",
		body="This is a test email from your Student Planner notifications."
	)
	if error:
		flash(f"Failed to send test email: {error}", "error")
	else:
		flash("Test email sent successfully.", "success")
	return redirect(url_for("profile"))


@app.route("/email/daily-summary", methods=["POST"])
@login_required
def send_daily_summary():
	body = _build_daily_summary_email(current_user.id)
	if not body:
		flash("No upcoming tasks to include in the summary.", "warning")
		return redirect(url_for("profile"))
	error = _send_reminder_email(
		to_email=current_user.email,
		subject="Your daily study summary",
		body=body
	)
	if error:
		flash(f"Failed to send daily summary: {error}", "error")
	else:
		flash("Daily summary sent.", "success")
	return redirect(url_for("profile"))



# MAIN APPLICATION ROUTES


@app.route("/")
@login_required
def index():
	_generate_task_reminders(current_user.id)
	# Send daily summary once per day (if enabled)
	# Reference: ChatGPT (OpenAI) - Daily Summary Email Trigger Pattern
	# Date: 2026-01-22
	# Prompt: "I need to send a daily summary email once per day when the user loads the dashboard.
	# Can you help me design a last-sent check and update status flags?"
	# ChatGPT provided the trigger pattern and last-sent guard.
	try:
		student = sb_fetch_one(
			"""
			SELECT email, email_daily_summary_enabled, last_daily_summary_sent_at
			FROM students
			WHERE id = :student_id
			""",
			{"student_id": current_user.id}
		)
		if student and student.get("email") and student.get("email_daily_summary_enabled", True):
			last_sent = student.get("last_daily_summary_sent_at")
			should_send = True
			if last_sent:
				try:
					last_date = last_sent.date() if hasattr(last_sent, "date") else None
					should_send = last_date != datetime.now(timezone.utc).date()
				except Exception:
					should_send = True
			if should_send:
				body = _build_daily_summary_email(current_user.id)
				if body:
					error = _send_reminder_email(
						to_email=student.get("email"),
						subject="Your daily study summary",
						body=body
					)
					status = "sent" if error is None else "failed"
					sb_execute(
						"""
						UPDATE students
						SET last_daily_summary_sent_at = NOW(),
						    daily_summary_status = :status,
						    daily_summary_error = :error
						WHERE id = :student_id
						""",
						{
							"status": status,
							"error": error,
							"student_id": current_user.id
						}
					)
	except Exception as exc:
		print(f"[daily-summary] failed to send summary user={current_user.id} error={exc}")
	try:
		rows = sb_fetch_all(
			"""
			SELECT t.id, t.title, t.status, t.due_date, t.due_at, t.completed_at,
			       t.weight_percentage, t.canvas_score, t.canvas_possible,
			       m.code AS module_code
			FROM tasks t
			JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[dashboard] failed to load tasks user={current_user.id} error={exc}")
		rows = []
	
	dashboard_cards = []
	for item in upcoming_tasks_with_priority(rows, limit=3):
		hours_remaining = max(0, int(item.due_in.total_seconds() // 3600))
		days = hours_remaining // 24
		hours = hours_remaining % 24
		time_label = f"{days}d {hours}h" if days else f"{hours}h"
		score_text = None
		# Reference: ChatGPT (OpenAI) - Canvas Grade Display Formatting
		# Date: 2026-01-22
		# Prompt: "I want to show Canvas grades on the dashboard as score/possible and percentage,
		# with safe handling for missing values. Can you suggest a robust formatting pattern?"
		# ChatGPT provided the formatting pattern with safe fallbacks.
		if item.canvas_score is not None:
			try:
				if item.canvas_possible not in (None, 0):
					percentage = (float(item.canvas_score) / float(item.canvas_possible)) * 100
					score_text = f"{item.canvas_score:.1f}/{item.canvas_possible:.1f} ({percentage:.0f}%)"
				else:
					score_text = f"{item.canvas_score:.1f}"
			except Exception:
				score_text = f"{item.canvas_score}" if item.canvas_score is not None else None
		dashboard_cards.append({
			"id": item.id,
			"title": item.title,
			"module_code": item.module_code,
			"due_at": item.due_at,
			"time_label": time_label,
			"weight": item.weight,
			"priority": item.priority,
			"status": item.status,
			"score_text": score_text,
		})
	
	progress = assess_progress(rows)

	reminders = []
	try:
		reminders = sb_fetch_all(
			"""
			SELECT r.id, r.message, r.reminder_type, r.due_at, r.created_at, r.task_id,
			       t.title, m.code AS module_code
			FROM reminders r
			LEFT JOIN tasks t ON t.id = r.task_id
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE r.student_id = :student_id AND r.is_read = FALSE
			ORDER BY r.due_at NULLS LAST, r.created_at DESC
			LIMIT 5
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[reminders] failed to load reminders user={current_user.id} error={exc}")

	lecturers = []
	try:
		lecturers = sb_fetch_all(
			"""
			SELECT id, name, email, module_code
			FROM lecturers
			ORDER BY name ASC
			"""
		)
	except Exception as exc:
		print(f"[lecturers] failed to load lecturers user={current_user.id} error={exc}")

	lecturer_history = []
	try:
		lecturer_history = sb_fetch_all(
			"""
			SELECT lm.id, lm.subject, lm.message, lm.sent_at, lm.email_status, lm.email_error,
			       l.name AS lecturer_name, l.module_code
			FROM lecturer_messages lm
			LEFT JOIN lecturers l ON l.id = lm.lecturer_id
			WHERE lm.student_id = :student_id
			ORDER BY lm.sent_at DESC
			LIMIT 10
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[lecturers] history load failed user={current_user.id} err={exc}")

	lecturer_draft = session.pop("lecturer_draft", None)
	lecturer_subject = session.pop("lecturer_subject", "")
	lecturer_request = session.pop("lecturer_request", "")
	lecturer_id = session.pop("lecturer_id", "")

	return render_template(
		"index.html",
		upcoming_cards=dashboard_cards,
		progress=progress,
		reminders=reminders,
		lecturers=lecturers,
		lecturer_history=lecturer_history,
		lecturer_draft=lecturer_draft,
		lecturer_subject=lecturer_subject,
		lecturer_request=lecturer_request,
		lecturer_id=lecturer_id,
	)


@app.route("/semester")
@login_required
def active_semester():
	"""Active semester view filtered by a date window."""
	start_str = request.args.get("start")
	end_str = request.args.get("end")
	save_range = request.args.get("save") == "1"

	now = datetime.now(timezone.utc)
	default_start = (now - timedelta(days=120)).date()
	default_end = (now + timedelta(days=120)).date()

	def parse_date(value: str, fallback):
		try:
			return datetime.strptime(value, "%Y-%m-%d").date()
		except Exception:
			return fallback

	saved_start = session.get("semester_start")
	saved_end = session.get("semester_end")

	start_date = parse_date(start_str, parse_date(saved_start or "", default_start))
	end_date = parse_date(end_str, parse_date(saved_end or "", default_end))

	if save_range:
		session["semester_start"] = start_date.isoformat()
		session["semester_end"] = end_date.isoformat()
	start_dt = datetime.combine(start_date, time.min).replace(tzinfo=timezone.utc)
	end_dt = datetime.combine(end_date, time.max).replace(tzinfo=timezone.utc)

	tasks = sb_fetch_all(
		"""
		SELECT t.id, t.title, t.status, t.due_date, t.due_at, t.weight_percentage,
		       m.code AS module_code
		FROM tasks t
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE t.student_id = :student_id
		  AND (
			(t.due_at IS NOT NULL AND t.due_at BETWEEN :start_dt AND :end_dt)
			OR (t.due_at IS NULL AND t.due_date BETWEEN :start_date AND :end_date)
		  )
		ORDER BY t.due_at NULLS LAST, t.due_date NULLS LAST
		""",
		{
			"student_id": current_user.id,
			"start_dt": start_dt,
			"end_dt": end_dt,
			"start_date": start_date,
			"end_date": end_date,
		}
	)

	raw_events = sb_fetch_all(
		"""
		SELECT e.id, e.title, e.start_at, e.end_at, e.location, m.code AS module_code
		FROM events e
		LEFT JOIN modules m ON m.id = e.module_id
		WHERE e.student_id = :student_id
		  AND e.start_at BETWEEN :start_dt AND :end_dt
		ORDER BY e.start_at ASC
		""",
		{
			"student_id": current_user.id,
			"start_dt": start_dt,
			"end_dt": end_dt,
		}
	)

	events_map = {}
	for ev in raw_events:
		start_at = ev.get("start_at")
		time_key = None
		if start_at:
			try:
				time_key = start_at.strftime("%H:%M")
			except Exception:
				time_key = str(start_at)[11:16]
		key = (ev.get("title"), ev.get("module_code"), time_key)
		entry = events_map.get(key)
		if not entry:
			events_map[key] = {
				"title": ev.get("title"),
				"module_code": ev.get("module_code"),
				"time": time_key,
				"first": start_at,
				"last": start_at,
				"count": 1,
			}
		else:
			entry["count"] += 1
			if start_at and entry["first"] and start_at < entry["first"]:
				entry["first"] = start_at
			if start_at and entry["last"] and start_at > entry["last"]:
				entry["last"] = start_at

	events = sorted(events_map.values(), key=lambda item: (item.get("first") or datetime.min.replace(tzinfo=timezone.utc)))

	reviews = sb_fetch_all(
		"""
		SELECT ar.id, ar.filename, ar.ai_score_estimate, ar.reviewed_at,
		       t.title AS task_title, m.code AS module_code
		FROM assignment_reviews ar
		LEFT JOIN tasks t ON t.id = ar.task_id
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE ar.student_id = :student_id
		  AND ar.reviewed_at BETWEEN :start_dt AND :end_dt
		ORDER BY ar.reviewed_at DESC
		LIMIT 50
		""",
		{"student_id": current_user.id, "start_dt": start_dt, "end_dt": end_dt}
	)

	lecturer_messages = sb_fetch_all(
		"""
		SELECT lm.id, lm.subject, lm.message, lm.sent_at, lm.email_status,
		       l.name AS lecturer_name
		FROM lecturer_messages lm
		LEFT JOIN lecturers l ON l.id = lm.lecturer_id
		WHERE lm.student_id = :student_id
		  AND lm.sent_at BETWEEN :start_dt AND :end_dt
		ORDER BY lm.sent_at DESC
		LIMIT 50
		""",
		{"student_id": current_user.id, "start_dt": start_dt, "end_dt": end_dt}
	)

	subtasks = sb_fetch_all(
		"""
		SELECT s.id, s.title, s.is_completed, s.planned_start, s.planned_end,
		       t.title AS task_title, m.code AS module_code
		FROM subtasks s
		JOIN tasks t ON t.id = s.task_id
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE t.student_id = :student_id
		  AND (
			(s.planned_start IS NOT NULL AND s.planned_start BETWEEN :start_date AND :end_date)
			OR (s.planned_end IS NOT NULL AND s.planned_end BETWEEN :start_date AND :end_date)
			OR (s.planned_start IS NULL AND s.planned_end IS NULL AND s.created_at BETWEEN :start_dt AND :end_dt)
		  )
		ORDER BY s.created_at DESC
		LIMIT 100
		""",
		{
			"student_id": current_user.id,
			"start_date": start_date,
			"end_date": end_date,
			"start_dt": start_dt,
			"end_dt": end_dt,
		}
	)

	try:
		activity_row = sb_fetch_one(
			# Reference: ChatGPT (OpenAI) - Activity Log Query Pattern
			# Date: 2026-02-04
			# Prompt: "I need to show activity logs in a semester view. Provide SQL to count
			# activity in a date range and fetch the latest entries for display."
			# ChatGPT provided the COUNT query with date-range filtering.
			"""
			SELECT COUNT(*) as count
			FROM activity_logs
			WHERE student_id = :student_id
			  AND created_at BETWEEN :start_dt AND :end_dt
			""",
			{"student_id": current_user.id, "start_dt": start_dt, "end_dt": end_dt}
		)
		total_activity = activity_row["count"] if activity_row else 0

		activity_logs = sb_fetch_all(
			# Reference: ChatGPT (OpenAI) - Activity Log Query Pattern
			# Date: 2026-02-04
			# Prompt: "I need to show activity logs in a semester view. Provide SQL to count
			# activity in a date range and fetch the latest entries for display."
			# ChatGPT provided the SELECT query with date-range filtering and ordering.
			"""
			SELECT id, action_type, method, path, endpoint, status_code, duration_ms, created_at
			FROM activity_logs
			WHERE student_id = :student_id
			  AND created_at BETWEEN :start_dt AND :end_dt
			ORDER BY created_at DESC
			LIMIT 50
			""",
			{"student_id": current_user.id, "start_dt": start_dt, "end_dt": end_dt}
		)
		for item in activity_logs:
			item["title"] = _format_activity_title(item)
	except Exception:
		total_activity = 0
		activity_logs = []

	status_counts = {"pending": 0, "in_progress": 0, "completed": 0}
	for task in tasks:
		status = (task.get("status") or "pending").lower()
		if status in status_counts:
			status_counts[status] += 1
	total_tasks = sum(status_counts.values())
	completed_subtasks = sum(1 for s in subtasks if s.get("is_completed"))

	# Reference: ChatGPT (OpenAI) - Semester Group Progress + Attendance Aggregation
	# Date: 2026-02-11
	# Prompt: "In a semester overview page, I need SQL summaries for (1) group project
	# progress and (2) lecture attendance stats constrained to the selected date range.
	# Can you provide practical aggregation queries and safe defaults?"
	# ChatGPT provided the aggregation/query pattern below.
	try:
		group_project_progress = sb_fetch_all(
			"""
			SELECT gp.id, gp.title,
			       COUNT(gt.id) AS total_tasks,
			       SUM(CASE WHEN gt.status = 'done' THEN 1 ELSE 0 END) AS completed_tasks
			FROM group_projects gp
			LEFT JOIN group_project_tasks gt
			       ON gt.project_id = gp.id
			      AND (
			          (gt.due_date IS NOT NULL AND gt.due_date BETWEEN :start_date AND :end_date)
			          OR gt.created_at BETWEEN :start_dt AND :end_dt
			      )
			WHERE gp.owner_student_id = :student_id
			  AND (
			      gp.created_at BETWEEN :start_dt AND :end_dt
			      OR (gp.due_date IS NOT NULL AND gp.due_date BETWEEN :start_date AND :end_date)
			      OR EXISTS (
			          SELECT 1
			          FROM group_project_tasks gt2
			          WHERE gt2.project_id = gp.id
			            AND (
			                (gt2.due_date IS NOT NULL AND gt2.due_date BETWEEN :start_date AND :end_date)
			                OR gt2.created_at BETWEEN :start_dt AND :end_dt
			            )
			      )
			  )
			GROUP BY gp.id, gp.title
			ORDER BY gp.title
			""",
			{
				"student_id": current_user.id,
				"start_date": start_date,
				"end_date": end_date,
				"start_dt": start_dt,
				"end_dt": end_dt,
			},
		)
	except Exception:
		group_project_progress = []

	total_group_projects = len(group_project_progress)
	total_group_tasks = sum(int(item.get("total_tasks") or 0) for item in group_project_progress)
	completed_group_tasks = sum(int(item.get("completed_tasks") or 0) for item in group_project_progress)
	group_completion_rate = round((completed_group_tasks * 100.0 / total_group_tasks), 1) if total_group_tasks else 0.0
	for item in group_project_progress:
		total = int(item.get("total_tasks") or 0)
		done = int(item.get("completed_tasks") or 0)
		item["completion_rate"] = round((done * 100.0 / total), 1) if total else 0.0

	try:
		lecture_attendance_summary = sb_fetch_one(
			"""
			SELECT
				COUNT(e.id) AS total_lectures,
				SUM(CASE WHEN COALESCE(la.attended, FALSE) THEN 1 ELSE 0 END) AS attended_lectures
			FROM events e
			LEFT JOIN lecture_attendance la
			       ON la.event_id = e.id
			      AND la.student_id = :student_id
			WHERE e.student_id = :student_id
			  AND e.canvas_event_id IS NOT NULL
			  AND e.start_at BETWEEN :start_dt AND LEAST(:end_dt, NOW())
			""",
			{"student_id": current_user.id, "start_dt": start_dt, "end_dt": end_dt},
		) or {"total_lectures": 0, "attended_lectures": 0}
	except Exception:
		lecture_attendance_summary = {"total_lectures": 0, "attended_lectures": 0}
	total_lectures = int(lecture_attendance_summary.get("total_lectures") or 0)
	attended_lectures = int(lecture_attendance_summary.get("attended_lectures") or 0)
	lecture_attendance_rate = round((attended_lectures * 100.0 / total_lectures), 1) if total_lectures else 0.0

	return render_template(
		"semester.html",
		start_date=start_date,
		end_date=end_date,
		tasks=tasks,
		events=events,
		reviews=reviews,
		lecturer_messages=lecturer_messages,
		subtasks=subtasks,
		status_counts=status_counts,
		total_tasks=total_tasks,
		total_events=len(events),
		total_reviews=len(reviews),
		total_subtasks=len(subtasks),
		completed_subtasks=completed_subtasks,
		total_activity=total_activity,
		activity_logs=activity_logs,
		group_project_progress=group_project_progress,
		total_group_projects=total_group_projects,
		total_group_tasks=total_group_tasks,
		completed_group_tasks=completed_group_tasks,
		group_completion_rate=group_completion_rate,
		total_lectures=total_lectures,
		attended_lectures=attended_lectures,
		lecture_attendance_rate=lecture_attendance_rate,
	)


@app.route("/reminders/<int:reminder_id>/read", methods=["POST"])
@login_required
def mark_reminder_read(reminder_id: int):
	try:
		sb_execute(
			"""
			UPDATE reminders
			SET is_read = TRUE
			WHERE id = :reminder_id AND student_id = :student_id
			""",
			{"reminder_id": reminder_id, "student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[reminders] mark read failed user={current_user.id} err={exc}")
	return redirect(url_for("index"))


@app.route("/reminders/clear", methods=["POST"])
@login_required
def clear_reminders():
	try:
		sb_execute(
			"""
			UPDATE reminders
			SET is_read = TRUE
			WHERE student_id = :student_id
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[reminders] clear failed user={current_user.id} err={exc}")
	return redirect(url_for("index"))


@app.route("/debug/db")
def debug_db():
	try:
		row = sb_fetch_one("SELECT 1 AS ok")
		ver_row = sb_fetch_one("SELECT version() AS version")
		status = {
			"connected": bool(row and row.get("ok") == 1),
			"database": "supabase_postgres",
			"version": ver_row.get("version") if ver_row else None,
		}
		return jsonify(status), 200
	except Exception as exc:
		return jsonify({"connected": False, "error": str(exc)}), 500


@app.route("/tasks")
@login_required
def tasks():
	sql = """
		SELECT t.id, t.title, t.status, t.due_date, 
		       t.canvas_assignment_id, t.canvas_course_id,
		       s.name AS student_name, 
		       m.code AS module_code
		FROM tasks t
		JOIN students s ON s.id = t.student_id
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE t.student_id = :student_id
		ORDER BY t.due_date ASC
		LIMIT 200
	"""
	try:
		rows: List[Dict] = sb_fetch_all(sql, {"student_id": current_user.id})
	except Exception:
		rows = []
	return render_template("tasks.html", tasks=rows)


@app.route("/tasks/<int:task_id>/delete", methods=["POST"])
@login_required
def delete_task(task_id: int):
	"""Delete a task and its associated subtasks"""
	try:
		# First verify the task belongs to the current user
		task = sb_fetch_one(
			"SELECT id FROM tasks WHERE id = :task_id AND student_id = :student_id",
			{"task_id": task_id, "student_id": current_user.id}
		)
		if not task:
			flash("Task not found or permission denied.", "warning")
			return redirect(url_for("tasks"))
		
		# Delete subtasks first (if any)
		sb_execute(
			"DELETE FROM subtasks WHERE task_id = :task_id",
			{"task_id": task_id}
		)
		
		# Delete the task
		result = sb_execute(
			"DELETE FROM tasks WHERE id = :task_id AND student_id = :student_id",
			{"task_id": task_id, "student_id": current_user.id}
		)
		
		if result > 0:
			flash("Task and all associated subtasks deleted successfully.", "success")
		else:
			flash("Task not found or already deleted.", "warning")
	except Exception as exc:
		print(f"[tasks] delete failed user={current_user.id} task={task_id} err={exc}")
		flash("Failed to delete the task. Please try again.", "danger")
	return redirect(url_for("tasks"))


@app.route("/tasks/<int:task_id>", methods=["GET", "POST"])
@login_required
def task_detail(task_id: int):
	"""Display a single task with its micro-task plan and allow creating subtasks."""
	task = sb_fetch_one(
		"""
		SELECT t.id, t.title, t.status, t.due_date, t.due_at, t.description,
		       t.weight_percentage, t.canvas_score, t.canvas_possible, t.canvas_graded_at,
		       t.canvas_assignment_id,
		       m.code AS module_code, m.id AS module_id
		FROM tasks t
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE t.id = :task_id AND t.student_id = :student_id
		""",
		{"task_id": task_id, "student_id": current_user.id}
	)
	if not task:
		flash("Task not found or permission denied.", "error")
		return redirect(url_for("tasks"))
	
	if request.method == "POST":
		action = request.form.get("action", "add_subtask")
		if action == "update_task":
			# Reference: ChatGPT (OpenAI) - Task Update Form Validation and Date Parsing
			# Date: 2025-11-18
			# Prompt: "I need to handle form submission for updating task details. The form can have 
			# title, description, weight, due_date (date only), and due_at (datetime). I need to 
			# validate the title is required, parse dates with multiple formats, handle timezone 
			# conversions, and update the database. Can you help me write this validation and parsing 
			# logic?"
			# ChatGPT provided the form validation and date parsing logic that handles multiple date 
			# formats, timezone conversions, validates required fields, and safely updates the database 
			# with proper error handling.
			title = request.form.get("task_title", "").strip()
			description = request.form.get("task_description", "").strip() or None
			weight = request.form.get("weight_percentage")
			due_date_str = request.form.get("due_date", "").strip() or None
			due_at_str = request.form.get("due_at", "").strip() or None
			
			if not title:
				flash("Title is required.", "error")
				return redirect(url_for("task_detail", task_id=task_id))
			
			weight_value = None
			if weight:
				try:
					weight_value = float(weight)
				except ValueError:
					flash("Weighting must be numeric.", "warning")
					weight_value = None
			
			due_date = None
			due_at = None
			if due_date_str:
				try:
					from datetime import datetime
					due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()
				except ValueError:
					flash("Invalid due date format.", "warning")
			if due_at_str:
				try:
					from datetime import datetime
					due_at = datetime.fromisoformat(due_at_str.replace('Z', '+00:00'))
				except ValueError:
					try:
						# Try parsing as date-time string
						due_at = datetime.strptime(due_at_str, "%Y-%m-%dT%H:%M")
					except ValueError:
						flash("Invalid due date/time format.", "warning")
			
			try:
				sb_execute(
					"""
					UPDATE tasks
					SET title = :title,
					    description = :description,
					    weight_percentage = :weight,
					    due_date = :due_date,
					    due_at = :due_at
					WHERE id = :task_id AND student_id = :student_id
					""",
					{
						"title": title,
						"description": description,
						"weight": weight_value,
						"due_date": due_date,
						"due_at": due_at,
						"task_id": task_id,
						"student_id": current_user.id
					}
				)
				flash("Assignment details updated successfully.", "success")
			except Exception as exc:
				flash(f"Failed to update assignment: {exc}", "error")
			return redirect(url_for("task_detail", task_id=task_id))

		title = request.form.get("title", "").strip()
		description = request.form.get("description", "").strip() or None
		sequence = request.form.get("sequence") or "1"
		est_hours = request.form.get("estimated_hours") or None
		planned_week = request.form.get("planned_week") or None
		planned_start = request.form.get("planned_start") or None
		planned_end = request.form.get("planned_end") or None
		
		if not title:
			flash("Subtask title is required.", "error")
		else:
			try:
				sequence_int = int(sequence)
			except ValueError:
				sequence_int = 1
			try:
				sb_execute(
					"""
					INSERT INTO subtasks (
						task_id, title, description, sequence, estimated_hours,
						planned_week, planned_start, planned_end
					) VALUES (
						:task_id, :title, :description, :sequence, :estimated_hours,
						:planned_week, :planned_start, :planned_end
					)
					""",
					{
						"task_id": task_id,
						"title": title,
						"description": description,
						"sequence": sequence_int,
						"estimated_hours": float(est_hours) if est_hours else None,
						"planned_week": int(planned_week) if planned_week else None,
						"planned_start": planned_start or None,
						"planned_end": planned_end or None,
					}
				)
				sb_execute(
					"UPDATE subtasks SET updated_at = NOW() WHERE task_id = :task_id",
					{"task_id": task_id}
				)
				flash("Subtask added successfully!", "success")
			except Exception as exc:
				flash(f"Failed to add subtask: {exc}", "error")
		return redirect(url_for("task_detail", task_id=task_id))
	
	subtasks = sb_fetch_all(
		"""
		SELECT id, title, description, sequence, estimated_hours,
		       planned_week, planned_start, planned_end,
		       is_completed, completed_at
		FROM subtasks
		WHERE task_id = :task_id
		ORDER BY sequence ASC, id ASC
		""",
		{"task_id": task_id}
	)
	print(f"[task_detail] task_id={task_id} found {len(subtasks)} subtasks")
	

	try:
		reviews = sb_fetch_all(
			"""
			SELECT id, filename, ai_feedback, ai_score_estimate, ai_possible_score,
			       reviewed_at, created_at
			FROM assignment_reviews
			WHERE task_id = :task_id AND student_id = :student_id
			ORDER BY reviewed_at DESC
			LIMIT 5
			""",
			{"task_id": task_id, "student_id": current_user.id}
		)
	except Exception as exc:
		# Table might not exist yet - return empty list
		print(f"[task_detail] Could not fetch reviews (table may not exist): {exc}")
		reviews = []
	
	return render_template("task_detail.html", task=task, subtasks=subtasks, reviews=reviews)


@app.route("/tasks/<int:task_id>/review", methods=["POST"])
@login_required
def review_assignment(task_id: int):
	"""Review and grade an assignment draft using AI"""
	# Reference: ChatGPT (OpenAI) - Assignment Review Orchestration
	# Date: 2026-01-22
	# Prompt: "I want a flow that accepts a file or text input, extracts content, sends it to
	# an AI reviewer, and stores feedback + score. Can you outline a clean end-to-end flow?"
	# ChatGPT provided the end-to-end flow and error-handling pattern.
	# Verify task exists and belongs to user
	task = sb_fetch_one(
		"""
		SELECT t.id, t.title, t.description, t.student_id
		FROM tasks t
		WHERE t.id = :task_id AND t.student_id = :student_id
		""",
		{"task_id": task_id, "student_id": current_user.id}
	)
	if not task:
		flash("Task not found or permission denied.", "error")
		return redirect(url_for("tasks"))
	
	error: Optional[str] = None
	review_result: Optional[AssignmentReviewResponse] = None
	
	# Get assignment text from file or text input
	assignment_text: Optional[str] = None
	filename: Optional[str] = None
	filepath: Optional[str] = None
	
	# Check for file upload
	file_storage = request.files.get("assignment_file")
	if file_storage and file_storage.filename:
		filename = file_storage.filename
		try:
			payload = file_storage.read()
			if not payload:
				raise ValueError("The uploaded file was empty.")
			if len(payload) > _AI_MAX_UPLOAD_BYTES:
				raise ValueError("The file is larger than 4 MB. Upload a smaller file.")
			
			# Extract text from file
			assignment_text = _extract_text_from_brief(filename, payload)
			
			# Store file on filesystem
			uploads_dir = os.path.join(app.root_path, "uploads", "assignments")
			os.makedirs(uploads_dir, exist_ok=True)
			# Create unique filename: task_id_timestamp_originalname
			from datetime import datetime
			timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
			safe_filename = f"{task_id}_{timestamp}_{filename}"
			filepath = os.path.join(uploads_dir, safe_filename)
			with open(filepath, "wb") as f:
				f.write(payload)
			
		except ValueError as exc:
			error = str(exc)
			print(f"[review-assignment] file validation error user={current_user.id} task={task_id} err={exc}")
		except Exception as exc:
			error = f"Failed to process uploaded file: {str(exc)}"
			print(f"[review-assignment] file processing failed user={current_user.id} task={task_id} err={exc}")
			import traceback
			traceback.print_exc()
	
	# Check for text input
	if not assignment_text:
		assignment_text = request.form.get("assignment_text", "").strip()
	
	if not assignment_text:
		error = "Please provide assignment content by uploading a file or pasting text."
	
	if not error:
		try:
			# Get assignment brief from task description
			assignment_brief = task.get("description")
			
			# Call ChatGPT to review and grade
			service = get_chatgpt_service()
			review_result = service.review_and_grade_assignment(
				assignment_text=assignment_text,
				task_title=task.get("title", "Assignment"),
				assignment_brief=assignment_brief
			)
			# Enforce conservative scoring caps based on weaknesses/suggestions
			if review_result and review_result.score_estimate is not None:
				try:
					score_value = float(review_result.score_estimate)
					if review_result.weaknesses or review_result.suggestions:
						score_value = min(score_value, 74.0)
					else:
						score_value = min(score_value, 79.0)
					review_result.score_estimate = score_value
				except Exception:
					pass
			
			# Format comprehensive feedback including strengths, weaknesses, and suggestions
			feedback_parts = [review_result.feedback]
			if review_result.strengths:
				feedback_parts.append("\n\n**Strengths:**\n" + "\n".join(f"• {s}" for s in review_result.strengths))
			if review_result.weaknesses:
				feedback_parts.append("\n\n**Areas for Improvement:**\n" + "\n".join(f"• {w}" for w in review_result.weaknesses))
			if review_result.suggestions:
				feedback_parts.append("\n\n**Suggestions:**\n" + "\n".join(f"• {s}" for s in review_result.suggestions))
			comprehensive_feedback = "\n".join(feedback_parts)
			
			# Save review to database
			sb_execute(
				"""
				INSERT INTO assignment_reviews (
					task_id, student_id, filename, filepath, original_text,
					ai_feedback, ai_score_estimate, ai_possible_score, reviewed_at
				) VALUES (
					:task_id, :student_id, :filename, :filepath, :original_text,
					:ai_feedback, :ai_score_estimate, :ai_possible_score, NOW()
				)
				""",
				{
					"task_id": task_id,
					"student_id": current_user.id,
					"filename": filename,
					"filepath": filepath,
					"original_text": assignment_text[:5000] if assignment_text else None,  # Store first 5000 chars
					"ai_feedback": comprehensive_feedback,
					"ai_score_estimate": review_result.score_estimate,
					"ai_possible_score": review_result.possible_score,
				}
			)
			flash("Assignment reviewed successfully! Check the review section below.", "success")
		except ChatGPTClientError as exc:
			error = f"AI review failed: {str(exc)}"
			print(f"[review-assignment] ChatGPT error user={current_user.id} task={task_id} err={exc}")
		except Exception as exc:
			error = f"Unexpected error: {str(exc)}"
			print(f"[review-assignment] Unexpected error user={current_user.id} task={task_id} err={exc}")
			import traceback
			traceback.print_exc()
	
	if error:
		flash(error, "error")
	
	return redirect(url_for("task_detail", task_id=task_id))


@app.route("/tasks/<int:task_id>/subtasks/<int:subtask_id>/toggle", methods=["POST"])
@login_required
def toggle_subtask(task_id: int, subtask_id: int):
	"""Toggle completion status of a subtask"""
	# Verify the subtask belongs to a task owned by the current user
	subtask = sb_fetch_one(
		"""
		SELECT s.id, s.task_id, s.is_completed
		FROM subtasks s
		JOIN tasks t ON t.id = s.task_id
		WHERE s.id = :subtask_id 
		  AND s.task_id = :task_id
		  AND t.student_id = :student_id
		""",
		{"subtask_id": subtask_id, "task_id": task_id, "student_id": current_user.id}
	)
	
	if not subtask:
		flash("Subtask not found or access denied.", "error")
		return redirect(url_for("task_detail", task_id=task_id))
	
	# Toggle completion status
	new_status = not subtask.get("is_completed", False)
	from datetime import datetime, timezone
	completed_at = datetime.now(timezone.utc) if new_status else None
	
	try:
		sb_execute(
			"""
			UPDATE subtasks 
			SET is_completed = :is_completed, 
			    completed_at = :completed_at,
			    updated_at = NOW()
			WHERE id = :subtask_id
			""",
			{
				"is_completed": new_status,
				"completed_at": completed_at,
				"subtask_id": subtask_id
			}
		)
		flash(f"Subtask marked as {'completed' if new_status else 'pending'}.", "success")
	except Exception as exc:
		flash(f"Failed to update subtask: {exc}", "error")
	
	return redirect(url_for("task_detail", task_id=task_id))


def _iso_date(value):
	if value is None:
		return None
	try:
		return value.strftime("%Y-%m-%d")
	except Exception:
		return str(value)


def _iso_datetime(value):
	if value is None:
		return None
	try:
		return value.isoformat()
	except Exception:
		return str(value)


# Reference: ChatGPT (OpenAI) - Text Summarization Algorithm
# Date: 2025-11-14
# Prompt: "I need a function that summarizes text to a maximum length by extracting complete 
# sentences. It should split text into sentences, add sentences until the max length is 
# reached, and handle edge cases like single very long sentences. Can you help me write this?"
# ChatGPT provided the algorithm for summarizing text by sentence extraction. It uses regex 
# to split sentences, accumulates sentences until the max length, handles single long 
# sentences, and ensures the summary doesn't cut words mid-sentence.
def _summarize_text(raw: str, max_len: int = 400) -> Optional[str]:
	text = raw.strip()
	if not text:
		return None
	try:
		import re
		sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
	except Exception:
		sentences = [s.strip() for s in text.split(".") if s.strip()]
	if not sentences:
		return None
	summary_parts = []
	total_len = 0
	for sentence in sentences:
		if summary_parts:
			add_len = len(sentence) + 1
		else:
			add_len = len(sentence)
		if summary_parts and total_len + add_len > max_len:
			break
		if not summary_parts and len(sentence) > max_len:
			sentence = sentence[:max_len]
			summary_parts.append(sentence)
			total_len = len(sentence)
			break
		summary_parts.append(sentence)
		total_len += add_len
		if total_len >= max_len:
			break
	if not summary_parts:
		return None
	combined = " ".join(summary_parts).strip()
	if len(combined) > max_len:
		combined = combined[:max_len].rsplit(" ", 1)[0] + " ..."
	return combined or None


def _build_reminder_message(task: Dict[str, Any], reminder_type: str) -> str:
	title = task.get("title") or "Untitled task"
	module_code = task.get("module_code")
	module_part = f" ({module_code})" if module_code else ""
	if reminder_type == "overdue":
		return f"Overdue: {title}{module_part}"
	if reminder_type == "due_24h":
		return f"Due in 24h: {title}{module_part}"
	return f"Due soon: {title}{module_part}"


def _send_reminder_email(*, to_email: str, subject: str, body: str) -> Optional[str]:
	# Reference: ChatGPT (OpenAI) - SMTP Email Sender Pattern
	# Date: 2026-01-22
	# Prompt: "Can you help me build a simple SMTP email sender that supports TLS and
	# optional login credentials, returning errors as strings?"
	# ChatGPT provided the SMTP send pattern used here.
	# Reference: ChatGPT (OpenAI) - Resend API Email Sender Fallback
	# Date: 2026-02-12
	# Prompt: "I need a Flask email helper that sends via Resend API first when
	# RESEND_API_KEY is configured, then falls back to SMTP. It should return
	# an error string instead of raising and use request timeouts."
	# ChatGPT provided the provider-priority + graceful-fallback pattern below.
	resend_api_key = (os.getenv("RESEND_API_KEY") or "").strip()
	resend_from_email = (os.getenv("RESEND_FROM_EMAIL") or "").strip()
	if resend_api_key:
		if not resend_from_email:
			return "RESEND_FROM_EMAIL is required when RESEND_API_KEY is set."
		try:
			html_body = "<br>".join(
				line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
				for line in body.splitlines()
			) or body
			response = requests.post(
				"https://api.resend.com/emails",
				headers={
					"Authorization": f"Bearer {resend_api_key}",
					"Content-Type": "application/json",
				},
				json={
					"from": resend_from_email,
					"to": [to_email],
					"subject": subject,
					"html": f"<p>{html_body}</p>",
				},
				timeout=8,
			)
			if response.status_code in {200, 201, 202}:
				return None
			try:
				error_payload = response.json()
			except Exception:
				error_payload = response.text
			return f"Resend API error ({response.status_code}): {error_payload}"
		except Exception as exc:
			return f"Resend API request failed: {exc}"

	smtp_config = get_smtp_config()
	if not smtp_config:
		return "SMTP is not configured and RESEND_API_KEY is missing."
	# Render free instances can block/timeout outbound SMTP sockets.
	# Allow explicitly disabling SMTP attempts so invite/task actions
	# complete without hanging the request worker.
	disable_outbound_email = os.getenv("DISABLE_OUTBOUND_EMAIL", "").strip()
	is_render = os.getenv("RENDER", "").strip().lower() in {"1", "true", "yes"}
	if disable_outbound_email in {"1", "true", "True"} or (is_render and disable_outbound_email == ""):
		return "Outbound email is disabled by configuration (DISABLE_OUTBOUND_EMAIL=1)."
	try:
		# Reference: ChatGPT (OpenAI) - SMTP Fail-Fast Timeout Guard
		# Date: 2026-02-12
		# Prompt: "My Flask request can timeout when SMTP host is slow/unreachable.
		# I need a fail-fast timeout strategy so email attempts return quickly instead
		# of blocking the web worker. Can you provide a practical pattern?"
		# ChatGPT recommended using a short, bounded timeout for all SMTP socket ops.
		timeout_raw = (os.getenv("SMTP_TIMEOUT_SECONDS", "4") or "4").strip()
		try:
			smtp_timeout = float(timeout_raw)
		except ValueError:
			smtp_timeout = 4.0
		smtp_timeout = max(2.0, min(smtp_timeout, 8.0))

		message = EmailMessage()
		message["Subject"] = subject
		message["From"] = smtp_config.from_email
		message["To"] = to_email
		message.set_content(body)

		with smtplib.SMTP(smtp_config.host, smtp_config.port, timeout=smtp_timeout) as server:
			# Ensure TLS/login/send use the same bounded timeout.
			server.sock.settimeout(smtp_timeout)
			if smtp_config.use_tls:
				server.starttls()
				if server.sock:
					server.sock.settimeout(smtp_timeout)
			if smtp_config.username and smtp_config.password:
				server.login(smtp_config.username, smtp_config.password)
			server.send_message(message)
		return None
	except Exception as exc:
		return str(exc)


def _fetch_lecturer_replies(student_id: int) -> Dict[str, Any]:
	"""
	Fetch emails from known lecturer addresses via IMAP.
	Returns dict with 'success', 'new_count', 'error' keys.
	"""
	import imaplib
	import email
	from email.header import decode_header
	from email.utils import parsedate_to_datetime
	
	result = {"success": False, "new_count": 0, "error": None}
	
	imap_config = get_imap_config()
	if not imap_config:
		result["error"] = "IMAP is not configured. Add IMAP_HOST, IMAP_USERNAME, IMAP_PASSWORD to your environment."
		return result
	
	# Get list of known lecturer emails for this student
	try:
		lecturers = sb_fetch_all(
			"SELECT id, email, name FROM lecturers",
			{}
		)
	except Exception as exc:
		result["error"] = f"Failed to fetch lecturers: {exc}"
		return result
	
	if not lecturers:
		result["error"] = "No lecturers found in database."
		return result
	
	lecturer_emails = {l["email"].lower(): l for l in lecturers if l.get("email")}
	
	if not lecturer_emails:
		result["error"] = "No lecturer email addresses found."
		return result
	
	try:
		# Connect to IMAP server
		if imap_config.use_ssl:
			mail = imaplib.IMAP4_SSL(imap_config.host, imap_config.port)
		else:
			mail = imaplib.IMAP4(imap_config.host, imap_config.port)
		
		mail.login(imap_config.username, imap_config.password)
		mail.select(imap_config.folder)
		
		new_count = 0
		
		# Search for emails from each lecturer
		for lecturer_email, lecturer_info in lecturer_emails.items():
			# Search for emails from this lecturer
			status, messages = mail.search(None, f'FROM "{lecturer_email}"')
			
			if status != "OK":
				continue
			
			email_ids = messages[0].split()
			
			# Process each email (limit to last 20 per lecturer)
			for email_id in email_ids[-20:]:
				status, msg_data = mail.fetch(email_id, "(RFC822)")
				
				if status != "OK":
					continue
				
				for response_part in msg_data:
					if isinstance(response_part, tuple):
						msg = email.message_from_bytes(response_part[1])
						
						# Get Message-ID to check for duplicates
						message_id = msg.get("Message-ID", "")
						if not message_id:
							message_id = f"{email_id.decode()}-{lecturer_email}"
						
						# Check if we already have this message
						existing = sb_fetch_one(
							"SELECT id FROM lecturer_replies WHERE message_id = :message_id",
							{"message_id": message_id}
						)
						
						if existing:
							continue  # Skip duplicate
						
						# Decode subject
						subject = msg.get("Subject", "")
						if subject:
							decoded_parts = decode_header(subject)
							subject = ""
							for part, encoding in decoded_parts:
								if isinstance(part, bytes):
									subject += part.decode(encoding or "utf-8", errors="replace")
								else:
									subject += part
						
						# Get sender name
						from_header = msg.get("From", "")
						from_name = lecturer_info.get("name", "")
						
						# Parse date
						date_str = msg.get("Date", "")
						received_at = None
						if date_str:
							try:
								received_at = parsedate_to_datetime(date_str)
							except Exception:
								received_at = datetime.now()
						
						# Get body
						body = ""
						if msg.is_multipart():
							for part in msg.walk():
								content_type = part.get_content_type()
								if content_type == "text/plain":
									try:
										payload = part.get_payload(decode=True)
										charset = part.get_content_charset() or "utf-8"
										body = payload.decode(charset, errors="replace")
										break
									except Exception:
										pass
						else:
							try:
								payload = msg.get_payload(decode=True)
								charset = msg.get_content_charset() or "utf-8"
								body = payload.decode(charset, errors="replace")
							except Exception:
								body = str(msg.get_payload())
						
						# Store in database
						try:
							sb_execute(
								"""
								INSERT INTO lecturer_replies (
									student_id, lecturer_id, from_email, from_name,
									subject, body, received_at, message_id, is_read
								) VALUES (
									:student_id, :lecturer_id, :from_email, :from_name,
									:subject, :body, :received_at, :message_id, FALSE
								)
								""",
								{
									"student_id": student_id,
									"lecturer_id": lecturer_info.get("id"),
									"from_email": lecturer_email,
									"from_name": from_name,
									"subject": subject[:500] if subject else None,
									"body": body,
									"received_at": received_at,
									"message_id": message_id,
								}
							)
							new_count += 1
						except Exception as db_err:
							print(f"[imap] failed to store reply: {db_err}")
		
		mail.logout()
		result["success"] = True
		result["new_count"] = new_count
		
	except imaplib.IMAP4.error as imap_err:
		result["error"] = f"IMAP error: {imap_err}"
	except Exception as exc:
		result["error"] = f"Failed to fetch emails: {exc}"
	
	return result


def _build_daily_summary_email(student_id: int) -> Optional[str]:
	# Reference: ChatGPT (OpenAI) - Daily Summary Grouping
	# Date: 2026-01-22
	# Prompt: "Can you help me build a daily summary email that groups tasks into overdue,
	# due in 24h, and due in 72h?"
	# ChatGPT provided the grouping and formatting pattern.
	try:
		rows = sb_fetch_all(
			"""
			SELECT t.id, t.title, t.status, t.due_date, t.due_at,
			       m.code AS module_code
			FROM tasks t
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			  AND t.status IN ('pending', 'in_progress')
			""",
			{"student_id": student_id}
		)
	except Exception as exc:
		print(f"[daily-summary] failed to load tasks user={student_id} error={exc}")
		return None

	now = datetime.now(timezone.utc)
	overdue = []
	due_24h = []
	due_72h = []
	for task in rows:
		due_at = normalise_due_datetime(task, now)
		if not due_at:
			continue
		hours_remaining = (due_at - now).total_seconds() / 3600
		entry = f"- {task.get('title')} ({task.get('module_code') or 'Module'}) due {due_at.strftime('%d/%m/%Y %H:%M')}"
		if hours_remaining <= 0:
			overdue.append(entry)
		elif hours_remaining <= 24:
			due_24h.append(entry)
		elif hours_remaining <= 72:
			due_72h.append(entry)

	if not (overdue or due_24h or due_72h):
		return None

	lines = ["Your daily study summary:", ""]
	if overdue:
		lines.append("Overdue:")
		lines.extend(overdue[:5])
		lines.append("")
	if due_24h:
		lines.append("Due within 24h:")
		lines.extend(due_24h[:5])
		lines.append("")
	if due_72h:
		lines.append("Due within 72h:")
		lines.extend(due_72h[:5])
		lines.append("")
	lines.append("Open your dashboard to plan micro-tasks and stay on track.")
	return "\n".join(lines).strip()


def _generate_task_reminders(student_id: int) -> None:
	try:
		rows = sb_fetch_all(
			"""
			SELECT t.id, t.title, t.status, t.due_date, t.due_at,
			       m.code AS module_code
			FROM tasks t
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			  AND t.status IN ('pending', 'in_progress')
			""",
			{"student_id": student_id}
		)
	except Exception as exc:
		print(f"[reminders] failed to load tasks user={student_id} error={exc}")
		return

	student_email = None
	email_notifications_enabled = True
	email_daily_summary_enabled = True
	try:
		student = sb_fetch_one(
			"""
			SELECT email, email_notifications_enabled, email_daily_summary_enabled
			FROM students
			WHERE id = :student_id
			""",
			{"student_id": student_id}
		)
		student_email = (student or {}).get("email")
		email_notifications_enabled = bool((student or {}).get("email_notifications_enabled", True))
		email_daily_summary_enabled = bool((student or {}).get("email_daily_summary_enabled", True))
	except Exception as exc:
		print(f"[reminders] failed to load student email user={student_id} error={exc}")

	now = datetime.now(timezone.utc)
	for task in rows:
		due_at = normalise_due_datetime(task, now)
		if not due_at:
			continue
		hours_remaining = (due_at - now).total_seconds() / 3600
		if hours_remaining <= 0:
			reminder_type = "overdue"
		elif hours_remaining <= 24:
			reminder_type = "due_24h"
		elif hours_remaining <= 72:
			reminder_type = "due_72h"
		else:
			continue

		message = _build_reminder_message(task, reminder_type)
		try:
			inserted = sb_execute(
				"""
				INSERT INTO reminders (
					student_id, task_id, reminder_type, message, due_at, created_at, is_read
				) VALUES (
					:student_id, :task_id, :reminder_type, :message, :due_at, NOW(), FALSE
				)
				ON CONFLICT (student_id, task_id, reminder_type) DO NOTHING
				""",
				{
					"student_id": student_id,
					"task_id": task.get("id"),
					"reminder_type": reminder_type,
					"message": message,
					"due_at": due_at,
				}
			)
			if inserted and student_email and email_notifications_enabled:
				subject = message
				body = (
					f"{message}\n"
					f"Due: {due_at.strftime('%d/%m/%Y %H:%M')}\n\n"
					"Open your dashboard to view task details."
				)
				error = _send_reminder_email(
					to_email=student_email,
					subject=subject,
					body=body
				)
				status = "sent" if error is None else "failed"
				sb_execute(
					"""
					UPDATE reminders
					SET email_sent_at = NOW(),
					    email_status = :status,
					    email_error = :error
					WHERE student_id = :student_id
					  AND task_id = :task_id
					  AND reminder_type = :reminder_type
					""",
					{
						"status": status,
						"error": error,
						"student_id": student_id,
						"task_id": task.get("id"),
						"reminder_type": reminder_type,
					}
				)
		except Exception as exc:
			print(f"[reminders] insert failed user={student_id} task={task.get('id')} err={exc}")


def _extract_query_terms(question: str) -> List[str]:
	import re
	stopwords = {
		"the", "and", "for", "with", "that", "this", "from", "into", "onto", "over", "under",
		"about", "what", "when", "where", "which", "who", "whom", "why", "how", "can", "could",
		"should", "would", "will", "may", "might", "must", "is", "are", "was", "were", "be",
		"been", "being", "a", "an", "of", "to", "in", "on", "at", "by", "as", "it", "its",
		"your", "our", "their", "they", "we", "you", "i", "me", "my", "mine", "ours"
	}
	tokens = re.findall(r"[a-z0-9']+", question.lower())
	return [t for t in tokens if len(t) > 2 and t not in stopwords]


def _score_text_for_query(text: str, terms: List[str]) -> int:
	if not text or not terms:
		return 0
	text_lower = text.lower()
	return sum(text_lower.count(term) for term in terms)


def _select_course_materials(
	documents: List[Dict[str, Any]],
	question: str,
	max_docs: int = 3,
	max_chars: int = 6000
) -> List[Dict[str, str]]:
	terms = _extract_query_terms(question)
	scored = []
	for doc in documents:
		text = (doc.get("extracted_text") or "").strip()
		score = _score_text_for_query(text, terms)
		uploaded = doc.get("uploaded_at") or datetime.min
		scored.append((score, uploaded, doc))
	scored.sort(key=lambda item: (item[0], item[1]), reverse=True)

	selected: List[Dict[str, str]] = []
	total_chars = 0
	for score, _uploaded, doc in scored:
		if len(selected) >= max_docs:
			break
		raw_text = (doc.get("extracted_text") or "").strip()
		if not raw_text:
			continue
		clean = " ".join(raw_text.split())
		remaining = max_chars - total_chars
		if remaining <= 0:
			break
		snippet = clean[:min(2500, remaining)]
		source_name = doc.get("filename") or f"Document {doc.get('id')}"
		selected.append({"source": source_name, "content": snippet})
		total_chars += len(snippet)
	return selected


# Reference: ChatGPT (OpenAI) - Prompt Text Sanitization
# Date: 2025-11-14
# Prompt: "I need to sanitize assignment briefs before sending them to AI. I need to remove 
# sensitive words (plagiarism, harassment, etc.) and replace them with safer alternatives. 
# Also strip content after certain markers like 'Rubric:' or 'Required Format:'. Can you 
# help me write this sanitization function?"
# ChatGPT provided the text sanitization logic that replaces sensitive keywords, strips 
# content after specific markers, and calls the summarization function to keep text concise. 
# This prevents AI safety blocks and keeps prompts focused.
def _sanitize_prompt_text(value: Optional[str]) -> Optional[str]:
	if not value:
		return None
	text = str(value)
	for marker in [
		"Rubric:",
		"Required Format:",
		"Referencing:",
		"Assignments should be",
		"Applications must be made",
	]:
		if marker in text:
			text = text.split(marker)[0]
			break
	replacements = {
		"plagiarism": "academic integrity",
		"harassment": "unwanted behaviour",
		"violence": "serious behaviour issue",
		"suicide": "self-care concern",
		"self-harm": "self-care concern",
	}
	for target, replacement in list(replacements.items()):
		text = text.replace(target, replacement)
		text = text.replace(target.capitalize(), replacement.capitalize())
		text = text.replace(target.upper(), replacement.upper())
	clean = " ".join(text.split())
	summary = _summarize_text(clean, max_len=400)
	return summary


# Reference: ChatGPT (OpenAI) - Multi-Format File Text Extraction
# Date: 2025-11-14
# Prompt: "I need a function that extracts text from multiple file formats: plain text (.txt, 
# .md), DOCX files, and PDF files. It should handle encoding issues, extract text from DOCX 
# paragraphs and tables, and extract text from PDF pages. Can you help me write this with 
# proper error handling?"
# ChatGPT provided the multi-format text extraction logic with encoding fallbacks for text 
# files, paragraph and table extraction for DOCX files, and page-by-page extraction for PDFs. 
# Includes comprehensive error handling and validation.
def _extract_text_from_brief(filename: str, payload: bytes) -> str:
	suffix = Path(filename or "").suffix.lower()
	if suffix not in _AI_ALLOWED_SUFFIXES:
		allowed = ", ".join(sorted(_AI_ALLOWED_SUFFIXES))
		raise ValueError(f"Unsupported file type '{suffix or 'unknown'}'. Upload one of: {allowed}.")
	if suffix in {".txt", ".md", ".markdown"}:
		for encoding in ("utf-8", "utf-16", "latin-1"):
			try:
				return payload.decode(encoding)
			except UnicodeDecodeError:
				continue
		return payload.decode("utf-8", errors="ignore")
	if suffix == ".docx":
		try:
			from docx import Document  # type: ignore
		except ImportError as exc:  # pragma: no cover - dependency guard
			raise ValueError("python-docx is required to read DOCX files. Install python-docx.") from exc
		try:
			document = Document(BytesIO(payload))
			# Extract text from all paragraphs
			lines = []
			for paragraph in document.paragraphs:
				if paragraph.text and paragraph.text.strip():
					lines.append(paragraph.text.strip())
			# Also try to extract text from tables
			for table in document.tables:
				for row in table.rows:
					for cell in row.cells:
						if cell.text and cell.text.strip():
							lines.append(cell.text.strip())
			result = "\n".join(lines).strip()
			if not result:
				raise ValueError("The DOCX file appears to be empty or contains only formatting.")
			return result
		except ValueError:
			raise  # Re-raise ValueError as-is
		except Exception as exc:
			print(f"[docx-extract] Error reading DOCX file: {exc}")
			import traceback
			traceback.print_exc()
			raise ValueError(f"Could not read the DOCX file: {str(exc)}. Ensure it is not password protected or corrupted.") from exc
	if suffix == ".pdf":
		try:
			from pypdf import PdfReader  # type: ignore
		except ImportError as exc:  # pragma: no cover
			raise ValueError("pypdf is required to read PDF files. Install pypdf.") from exc
		try:
			reader = PdfReader(BytesIO(payload))
			pages = []
			for page in reader.pages:
				text = page.extract_text() or ""
				text = text.strip()
				if text:
					pages.append(text)
			content = "\n\n".join(pages).strip()
			if not content:
				raise ValueError("The PDF did not contain extractable text.")
			return content
		except ValueError:
			raise
		except Exception as exc:
			raise ValueError("Could not read the PDF file. Ensure it is not locked or scanned.") from exc
	return ""


def _format_future_date(value) -> str:
	if value is None:
		return "unscheduled"
	try:
		if isinstance(value, datetime):
			return value.strftime("%d %b %Y %H:%M")
		return value.strftime("%d %b %Y")  # date
	except Exception:
		return str(value)


# Reference: ChatGPT (OpenAI) - Calendar Context Building for AI Prompts
# Date: 2025-11-14
# Prompt: "I need to build a summary of a student's upcoming assignments and calendar events 
# to include in the AI prompt. This helps the AI suggest realistic timing for new task breakdowns 
# that don't conflict with existing commitments. Can you help me format this context information?"
# ChatGPT provided the structure for building a context summary that includes upcoming assignments 
# with due dates, weights, and status, plus calendar events with locations. This context is 
# included in AI prompts to enable smarter scheduling suggestions.
def _build_schedule_context(student_id: int) -> str:
	"""Summarise upcoming assignments and events so the AI can propose realistic timings."""
	lines: List[str] = []
	try:
		upcoming_tasks = sb_fetch_all(
			"""
			SELECT t.title,
			       t.due_date,
			       t.due_at,
			       t.status,
			       t.weight_percentage,
			       m.code AS module_code
			FROM tasks t
			JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			  AND (t.status IS NULL OR t.status <> 'completed')
			  AND (t.due_date IS NULL OR t.due_date >= CURRENT_DATE)
			ORDER BY t.due_date ASC NULLS LAST
			LIMIT 5
			""",
			{"student_id": student_id}
		)
	except Exception as exc:
		print(f"[ai-workspace] schedule task fetch failed user={student_id} err={exc}")
		upcoming_tasks = []
	if upcoming_tasks:
		lines.append("Upcoming assignments:")
		for task in upcoming_tasks:
			title = task.get("title") or "Assignment"
			module = task.get("module_code") or "Module"
			due_date = _format_future_date(task.get("due_at") or task.get("due_date"))
			status = task.get("status") or "pending"
			weight = task.get("weight_percentage")
			weight_str = f" · weight {weight:.0f}%" if weight else ""
			lines.append(f"- {title} ({module}) due {due_date} · status {status}{weight_str}")
	else:
		lines.append("No other upcoming assignments found in the database.")
	try:
		upcoming_events = sb_fetch_all(
			"""
			SELECT title, start_at, end_at, location
			FROM events
			WHERE student_id = :student_id
			  AND start_at >= NOW()
			ORDER BY start_at ASC
			LIMIT 4
			""",
			{"student_id": student_id}
		)
	except Exception as exc:
		print(f"[ai-workspace] schedule events fetch failed user={student_id} err={exc}")
		upcoming_events = []
	if upcoming_events:
		lines.append("Calendar events:")
		for event in upcoming_events:
			title = event.get("title") or "Event"
			start = _format_future_date(event.get("start_at"))
			location = event.get("location")
			loc = f" @ {location}" if location else ""
			lines.append(f"- {title} on {start}{loc}")
	else:
		lines.append("No upcoming events in the calendar.")
	return "\n".join(lines)


# Reference: ChatGPT (OpenAI) - Datetime Rounding to Nearest Half Hour
# Date: 2025-11-14
# Prompt: "I need a function that rounds a datetime to the nearest hour or 30-minute mark, 
# always rounding up. So 2:15 PM becomes 2:30 PM, 2:45 PM becomes 3:00 PM. Can you help me 
# write this rounding logic?"
# ChatGPT provided the datetime rounding algorithm that rounds to the nearest hour or 
# 30-minute mark, always rounding up. This ensures calendar events start at clean time 
# boundaries (hour or half-hour marks).
def _round_to_nearest_half_hour(dt: datetime) -> datetime:
	"""Round datetime to nearest hour or 30-minute mark, always rounding up"""
	minutes = dt.minute
	seconds = dt.second
	microseconds = dt.microsecond
	
	if minutes < 15 and seconds == 0 and microseconds == 0:
		return dt.replace(second=0, microsecond=0)
	elif minutes < 15:
		return dt.replace(minute=30, second=0, microsecond=0)
	elif minutes < 45:
		return dt.replace(minute=30, second=0, microsecond=0)
	else:
		# Round up to next hour
		return (dt.replace(minute=0, second=0, microsecond=0) + timedelta(hours=1))


def _ensure_datetime(value, tzinfo) -> Optional[datetime]:
	if value is None:
		return None
	if isinstance(value, datetime):
		if value.tzinfo is None:
			return value.replace(tzinfo=tzinfo)
		return value
	if isinstance(value, str):
		try:
			dt = datetime.fromisoformat(value)
		except ValueError:
			return None
		if dt.tzinfo is None:
			return dt.replace(tzinfo=tzinfo)
		return dt
	return None


# Reference: ChatGPT (OpenAI) - Interval Arithmetic for Time Slots
# Date: 2025-11-14
# Prompt: "I need a function that subtracts a busy time interval from a list of free time segments. 
# If a busy interval overlaps with a free segment, I need to split the free segment into parts 
# that don't overlap. Can you help me write this interval subtraction logic?"
# ChatGPT provided the algorithm for subtracting busy intervals from free time segments by 
# checking overlap conditions and splitting segments accordingly. This is used to find available 
# time slots after accounting for existing calendar events.
def _subtract_interval(segments: List[Tuple[datetime, datetime]], busy: Tuple[datetime, datetime]) -> List[Tuple[datetime, datetime]]:
	result: List[Tuple[datetime, datetime]] = []
	busy_start, busy_end = busy
	for start, end in segments:
		if busy_end <= start or busy_start >= end:
			result.append((start, end))
			continue
		if busy_start > start:
			result.append((start, busy_start))
		if busy_end < end:
			result.append((busy_end, end))
	return [(s, e) for s, e in result if s < e]


# Reference: ChatGPT (OpenAI) - Free Time Slot Generation Algorithm
# Date: 2025-11-14
# Prompt: "I need to find free time slots in a student's calendar. I have a working day window 
# (9 AM to 9 PM), existing calendar events (busy intervals), and I need to find all available 
# 30+ minute slots over the next few weeks. Can you help me write an algorithm that subtracts 
# busy intervals from working hours and returns free slots?"
# ChatGPT provided the algorithm for generating free time slots by defining a working window 
# (9 AM - 9 PM), iterating through each day, and subtracting busy intervals from each day's 
# working hours using interval arithmetic. This ensures AI-generated tasks can be scheduled 
# in available time.
def _generate_free_slots(student_id: int, tzinfo, horizon_days: int = 21) -> List[Tuple[datetime, datetime]]:
	now = datetime.now(tzinfo)
	start_date = now.date()
	try:
		rows = sb_fetch_all(
			"""
			SELECT start_at, end_at
			FROM events
			WHERE student_id = :student_id
			  AND start_at IS NOT NULL
			  AND end_at IS NOT NULL
			  AND start_at >= NOW()
			  AND start_at <= NOW() + interval '30 days'
			""",
			{"student_id": student_id}
		)
	except Exception:
		rows = []
	busy_intervals: List[Tuple[datetime, datetime]] = []
	for row in rows:
		start = _ensure_datetime(row.get("start_at"), tzinfo)
		end = _ensure_datetime(row.get("end_at"), tzinfo)
		if start and end and end > now:
			busy_intervals.append((start, end))
	busy_intervals.sort(key=lambda item: item[0])
	free_slots: List[Tuple[datetime, datetime]] = []
	work_start = time(hour=9)
	work_end = time(hour=21)
	for day_offset in range(horizon_days):
		day = start_date + timedelta(days=day_offset)
		day_start = datetime.combine(day, work_start, tzinfo)
		day_end = datetime.combine(day, work_end, tzinfo)
		if day_end <= now:
			continue
		if day_start < now:
			day_start = now
		segments = [(day_start, day_end)]
		for busy in busy_intervals:
			if busy[0].date() != day:
				continue
			segments = _subtract_interval(segments, busy)
			if not segments:
				break
		for segment in segments:
			if segment[1] - segment[0] >= timedelta(minutes=30):
				free_slots.append(segment)
	return free_slots


def _parse_due_datetime(due_date_str: Optional[str], tzinfo) -> Optional[datetime]:
	if not due_date_str:
		return None
	try:
		dt = datetime.strptime(due_date_str, "%Y-%m-%d")
	except ValueError:
		return None
	return datetime.combine(dt.date(), time(hour=23, minute=59), tzinfo)


def _lookup_module_id(module_code: Optional[str], student_id: int) -> Optional[int]:
	if not module_code:
		return None
	try:
		row = sb_fetch_one(
			"""
			SELECT id FROM modules
			WHERE code ILIKE :code
			LIMIT 1
			""",
			{"code": module_code}
		)
	except Exception:
		row = None
	if row:
		return row.get("id")
	return None


# Reference: ChatGPT (OpenAI) - Task Lookup and Creation Pattern
# Date: 2025-11-14
# Prompt: "I need a function that looks up an existing task by title and module, and if 
# not found, creates a new task. It should handle both tasks with and without modules, 
# parse due dates from strings, and set default due times. Can you help me write this?"
# ChatGPT provided the lookup-or-create pattern for tasks. It searches by title and module 
# (or NULL module), handles date parsing with multiple formats, sets default due times, and 
# returns the task ID for linking subtasks.
def _lookup_or_create_task(assignment_title: str, module_id: Optional[int], module_code: Optional[str], 
                           student_id: int, due_date_str: Optional[str]) -> Optional[int]:
	"""Look up existing task or create a new one for AI-generated subtasks"""
	if not assignment_title:
		return None
	
	# First, try to find an existing task by title and module
	try:
		if module_id:
			task = sb_fetch_one(
				"""
				SELECT id FROM tasks
				WHERE student_id = :student_id
				  AND module_id = :module_id
				  AND title ILIKE :title
				LIMIT 1
				""",
				{"student_id": student_id, "module_id": module_id, "title": assignment_title}
			)
		else:
			task = sb_fetch_one(
				"""
				SELECT id FROM tasks
				WHERE student_id = :student_id
				  AND module_id IS NULL
				  AND title ILIKE :title
				LIMIT 1
				""",
				{"student_id": student_id, "title": assignment_title}
			)
		
		if task:
			return task.get("id")
	except Exception as exc:
		print(f"[ai-workspace] task lookup failed: {exc}")
		import traceback
		traceback.print_exc()
		pass
	
	# Create new task if not found
	try:
		print(f"[ai-workspace] Creating new task: title='{assignment_title}', module_id={module_id}, due_date={due_date_str}")
		from datetime import datetime, date, timezone
		due_date = None
		due_at_value = None
		
		if due_date_str:
			try:
				# Try parsing as date first
				due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()
				# Set due_at to end of day (5 PM) if only date provided
				due_at_value = datetime.combine(due_date, datetime.min.time().replace(hour=17, minute=0))
				if due_at_value.tzinfo is None:
					due_at_value = due_at_value.replace(tzinfo=timezone.utc)
			except ValueError:
				try:
					due_at_value = datetime.fromisoformat(due_date_str.replace('Z', '+00:00'))
					due_date = due_at_value.date()
				except:
					pass
		
		sb_execute(
			"""INSERT INTO tasks (title, student_id, module_id, due_date, due_at, status) 
			   VALUES (:title, :student_id, :module_id, :due_date, :due_at, 'pending')""",
			{
				"title": assignment_title,
				"student_id": student_id,
				"module_id": module_id,
				"due_date": due_date,
				"due_at": due_at_value
			}
		)
		
		# Fetch the newly created task ID
		if module_id:
			task = sb_fetch_one(
				"""
				SELECT id FROM tasks
				WHERE student_id = :student_id
				  AND module_id = :module_id
				  AND title ILIKE :title
				LIMIT 1
				""",
				{"student_id": student_id, "module_id": module_id, "title": assignment_title}
			)
		else:
			task = sb_fetch_one(
				"""
				SELECT id FROM tasks
				WHERE student_id = :student_id
				  AND module_id IS NULL
				  AND title ILIKE :title
				LIMIT 1
				""",
				{"student_id": student_id, "title": assignment_title}
			)
		
		if task:
			print(f"[ai-workspace] Successfully created task_id={task.get('id')}")
			return task.get("id")
		else:
			print(f"[ai-workspace] WARNING: Task created but could not retrieve task_id")
	except Exception as exc:
		print(f"[ai-workspace] failed to create task: {exc}")
		import traceback
		traceback.print_exc()
		pass
	
	return None


# Reference: ChatGPT (OpenAI) - Natural Language Date/Time Parsing
# Date: 2025-11-14
# Prompt: "The AI sometimes returns time hints like 'Tuesday evening', 'next week afternoon', 
# or ISO dates. I need a function that can parse these natural language hints and convert them 
# to datetime objects. It should handle multiple date formats and extract time-of-day hints 
# (morning, afternoon, evening, night) to set appropriate hours."
# ChatGPT provided the parsing logic for multiple date formats and natural language time-of-day 
# hints. The function tries various date formats, extracts time-of-day keywords to set hours, 
# and falls back to ISO parsing. This allows the scheduling system to respect AI-suggested 
# preferred times.
def _parse_plan_hint(text: Optional[str], tzinfo) -> Optional[datetime]:
	if not text:
		return None
	value = text.strip()
	if not value:
		return None
	candidates = [
		"%Y-%m-%d %H:%M",
		"%Y-%m-%d",
		"%d %b %Y",
		"%d %B %Y",
		"%b %d %Y",
		"%B %d %Y",
	]
	lower = value.lower()
	hour = 9
	if "evening" in lower:
		hour = 19
	elif "afternoon" in lower:
		hour = 14
	elif "morning" in lower:
		hour = 10
	elif "night" in lower:
		hour = 21
	for fmt in candidates:
		try:
			dt = datetime.strptime(value.split(" evening")[0].split(" afternoon")[0].split(" morning")[0].split(" night")[0].strip(), fmt)
			if "hour" in fmt.lower():
				return dt.replace(tzinfo=tzinfo)
			return datetime.combine(dt.date(), time(hour=hour), tzinfo)
		except ValueError:
			continue
	# Try ISO parse
	try:
		dt = datetime.fromisoformat(value)
		if dt.tzinfo is None:
			dt = dt.replace(tzinfo=tzinfo)
		return dt
	except ValueError:
		return None


# Reference: ChatGPT (OpenAI) - Even Task Distribution Scheduling Algorithm
# Date: 2025-11-14
# Prompt: "I need to schedule AI-generated subtasks into calendar slots. The AI suggests preferred 
# times, but I need to distribute tasks evenly across available time until the deadline. If a task 
# can't fit in the preferred time slot, find the nearest available slot. Tasks should be spread out 
# evenly, not clustered on one day. Can you help me write this scheduling algorithm?"
# ChatGPT provided the algorithm for distributing subtasks evenly across available time slots. 
# It prioritizes AI-suggested times when possible, but evenly spreads tasks across the available 
# window until the deadline. The algorithm uses slot consumption to prevent double-booking and 
# finds nearest available slots when preferred times are unavailable.
def _schedule_ai_subtasks(
	*,
	subtasks: List[Dict[str, Any]],
	student_id: int,
	module_id: Optional[int],
	assignment_title: str,
	due_date_str: Optional[str]
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
	tzinfo = datetime.now().astimezone().tzinfo
	now = datetime.now(tzinfo)
	free_slots = _generate_free_slots(student_id, tzinfo)
	if not free_slots:
		return [], subtasks
	due_deadline = _parse_due_datetime(due_date_str, tzinfo)
	if due_deadline:
		before = [slot for slot in free_slots if slot[0] <= due_deadline]
		after = [slot for slot in free_slots if slot[0] > due_deadline]
		slot_queue = before + after
	else:
		slot_queue = free_slots[:]
	slot_queue.sort(key=lambda s: s[0])
	scheduled: List[Dict[str, Any]] = []
	unscheduled: List[Dict[str, Any]] = []

	def consume_slot(idx: int, duration: timedelta) -> Optional[Tuple[datetime, datetime]]:
		start, end = slot_queue[idx]
		# Need duration + 30 min buffer for spacing
		if end - start < duration + timedelta(minutes=30):
			return None
		
		# Round start time to nearest hour or 30 minutes, but never below slot start
		assigned_start = _round_to_nearest_half_hour(start)
		# If rounding went below slot start, round up to next half hour
		if assigned_start < start:
			# Round up to next :00 or :30
			if start.minute < 30:
				assigned_start = start.replace(minute=30, second=0, microsecond=0)
			else:
				assigned_start = (start.replace(minute=0, second=0, microsecond=0) + timedelta(hours=1))
			# If still below, just use the slot start
			if assigned_start < start:
				assigned_start = start.replace(second=0, microsecond=0)
		
		if assigned_start >= end:
			return None
		
		# Calculate end time
		assigned_end = assigned_start + duration
		# Round end time to nearest half hour, but ensure it doesn't exceed slot end
		assigned_end_rounded = _round_to_nearest_half_hour(assigned_end)
		if assigned_end_rounded > end:

			assigned_end = end.replace(second=0, microsecond=0)
		else:
			assigned_end = assigned_end_rounded
		
		# Ensure end is after start
		if assigned_end <= assigned_start:
			assigned_end = assigned_start + duration
			if assigned_end > end:
				return None
		
		# Add 30 minute buffer after event to prevent overlap (increased from 15 to 30)
		buffer_end = assigned_end + timedelta(minutes=30)
		
		if buffer_end <= end:
			slot_queue[idx] = (buffer_end, end)
		elif assigned_end < end:
			slot_queue[idx] = (assigned_end, end)
		else:
			slot_queue.pop(idx)
		return assigned_start, assigned_end

	total_subtasks = len(subtasks)
	if total_subtasks <= 0:
		return [], []

	for idx, subtask in enumerate(subtasks):
		estimated_hours = subtask.get("estimated_hours")
		try:
			duration_hours = float(estimated_hours) if estimated_hours is not None else 2.0
		except (ValueError, TypeError):
			duration_hours = 2.0
		duration_hours = max(0.5, min(duration_hours, 6.0))
		duration = timedelta(hours=duration_hours)
		target_time = _parse_plan_hint(subtask.get("planned_start"), tzinfo)
		if not target_time:
			if due_deadline:
				window = max(due_deadline - now, timedelta(days=1))
				# Spread more evenly across available time
				target_time = now + (window * (idx / max(total_subtasks - 1, 1)))
			else:
				# Spread across days: one task per day
				target_time = now + timedelta(days=idx)

		# Build candidates with preference for spreading across days
		candidate_indices = list(range(len(slot_queue)))
		# Count how many tasks are already scheduled on each day
		day_counts = {}
		for scheduled_item in scheduled:
			start_dt = scheduled_item.get("start")
			if isinstance(start_dt, datetime):
				day_key = start_dt.date()
				day_counts[day_key] = day_counts.get(day_key, 0) + 1
		
		candidate_indices.sort(
			key=lambda i: (
				0 if slot_queue[i][0] >= target_time else 1,  # Prefer future slots
				day_counts.get(slot_queue[i][0].date(), 0),  # Prefer days with fewer tasks
				abs((slot_queue[i][0] - target_time).total_seconds())  # Then closest to target
			)
		)

		assignment_slot = None
		for slot_idx in candidate_indices:
			if slot_idx >= len(slot_queue):
				continue
			slot_start, slot_end = slot_queue[slot_idx]
			if slot_end - slot_start < duration + timedelta(minutes=30):
				continue
			
			# Calculate what the assigned times would be (before consuming)
			# Round start time to nearest hour or 30 minutes, but never below slot start
			potential_start = _round_to_nearest_half_hour(slot_start)
			if potential_start < slot_start:
				if slot_start.minute < 30:
					potential_start = slot_start.replace(minute=30, second=0, microsecond=0)
				else:
					potential_start = (slot_start.replace(minute=0, second=0, microsecond=0) + timedelta(hours=1))
				if potential_start < slot_start:
					potential_start = slot_start.replace(second=0, microsecond=0)
			
			if potential_start >= slot_end:
				continue
			
			potential_end = potential_start + duration
			potential_end_rounded = _round_to_nearest_half_hour(potential_end)
			if potential_end_rounded > slot_end:
				potential_end = slot_end.replace(second=0, microsecond=0)
			else:
				potential_end = potential_end_rounded
			
			if potential_end <= potential_start:
				potential_end = potential_start + duration
				if potential_end > slot_end:
					continue
			
			# Reference: ChatGPT (OpenAI) - Calendar Event Overlap Detection
			# Date: 2025-11-14
			# Prompt: "I need to check if a new calendar event would overlap with existing events, 
			# including a buffer time (30 minutes) between events. The overlap check should account 
			# for both the new event's time and the buffer. Can you help me write this overlap 
			# detection logic?"
			# ChatGPT provided the overlap detection algorithm that checks if a new event (with buffer) 
			# overlaps with existing events. It compares start and end times with buffers to prevent 
			# double-booking and ensures events have proper spacing.
			# Check for overlaps with already-scheduled events (with 30 min buffer)
			overlaps = False
			for existing in scheduled:
				existing_start = existing.get("start")
				existing_end = existing.get("end")
				if isinstance(existing_start, datetime) and isinstance(existing_end, datetime):
					# Check if new event overlaps with existing event (with buffers)
					new_start = potential_start
					new_end = potential_end + timedelta(minutes=30)  # Include buffer in check
					existing_end_buffered = existing_end + timedelta(minutes=30)
					# Overlap if: new_start < existing_end_buffered AND new_end > existing_start
					if new_start < existing_end_buffered and new_end > existing_start:
						overlaps = True
						break
			
			if overlaps:
				# Skip this slot and try next one
				continue
			
			# No overlap found, consume the slot
			assignment_slot = consume_slot(slot_idx, duration)
			if assignment_slot:
				start_ts, end_ts = assignment_slot
				scheduled.append({
					"title": f"{assignment_title}: {subtask.get('title') or 'Subtask'}",
					"start": start_ts,
					"end": end_ts,
					"focus": subtask.get("focus"),
					"module_id": module_id,
				})
				break
		
		if not assignment_slot:
			unscheduled.append(subtask)
	return scheduled, unscheduled


# Reference: ChatGPT (OpenAI) - Group Workspace Backend Flow
# Date: 2026-02-11
# Prompt: "I need a group project workspace in Flask where students can create
# projects, add members, assign tasks manually or with AI, track progress, and
# email members their assigned items. Can you provide a robust route/action
# pattern with ownership checks, safe parsing, and summary metrics?"
# ChatGPT provided the helper + route architecture below, adapted to this
# codebase's Supabase helpers and existing AI/email services.
def _parse_iso_date_value(value: Optional[str]) -> Optional[datetime.date]:
	raw = (value or "").strip()
	if not raw:
		return None
	for parser in (
		lambda x: datetime.strptime(x, "%Y-%m-%d").date(),
		lambda x: datetime.fromisoformat(x.replace("Z", "+00:00")).date(),
	):
		try:
			return parser(raw)
		except Exception:
			continue
	return None


def _coerce_int(value: Any) -> Optional[int]:
	try:
		return int(value)
	except Exception:
		return None


def _get_owned_group_project(project_id: int, student_id: int) -> Optional[Dict[str, Any]]:
	return sb_fetch_one(
		"""
		SELECT id, owner_student_id, title, module_code, description, due_date, created_at
		FROM group_projects
		WHERE id = :project_id AND owner_student_id = :student_id
		""",
		{"project_id": project_id, "student_id": student_id},
	)


def _get_group_membership_project(project_id: int, student_id: int, student_email: Optional[str]) -> Optional[Dict[str, Any]]:
	email = (student_email or "").strip().lower()
	return sb_fetch_one(
		"""
		SELECT p.id, p.owner_student_id, p.title, p.module_code, p.description, p.due_date, p.created_at
		FROM group_projects p
		JOIN group_project_members m ON m.project_id = p.id
		WHERE p.id = :project_id
		  AND m.project_id = :project_id
		  AND m.invite_status = 'accepted'
		  AND LOWER(m.member_email) = :member_email
		LIMIT 1
		""",
		{"project_id": project_id, "member_email": email},
	)


def _get_accessible_group_project(project_id: int, student_id: int, student_email: Optional[str]) -> Optional[Dict[str, Any]]:
	owned = _get_owned_group_project(project_id, student_id)
	if owned:
		return owned
	return _get_group_membership_project(project_id, student_id, student_email)


def _group_invite_url(invite_token: Optional[str]) -> Optional[str]:
	if not invite_token:
		return None
	try:
		return url_for("group_workspace_invite", invite_token=invite_token, _external=True)
	except Exception:
		return f"/group-workspace/invite/{invite_token}"


@app.route("/group-workspace", methods=["GET", "POST"])
@login_required
def group_workspace():
	selected_project_id = request.args.get("project_id")
	if request.method == "POST":
		action = (request.form.get("action") or "").strip()
		selected_project_id = request.form.get("project_id") or selected_project_id

		if action == "create_project":
			title = (request.form.get("title") or "").strip()
			module_code = (request.form.get("module_code") or "").strip().upper() or None
			description = (request.form.get("description") or "").strip() or None
			due_date = _parse_iso_date_value(request.form.get("due_date"))
			if not title:
				flash("Project title is required.", "error")
				return redirect(url_for("group_workspace"))
			try:
				sb_execute(
					"""
					INSERT INTO group_projects (
						owner_student_id, title, module_code, description, due_date, created_at, updated_at
					) VALUES (
						:owner_student_id, :title, :module_code, :description, :due_date, NOW(), NOW()
					)
					""",
					{
						"owner_student_id": current_user.id,
						"title": title[:255],
						"module_code": module_code,
						"description": description,
						"due_date": due_date,
					},
				)
				project = sb_fetch_one(
					"""
					SELECT id
					FROM group_projects
					WHERE owner_student_id = :student_id AND title = :title
					ORDER BY created_at DESC, id DESC
					LIMIT 1
					""",
					{"student_id": current_user.id, "title": title[:255]},
				)
				new_id = project.get("id") if project else None
				flash("Group project created.", "success")
				if new_id:
					return redirect(url_for("group_workspace", project_id=new_id))
			except Exception as exc:
				print(f"[group-workspace] create project failed user={current_user.id} err={exc}")
				flash("Failed to create project.", "error")
			return redirect(url_for("group_workspace"))

		project_id = _coerce_int(selected_project_id)
		if not project_id:
			flash("Please select a project first.", "warning")
			return redirect(url_for("group_workspace"))
		project = _get_accessible_group_project(project_id, current_user.id, current_user.email)
		if not project:
			flash("Project not found or permission denied.", "error")
			return redirect(url_for("group_workspace"))
		can_manage_project = int(project.get("owner_student_id") or 0) == int(current_user.id)
		owner_only_actions = {
			"add_member",
			"add_task",
			"generate_ai_tasks",
			"update_task",
			"delete_task",
			"email_task",
			"email_all_assignments",
			"resend_invite",
			"add_milestone",
			"toggle_milestone",
			"upload_task_file",
			"delete_task_file",
			"upload_project_file",
			"delete_project_file",
			"post_message",
		}
		if action in owner_only_actions and not can_manage_project:
			flash("This project is view-only for members. Use the invite portal to update your assigned tasks.", "warning")
			return redirect(url_for("group_workspace", project_id=project_id))

		if action == "add_member":
			member_name = (request.form.get("member_name") or "").strip()
			member_email = (request.form.get("member_email") or "").strip()
			member_role = (request.form.get("member_role") or "").strip() or None
			if not member_name or not member_email:
				flash("Member name and email are required.", "error")
			else:
				try:
					invite_token = secrets.token_urlsafe(24)
					sb_execute(
						"""
						INSERT INTO group_project_members (
							project_id, member_name, member_email, member_role, notes,
							invite_token, invite_status, created_at
						) VALUES (
							:project_id, :member_name, :member_email, :member_role, :notes,
							:invite_token, :invite_status, NOW()
						)
						""",
						{
							"project_id": project_id,
							"member_name": member_name[:255],
							"member_email": member_email[:255],
							"member_role": member_role,
							"notes": None,
							"invite_token": invite_token,
							"invite_status": "pending",
						},
					)
					flash(f"Added {member_name} to the project.", "success")
				except Exception as exc:
					print(f"[group-workspace] add member failed user={current_user.id} project={project_id} err={exc}")
					flash("Failed to add member.", "error")

		elif action == "add_task":
			title = (request.form.get("task_title") or "").strip()
			description = (request.form.get("task_description") or "").strip() or None
			status = (request.form.get("task_status") or "todo").strip()
			priority = (request.form.get("task_priority") or "medium").strip()
			due_date = _parse_iso_date_value(request.form.get("task_due_date"))
			estimated_hours_raw = (request.form.get("task_estimated_hours") or "").strip()
			assigned_member_id = _coerce_int(request.form.get("assigned_member_id"))
			estimated_hours = None
			if estimated_hours_raw:
				try:
					estimated_hours = float(estimated_hours_raw)
				except ValueError:
					estimated_hours = None
			if status not in {"todo", "in_progress", "review", "done"}:
				status = "todo"
			if priority not in {"low", "medium", "high"}:
				priority = "medium"
			if assigned_member_id:
				member_exists = sb_fetch_one(
					"SELECT id FROM group_project_members WHERE id = :id AND project_id = :project_id",
					{"id": assigned_member_id, "project_id": project_id},
				)
				if not member_exists:
					assigned_member_id = None
			if not title:
				flash("Task title is required.", "error")
			else:
				try:
					sb_execute(
						"""
						INSERT INTO group_project_tasks (
							project_id, title, description, assigned_member_id, status, priority,
							due_date, estimated_hours, progress_percent, ai_generated,
							created_by_student_id, created_at, updated_at
						) VALUES (
							:project_id, :title, :description, :assigned_member_id, :status, :priority,
							:due_date, :estimated_hours, :progress_percent, :ai_generated,
							:created_by_student_id, NOW(), NOW()
						)
						""",
						{
							"project_id": project_id,
							"title": title[:255],
							"description": description,
							"assigned_member_id": assigned_member_id,
							"status": status,
							"priority": priority,
							"due_date": due_date,
							"estimated_hours": estimated_hours,
							"progress_percent": 100 if status == "done" else 0,
							"ai_generated": False,
							"created_by_student_id": current_user.id,
						},
					)
					flash("Task added to group project.", "success")
				except Exception as exc:
					print(f"[group-workspace] add task failed user={current_user.id} project={project_id} err={exc}")
					flash("Failed to add task.", "error")

		elif action == "generate_ai_tasks":
			brief = (request.form.get("ai_brief") or "").strip()
			task_count = _coerce_int(request.form.get("ai_task_count")) or 6
			task_count = max(3, min(15, task_count))
			replace_ai = request.form.get("replace_ai_tasks") == "on"
			if not brief:
				flash("Provide a project brief for AI breakdown.", "error")
			else:
				try:
					members = sb_fetch_all(
						"""
						SELECT id, member_name
						FROM group_project_members
						WHERE project_id = :project_id
						ORDER BY member_name ASC
						""",
						{"project_id": project_id},
					)
					member_labels = ", ".join(m.get("member_name") for m in members if m.get("member_name")) or "No members yet"
					service = get_chatgpt_service()
					breakdown = service.breakdown_task(
						task_title=f"{project.get('title')} (Group Project)",
						module_code=project.get("module_code"),
						due_date=project.get("due_date").isoformat() if getattr(project.get("due_date"), "isoformat", None) else None,
						due_at=None,
						status="in_progress",
						description=brief,
						additional_context=(
							f"Group members: {member_labels}. "
							f"Return {task_count} concrete tasks suitable for assignment."
						),
						schedule_context=None,
					)
					if replace_ai:
						sb_execute(
							"DELETE FROM group_project_tasks WHERE project_id = :project_id AND ai_generated = TRUE",
							{"project_id": project_id},
						)

					created = 0
					for idx, item in enumerate(breakdown.subtasks[:task_count], start=1):
						assigned_member_id = None
						if members:
							assigned_member_id = members[(idx - 1) % len(members)].get("id")
						due_hint = _parse_iso_date_value(item.planned_end) or _parse_iso_date_value(item.planned_start)
						if not due_hint and project.get("due_date"):
							due_hint = project.get("due_date")
						sb_execute(
							"""
							INSERT INTO group_project_tasks (
								project_id, title, description, assigned_member_id, status, priority,
								due_date, estimated_hours, progress_percent, ai_generated,
								created_by_student_id, created_at, updated_at
							) VALUES (
								:project_id, :title, :description, :assigned_member_id, :status, :priority,
								:due_date, :estimated_hours, :progress_percent, :ai_generated,
								:created_by_student_id, NOW(), NOW()
							)
							""",
							{
								"project_id": project_id,
								"title": (item.title or f"Task {idx}")[:255],
								"description": item.description or None,
								"assigned_member_id": assigned_member_id,
								"status": "todo",
								"priority": "medium",
								"due_date": due_hint,
								"estimated_hours": item.estimated_hours,
								"progress_percent": 0,
								"ai_generated": True,
								"created_by_student_id": current_user.id,
							},
						)
						created += 1
					flash(f"AI generated {created} project tasks.", "success")
				except ChatGPTClientError as exc:
					flash(f"AI breakdown failed: {exc}", "error")
				except Exception as exc:
					print(f"[group-workspace] ai breakdown failed user={current_user.id} project={project_id} err={exc}")
					flash("Failed to generate AI tasks.", "error")

		elif action == "update_task":
			task_id = _coerce_int(request.form.get("task_id"))
			if not task_id:
				flash("Invalid task.", "error")
			else:
				task = sb_fetch_one(
					"SELECT id FROM group_project_tasks WHERE id = :task_id AND project_id = :project_id",
					{"task_id": task_id, "project_id": project_id},
				)
				if not task:
					flash("Task not found.", "error")
				else:
					status = (request.form.get("status") or "todo").strip()
					if status not in {"todo", "in_progress", "review", "done"}:
						status = "todo"
					progress_raw = _coerce_int(request.form.get("progress_percent"))
					progress_percent = max(0, min(100, progress_raw if progress_raw is not None else 0))
					if status == "done":
						progress_percent = 100
					assigned_member_id = _coerce_int(request.form.get("assigned_member_id"))
					if assigned_member_id:
						member_exists = sb_fetch_one(
							"SELECT id FROM group_project_members WHERE id = :id AND project_id = :project_id",
							{"id": assigned_member_id, "project_id": project_id},
						)
						if not member_exists:
							assigned_member_id = None
					due_date = _parse_iso_date_value(request.form.get("due_date"))
					try:
						sb_execute(
							"""
							UPDATE group_project_tasks
							SET assigned_member_id = :assigned_member_id,
							    status = :status,
							    progress_percent = :progress_percent,
							    due_date = :due_date,
							    updated_at = NOW()
							WHERE id = :task_id AND project_id = :project_id
							""",
							{
								"assigned_member_id": assigned_member_id,
								"status": status,
								"progress_percent": progress_percent,
								"due_date": due_date,
								"task_id": task_id,
								"project_id": project_id,
							},
						)
						flash("Task updated.", "success")
					except Exception as exc:
						print(f"[group-workspace] update task failed user={current_user.id} task={task_id} err={exc}")
						flash("Failed to update task.", "error")

		elif action == "delete_task":
			task_id = _coerce_int(request.form.get("task_id"))
			if task_id:
				try:
					sb_execute(
						"DELETE FROM group_project_tasks WHERE id = :task_id AND project_id = :project_id",
						{"task_id": task_id, "project_id": project_id},
					)
					flash("Task deleted.", "success")
				except Exception as exc:
					print(f"[group-workspace] delete task failed user={current_user.id} task={task_id} err={exc}")
					flash("Failed to delete task.", "error")

		elif action == "email_task":
			task_id = _coerce_int(request.form.get("task_id"))
			task = None
			if task_id:
				task = sb_fetch_one(
					"""
					SELECT t.id, t.title, t.description, t.status, t.progress_percent, t.due_date,
					       m.member_name, m.member_email
					FROM group_project_tasks t
					LEFT JOIN group_project_members m ON m.id = t.assigned_member_id
					WHERE t.id = :task_id AND t.project_id = :project_id
					""",
					{"task_id": task_id, "project_id": project_id},
				)
			if not task:
				flash("Task not found.", "error")
			elif not task.get("member_email"):
				flash("Assign this task to a member with an email first.", "warning")
			else:
				subject = f"[Group Project] {project.get('title')} - Your assigned task"
				due_label = task.get("due_date").isoformat() if getattr(task.get("due_date"), "isoformat", None) else "No due date set"
				body = (
					f"Hi {task.get('member_name') or 'team member'},\n\n"
					f"You have been assigned a task in '{project.get('title')}'.\n\n"
					f"Task: {task.get('title')}\n"
					f"Status: {task.get('status')}\n"
					f"Progress: {task.get('progress_percent') or 0}%\n"
					f"Due date: {due_label}\n\n"
					f"Details:\n{task.get('description') or 'No additional notes.'}\n\n"
					f"Sent by: {current_user.name} ({current_user.email})"
				)
				error = _send_reminder_email(
					to_email=task.get("member_email"),
					subject=subject,
					body=body,
				)
				if error:
					flash(f"Failed to email {task.get('member_name')}: {error}", "error")
				else:
					flash(f"Email sent to {task.get('member_name')}.", "success")

		elif action == "email_all_assignments":
			rows = sb_fetch_all(
				"""
				SELECT m.id AS member_id, m.member_name, m.member_email,
				       t.title, t.status, t.progress_percent, t.due_date
				FROM group_project_members m
				LEFT JOIN group_project_tasks t
				       ON t.assigned_member_id = m.id
				      AND t.project_id = m.project_id
				      AND t.status != 'done'
				WHERE m.project_id = :project_id
				ORDER BY m.member_name ASC, t.due_date ASC NULLS LAST, t.id ASC
				""",
				{"project_id": project_id},
			)
			messages_sent = 0
			errors = 0
			by_member: Dict[int, Dict[str, Any]] = {}
			for row in rows:
				member_id = row.get("member_id")
				if not member_id:
					continue
				bucket = by_member.setdefault(member_id, {
					"name": row.get("member_name"),
					"email": row.get("member_email"),
					"tasks": [],
				})
				if row.get("title"):
					bucket["tasks"].append(row)
			for member in by_member.values():
				email = member.get("email")
				tasks_for_member = member.get("tasks") or []
				if not email or not tasks_for_member:
					continue
				lines = [
					f"Hi {member.get('name') or 'team member'},",
					"",
					f"Here are your current assigned items for '{project.get('title')}':",
					"",
				]
				for item in tasks_for_member:
					due_label = item.get("due_date").isoformat() if getattr(item.get("due_date"), "isoformat", None) else "No date"
					lines.append(
						f"- {item.get('title')} (status: {item.get('status')}, "
						f"progress: {item.get('progress_percent') or 0}%, due: {due_label})"
					)
				lines.extend([
					"",
					f"Sent by: {current_user.name} ({current_user.email})",
				])
				err = _send_reminder_email(
					to_email=email,
					subject=f"[Group Project] {project.get('title')} - Assigned Tasks Digest",
					body="\n".join(lines),
				)
				if err:
					errors += 1
				else:
					messages_sent += 1
			if messages_sent:
				flash(f"Sent assignment digests to {messages_sent} member(s).", "success")
			if errors:
				flash(f"{errors} email(s) failed to send.", "warning")

		# Reference: ChatGPT (OpenAI) - Group Workspace Collaboration Extensions
		# Date: 2026-02-11
		# Prompt: "Extend the group workspace with member invite resend emails, project
		# milestones, and per-task file attachments (upload/delete) while keeping action
		# handling in a single Flask route. Can you provide a clean branching pattern?"
		# ChatGPT provided the extension pattern adapted below.
		elif action == "resend_invite":
			member_id = _coerce_int(request.form.get("member_id"))
			member = None
			if member_id:
				member = sb_fetch_one(
					"""
					SELECT id, member_name, member_email, invite_token
					FROM group_project_members
					WHERE id = :member_id AND project_id = :project_id
					""",
					{"member_id": member_id, "project_id": project_id},
				)
			if not member:
				flash("Member not found.", "error")
			elif not member.get("member_email"):
				flash("Member does not have an email address.", "warning")
			else:
				invite_token = member.get("invite_token") or secrets.token_urlsafe(24)
				if not member.get("invite_token"):
					try:
						sb_execute(
							"""
							UPDATE group_project_members
							SET invite_token = :invite_token, invite_status = 'pending'
							WHERE id = :member_id AND project_id = :project_id
							""",
							{
								"invite_token": invite_token,
								"member_id": member_id,
								"project_id": project_id,
							},
						)
					except Exception:
						pass
				invite_url = _group_invite_url(invite_token)
				subject = f"[Group Project] Invitation to {project.get('title')}"
				body = (
					f"Hi {member.get('member_name')},\n\n"
					f"You have been added to the group project '{project.get('title')}'.\n"
					f"Use this link to open your task portal and update progress:\n\n"
					f"{invite_url}\n\n"
					"After accepting, you can mark your assigned tasks as in progress/done.\n\n"
					f"Invited by: {current_user.name} ({current_user.email})"
				)
				error = _send_reminder_email(
					to_email=member.get("member_email"),
					subject=subject,
					body=body,
				)
				if error:
					flash(f"Failed to send invite: {error}", "error")
				else:
					flash(f"Invite email sent to {member.get('member_name')}.", "success")

		elif action == "add_milestone":
			title = (request.form.get("milestone_title") or "").strip()
			target_date = _parse_iso_date_value(request.form.get("milestone_target_date"))
			notes = (request.form.get("milestone_notes") or "").strip() or None
			if not title:
				flash("Milestone title is required.", "error")
			elif not target_date:
				flash("Milestone target date is required.", "error")
			else:
				try:
					sb_execute(
						"""
						INSERT INTO group_project_milestones (
							project_id, title, target_date, notes, is_completed, completed_at, created_at
						) VALUES (
							:project_id, :title, :target_date, :notes, FALSE, NULL, NOW()
						)
						""",
						{
							"project_id": project_id,
							"title": title[:255],
							"target_date": target_date,
							"notes": notes,
						},
					)
					flash("Milestone added.", "success")
				except Exception as exc:
					print(f"[group-workspace] add milestone failed user={current_user.id} project={project_id} err={exc}")
					flash("Failed to add milestone.", "error")

		elif action == "toggle_milestone":
			milestone_id = _coerce_int(request.form.get("milestone_id"))
			if milestone_id:
				try:
					milestone = sb_fetch_one(
						"""
						SELECT id, is_completed
						FROM group_project_milestones
						WHERE id = :milestone_id AND project_id = :project_id
						""",
						{"milestone_id": milestone_id, "project_id": project_id},
					)
					if milestone:
						is_completed = bool(milestone.get("is_completed"))
						sb_execute(
							"""
							UPDATE group_project_milestones
							SET is_completed = :is_completed,
							    completed_at = CASE WHEN :is_completed THEN NOW() ELSE NULL END
							WHERE id = :milestone_id AND project_id = :project_id
							""",
							{
								"is_completed": not is_completed,
								"milestone_id": milestone_id,
								"project_id": project_id,
							},
						)
						flash("Milestone updated.", "success")
				except Exception as exc:
					print(f"[group-workspace] toggle milestone failed user={current_user.id} milestone={milestone_id} err={exc}")
					flash("Failed to update milestone.", "error")

		elif action == "upload_task_file":
			task_id = _coerce_int(request.form.get("task_id"))
			file_storage = request.files.get("task_file")
			task_row = None
			if task_id:
				task_row = sb_fetch_one(
					"SELECT id, title FROM group_project_tasks WHERE id = :task_id AND project_id = :project_id",
					{"task_id": task_id, "project_id": project_id},
				)
			if not task_row:
				flash("Task not found.", "error")
			elif not file_storage or not file_storage.filename:
				flash("Please choose a file to upload.", "warning")
			else:
				try:
					filename = os.path.basename((file_storage.filename or "").strip())
					if not filename:
						raise ValueError("Invalid filename.")
					payload = file_storage.read()
					if not payload:
						raise ValueError("The uploaded file was empty.")
					if len(payload) > _GROUP_MAX_UPLOAD_BYTES:
						raise ValueError("File is larger than 8 MB.")
					uploads_dir = os.path.join(app.root_path, "uploads", "group_tasks")
					os.makedirs(uploads_dir, exist_ok=True)
					timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
					stored_name = f"{project_id}_{task_id}_{timestamp}_{filename}"
					filepath = os.path.join(uploads_dir, stored_name)
					with open(filepath, "wb") as f:
						f.write(payload)
					sb_execute(
						"""
						INSERT INTO group_project_task_files (
							task_id, project_id, uploaded_by_student_id, filename, filepath, file_size_bytes, uploaded_at
						) VALUES (
							:task_id, :project_id, :uploaded_by_student_id, :filename, :filepath, :file_size_bytes, NOW()
						)
						""",
						{
							"task_id": task_id,
							"project_id": project_id,
							"uploaded_by_student_id": current_user.id,
							"filename": filename[:255],
							"filepath": filepath,
							"file_size_bytes": len(payload),
						},
					)
					flash("Task attachment uploaded.", "success")
				except ValueError as exc:
					flash(str(exc), "error")
				except Exception as exc:
					print(f"[group-workspace] upload attachment failed user={current_user.id} task={task_id} err={exc}")
					flash("Failed to upload attachment.", "error")

		elif action == "delete_task_file":
			file_id = _coerce_int(request.form.get("file_id"))
			if file_id:
				try:
					row = sb_fetch_one(
						"""
						SELECT id, filepath
						FROM group_project_task_files
						WHERE id = :file_id AND project_id = :project_id
						""",
						{"file_id": file_id, "project_id": project_id},
					)
					if row:
						sb_execute(
							"DELETE FROM group_project_task_files WHERE id = :file_id AND project_id = :project_id",
							{"file_id": file_id, "project_id": project_id},
						)
						filepath = row.get("filepath")
						if filepath and os.path.exists(filepath):
							try:
								os.remove(filepath)
							except Exception:
								pass
						flash("Attachment deleted.", "success")
				except Exception as exc:
					print(f"[group-workspace] delete attachment failed user={current_user.id} file={file_id} err={exc}")
					flash("Failed to delete attachment.", "error")

		elif action == "upload_project_file":
			file_storage = request.files.get("project_file")
			if not file_storage or not file_storage.filename:
				flash("Please choose a project file to upload.", "warning")
			else:
				try:
					filename = os.path.basename((file_storage.filename or "").strip())
					if not filename:
						raise ValueError("Invalid filename.")
					payload = file_storage.read()
					if not payload:
						raise ValueError("The uploaded file was empty.")
					if len(payload) > _GROUP_MAX_UPLOAD_BYTES:
						raise ValueError("Project file is larger than 8 MB.")
					uploads_dir = os.path.join(app.root_path, "uploads", "group_projects")
					os.makedirs(uploads_dir, exist_ok=True)
					timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
					stored_name = f"{project_id}_{timestamp}_{filename}"
					filepath = os.path.join(uploads_dir, stored_name)
					with open(filepath, "wb") as f:
						f.write(payload)
					sb_execute(
						"""
						INSERT INTO group_project_files (
							project_id, uploaded_by_student_id, filename, filepath, file_size_bytes, uploaded_at
						) VALUES (
							:project_id, :uploaded_by_student_id, :filename, :filepath, :file_size_bytes, NOW()
						)
						""",
						{
							"project_id": project_id,
							"uploaded_by_student_id": current_user.id,
							"filename": filename[:255],
							"filepath": filepath,
							"file_size_bytes": len(payload),
						},
					)
					flash("Project file uploaded.", "success")
				except ValueError as exc:
					flash(str(exc), "error")
				except Exception as exc:
					print(f"[group-workspace] upload project file failed user={current_user.id} project={project_id} err={exc}")
					flash("Failed to upload project file.", "error")

		elif action == "delete_project_file":
			file_id = _coerce_int(request.form.get("file_id"))
			if file_id:
				try:
					row = sb_fetch_one(
						"""
						SELECT id, filepath
						FROM group_project_files
						WHERE id = :file_id AND project_id = :project_id
						""",
						{"file_id": file_id, "project_id": project_id},
					)
					if row:
						sb_execute(
							"DELETE FROM group_project_files WHERE id = :file_id AND project_id = :project_id",
							{"file_id": file_id, "project_id": project_id},
						)
						filepath = row.get("filepath")
						if filepath and os.path.exists(filepath):
							try:
								os.remove(filepath)
							except Exception:
								pass
						flash("Project file deleted.", "success")
				except Exception as exc:
					print(f"[group-workspace] delete project file failed user={current_user.id} file={file_id} err={exc}")
					flash("Failed to delete project file.", "error")

		# Reference: ChatGPT (OpenAI) - Project Team Messaging Thread Flow
		# Date: 2026-02-11
		# Prompt: "I need a project-level team thread in Flask where the project owner
		# can post updates/comments and render messages newest-last in the workspace.
		# Can you provide a simple insert + fetch pattern with validation?"
		# ChatGPT provided the insertion/ordering pattern adapted below.
		elif action == "post_message":
			message_text = (request.form.get("message_text") or "").strip()
			if not message_text:
				flash("Message cannot be empty.", "warning")
			else:
				try:
					sb_execute(
						"""
						INSERT INTO group_project_messages (
							project_id, sender_student_id, sender_name, message, created_at
						) VALUES (
							:project_id, :sender_student_id, :sender_name, :message, NOW()
						)
						""",
						{
							"project_id": project_id,
							"sender_student_id": current_user.id,
							"sender_name": current_user.name or "Project Owner",
							"message": message_text[:2000],
						},
					)
					flash("Message posted.", "success")
				except Exception as exc:
					print(f"[group-workspace] post message failed user={current_user.id} project={project_id} err={exc}")
					flash("Failed to post message.", "error")

		return redirect(url_for("group_workspace", project_id=project_id))

	try:
		projects = sb_fetch_all(
			"""
			SELECT p.id, p.title, p.module_code, p.due_date, p.created_at, p.owner_student_id
			FROM group_projects p
			WHERE p.owner_student_id = :student_id

			UNION

			SELECT p.id, p.title, p.module_code, p.due_date, p.created_at, p.owner_student_id
			FROM group_projects p
			JOIN group_project_members m ON m.project_id = p.id
			WHERE m.invite_status = 'accepted'
			  AND LOWER(m.member_email) = :member_email

			ORDER BY created_at DESC, id DESC
			""",
			{"student_id": current_user.id, "member_email": (current_user.email or "").strip().lower()},
		)
	except Exception as exc:
		print(f"[group-workspace] project list failed user={current_user.id} err={exc}")
		projects = []

	selected_project = None
	if selected_project_id:
		pid = _coerce_int(selected_project_id)
		if pid:
			selected_project = _get_accessible_group_project(pid, current_user.id, current_user.email)
	if not selected_project and projects:
		selected_project = _get_accessible_group_project(projects[0].get("id"), current_user.id, current_user.email)

	members: List[Dict[str, Any]] = []
	tasks: List[Dict[str, Any]] = []
	member_stats: List[Dict[str, Any]] = []
	milestones: List[Dict[str, Any]] = []
	task_files_map: Dict[int, List[Dict[str, Any]]] = {}
	project_files: List[Dict[str, Any]] = []
	project_messages: List[Dict[str, Any]] = []
	progress = {
		"total": 0,
		"todo": 0,
		"in_progress": 0,
		"review": 0,
		"done": 0,
		"overdue": 0,
		"completion_percent": 0,
	}
	if selected_project:
		project_id = selected_project.get("id")
		try:
			members = sb_fetch_all(
				"""
				SELECT id, member_name, member_email, member_role, invite_token, invite_status, accepted_at, created_at
				FROM group_project_members
				WHERE project_id = :project_id
				ORDER BY member_name ASC
				""",
				{"project_id": project_id},
			)
		except Exception:
			members = []
		try:
			tasks = sb_fetch_all(
				"""
				SELECT t.id, t.title, t.description, t.assigned_member_id, t.status, t.priority,
				       t.due_date, t.estimated_hours, t.progress_percent, t.ai_generated,
				       m.member_name, m.member_email
				FROM group_project_tasks t
				LEFT JOIN group_project_members m ON m.id = t.assigned_member_id
				WHERE t.project_id = :project_id
				ORDER BY
					CASE t.priority WHEN 'high' THEN 1 WHEN 'medium' THEN 2 ELSE 3 END,
					t.due_date ASC NULLS LAST,
					t.id ASC
				""",
				{"project_id": project_id},
			)
		except Exception:
			tasks = []
		try:
			milestones = sb_fetch_all(
				"""
				SELECT id, title, target_date, notes, is_completed, completed_at, created_at
				FROM group_project_milestones
				WHERE project_id = :project_id
				ORDER BY target_date ASC, id ASC
				""",
				{"project_id": project_id},
			)
		except Exception:
			milestones = []
		try:
			file_rows = sb_fetch_all(
				"""
				SELECT id, task_id, filename, file_size_bytes, uploaded_at
				FROM group_project_task_files
				WHERE project_id = :project_id
				ORDER BY uploaded_at DESC, id DESC
				""",
				{"project_id": project_id},
			)
		except Exception:
			file_rows = []
		for row in file_rows:
			task_key = row.get("task_id")
			if not task_key:
				continue
			task_files_map.setdefault(task_key, []).append(row)
		try:
			project_files = sb_fetch_all(
				"""
				SELECT id, filename, file_size_bytes, uploaded_at
				FROM group_project_files
				WHERE project_id = :project_id
				ORDER BY uploaded_at DESC, id DESC
				""",
				{"project_id": project_id},
			)
		except Exception:
			project_files = []
		try:
			project_messages = sb_fetch_all(
				"""
				SELECT id, sender_name, message, created_at
				FROM group_project_messages
				WHERE project_id = :project_id
				ORDER BY created_at ASC, id ASC
				LIMIT 300
				""",
				{"project_id": project_id},
			)
		except Exception:
			project_messages = []

		progress["total"] = len(tasks)
		today = datetime.now().date()
		for task in tasks:
			status = (task.get("status") or "todo").strip()
			if status == "in_progress":
				progress["in_progress"] += 1
			elif status == "review":
				progress["review"] += 1
			elif status == "done":
				progress["done"] += 1
			else:
				progress["todo"] += 1
			due_date = task.get("due_date")
			if status != "done" and due_date and getattr(due_date, "__lt__", None) and due_date < today:
				progress["overdue"] += 1
		if progress["total"] > 0:
			progress["completion_percent"] = round((progress["done"] / progress["total"]) * 100)

		member_map: Dict[int, Dict[str, Any]] = {
			m.get("id"): {
				"id": m.get("id"),
				"name": m.get("member_name"),
				"email": m.get("member_email"),
				"assigned": 0,
				"done": 0,
				"in_progress": 0,
			}
			for m in members
		}
		for task in tasks:
			member_id = task.get("assigned_member_id")
			if not member_id or member_id not in member_map:
				continue
			member_map[member_id]["assigned"] += 1
			if task.get("status") == "done":
				member_map[member_id]["done"] += 1
			elif task.get("status") == "in_progress":
				member_map[member_id]["in_progress"] += 1
		member_stats = list(member_map.values())
		for item in member_stats:
			total = item.get("assigned") or 0
			item["completion_percent"] = round(((item.get("done") or 0) / total) * 100) if total else 0
		member_stats.sort(key=lambda x: (-(x.get("assigned") or 0), x.get("name") or ""))

	return render_template(
		"group_workspace.html",
		projects=projects,
		selected_project=selected_project,
		members=members,
		tasks=tasks,
		milestones=milestones,
		task_files_map=task_files_map,
		project_files=project_files,
		project_messages=project_messages,
		member_stats=member_stats,
		progress=progress,
		invite_base_url=request.host_url.rstrip("/"),
	)


# Reference: ChatGPT (OpenAI) - Invite Token Member Portal Pattern
# Date: 2026-02-11
# Prompt: "I need invited group members to open a secure token link, accept the
# invite, and update only tasks assigned to them without full account login.
# Can you provide a safe Flask route flow?"
# ChatGPT provided the token validation + restricted update flow adapted below.
@app.route("/group-workspace/invite/<invite_token>", methods=["GET", "POST"])
def group_workspace_invite(invite_token: str):
	member = sb_fetch_one(
		"""
		SELECT m.id, m.project_id, m.member_name, m.member_email, m.invite_status, m.accepted_at,
		       p.title AS project_title, p.module_code, p.due_date
		FROM group_project_members m
		JOIN group_projects p ON p.id = m.project_id
		WHERE m.invite_token = :invite_token
		LIMIT 1
		""",
		{"invite_token": invite_token},
	)
	if not member:
		return render_template("group_workspace_invite.html", invalid_link=True, invite_token=invite_token), 404

	if request.method == "POST":
		action = (request.form.get("action") or "").strip()
		if action == "accept_invite":
			try:
				sb_execute(
					"""
					UPDATE group_project_members
					SET invite_status = 'accepted',
					    accepted_at = COALESCE(accepted_at, NOW())
					WHERE id = :member_id
					""",
					{"member_id": member.get("id")},
				)
				member["invite_status"] = "accepted"
				flash("Invite accepted. You can now update your tasks.", "success")
			except Exception as exc:
				print(f"[group-workspace] accept invite failed member={member.get('id')} err={exc}")
				flash("Failed to accept invite.", "error")
		elif action == "update_member_task":
			task_id = _coerce_int(request.form.get("task_id"))
			status = (request.form.get("status") or "").strip()
			progress_percent = _coerce_int(request.form.get("progress_percent"))
			if status not in {"todo", "in_progress", "review", "done"}:
				status = "todo"
			progress_value = max(0, min(100, progress_percent if progress_percent is not None else 0))
			if status == "done":
				progress_value = 100
			if not task_id:
				flash("Invalid task.", "error")
			else:
				task = sb_fetch_one(
					"""
					SELECT id
					FROM group_project_tasks
					WHERE id = :task_id
					  AND project_id = :project_id
					  AND assigned_member_id = :member_id
					""",
					{
						"task_id": task_id,
						"project_id": member.get("project_id"),
						"member_id": member.get("id"),
					},
				)
				if not task:
					flash("Task not found.", "error")
				else:
					try:
						sb_execute(
							"""
							UPDATE group_project_tasks
							SET status = :status,
							    progress_percent = :progress_percent,
							    updated_at = NOW()
							WHERE id = :task_id
							  AND project_id = :project_id
							  AND assigned_member_id = :member_id
							""",
							{
								"status": status,
								"progress_percent": progress_value,
								"task_id": task_id,
								"project_id": member.get("project_id"),
								"member_id": member.get("id"),
							},
						)
						flash("Task updated.", "success")
					except Exception as exc:
						print(f"[group-workspace] member task update failed member={member.get('id')} task={task_id} err={exc}")
						flash("Failed to update task.", "error")
		return redirect(url_for("group_workspace_invite", invite_token=invite_token))

	tasks = sb_fetch_all(
		"""
		SELECT id, title, description, status, priority, due_date, progress_percent
		FROM group_project_tasks
		WHERE project_id = :project_id
		  AND assigned_member_id = :member_id
		ORDER BY due_date ASC NULLS LAST, id ASC
		""",
		{
			"project_id": member.get("project_id"),
			"member_id": member.get("id"),
		},
	)
	return render_template(
		"group_workspace_invite.html",
		invalid_link=False,
		invite_token=invite_token,
		member=member,
		tasks=tasks,
	)


# Reference: ChatGPT (OpenAI) - Secure Owner-Only File Download Route
# Date: 2026-02-11
# Prompt: "I need a Flask download endpoint for group-task attachments where only
# the project owner can download files. Can you provide a secure pattern with
# ownership checks and safe send_file handling?"
# ChatGPT provided the ownership-check + send_file route pattern adapted below.
@app.route("/group-workspace/files/<int:file_id>/download")
@login_required
def group_workspace_download_file(file_id: int):
	row = sb_fetch_one(
		"""
		SELECT f.id, f.filename, f.filepath
		FROM group_project_task_files f
		JOIN group_projects p ON p.id = f.project_id
		WHERE f.id = :file_id
		  AND p.owner_student_id = :student_id
		""",
		{"file_id": file_id, "student_id": current_user.id},
	)
	if not row:
		abort(404)
	filepath = row.get("filepath")
	if not filepath or not os.path.exists(filepath):
		abort(404)
	return send_file(filepath, as_attachment=True, download_name=row.get("filename") or "attachment")


@app.route("/group-workspace/project-files/<int:file_id>/download")
@login_required
def group_workspace_download_project_file(file_id: int):
	# Reference: ChatGPT (OpenAI) - Secure Project File Download Guard
	# Date: 2026-02-11
	# Prompt: "I need a secure Flask download route for project-level files where
	# only the project owner can access downloads. Can you provide a safe ownership
	# check and send_file pattern?"
	# ChatGPT provided the owner-check + safe send_file pattern adapted here.
	row = sb_fetch_one(
		"""
		SELECT f.id, f.filename, f.filepath
		FROM group_project_files f
		JOIN group_projects p ON p.id = f.project_id
		WHERE f.id = :file_id
		  AND p.owner_student_id = :student_id
		""",
		{"file_id": file_id, "student_id": current_user.id},
	)
	if not row:
		abort(404)
	filepath = row.get("filepath")
	if not filepath or not os.path.exists(filepath):
		abort(404)
	return send_file(filepath, as_attachment=True, download_name=row.get("filename") or "project_file")


@app.route("/ai-workspace", methods=["GET", "POST"])
@login_required
def ai_workspace():
	error: Optional[str] = None
	breakdown = None
	combined_text: Optional[str] = None
	sanitized_summary: Optional[str] = None
	uploaded_filename: Optional[str] = None
	form_data: Dict[str, str] = {}
	schedule_context_text: Optional[str] = None
	plan_payload: Optional[Dict[str, Any]] = None
	
	if request.method == "POST":
		form_data = request.form.to_dict()
		title = (form_data.get("assignment_title") or "").strip()
		module_code = (form_data.get("module_code") or "").strip() or None
		due_date = (form_data.get("due_date") or "").strip() or None
		status = (form_data.get("status") or "").strip() or None
		context_text = (form_data.get("additional_context") or "").strip()
		manual_text = (form_data.get("brief_text") or "").strip()
		
		segments: List[str] = []
		file_storage = request.files.get("brief_file")
		if file_storage and file_storage.filename:
			uploaded_filename = file_storage.filename
			try:
				payload = file_storage.read()
				if not payload:
					raise ValueError("The uploaded file was empty.")
				if len(payload) > _AI_MAX_UPLOAD_BYTES:
					raise ValueError("The file is larger than 4 MB. Upload a smaller brief.")
				segments.append(_extract_text_from_brief(uploaded_filename, payload))
			except ValueError as exc:
				error = str(exc)
				print(f"[ai-workspace] file validation error user={current_user.id} file={uploaded_filename} err={exc}")
			except Exception as exc:  # pragma: no cover - defensive guard
				print(f"[ai-workspace] file read failed user={current_user.id} file={uploaded_filename} err={exc}")
				import traceback
				traceback.print_exc()
				error = f"Failed to read the uploaded file: {str(exc)}"
		
		if manual_text:
			segments.append(manual_text)
		
		if not error:
			combined_text = "\n\n".join([segment for segment in segments if segment]).strip()
			if not combined_text:
				error = "Provide an assignment brief by uploading a file or pasting text."
			else:
				sanitized_summary = _sanitize_prompt_text(combined_text) or _summarize_text(combined_text, max_len=400)
				context_summary = _sanitize_prompt_text(context_text) or None
				if not sanitized_summary:
					error = "Could not prepare the assignment brief. Try adding more detail."
				else:
					schedule_context = _build_schedule_context(current_user.id)
					schedule_context_text = schedule_context
					try:
						service = get_chatgpt_service()
						breakdown = service.breakdown_task(
							task_title=title or "Assignment Plan",
							module_code=module_code,
							due_date=due_date,
							due_at=None,
							status=status,
							description=sanitized_summary,
							additional_context=context_summary,
							schedule_context=schedule_context
						)
					except ChatGPTClientError as exc:
						error = str(exc)
					except Exception as exc:  # pragma: no cover - defensive logging
						print(f"[ai-workspace] chatgpt breakdown failed user={current_user.id} err={exc}")
						error = "ChatGPT request failed. Try again later."
					else:
						if breakdown:
							plan_payload = {
								"subtasks": [
									{
										"sequence": item.sequence,
										"title": item.title,
										"description": item.description,
										"estimated_hours": item.estimated_hours,
										"planned_start": item.planned_start,
										"planned_end": item.planned_end,
										"focus": item.focus,
									}
									for item in breakdown.subtasks
								],
								"advice": breakdown.advice,
							}
	else:
		form_data = {
			"assignment_title": "",
			"module_code": "",
			"due_date": "",
			"status": "",
			"brief_text": "",
			"additional_context": "",
		}
	
	return render_template(
		"ai_workspace.html",
		breakdown=breakdown,
		error=error,
		form_data=form_data,
		source_text=combined_text,
		sanitized_summary=sanitized_summary,
		uploaded_filename=uploaded_filename,
		schedule_context_text=schedule_context_text,
		plan_payload=plan_payload,
	)


@app.route("/ai-workspace/save", methods=["POST"])
@login_required
def ai_workspace_save():
	plan_json = request.form.get("plan_json")
	if not plan_json:
		flash("No AI plan found to save.", "danger")
		return redirect(url_for("ai_workspace"))
	try:
		payload = json.loads(plan_json)
	except json.JSONDecodeError:
		flash("The AI plan data was invalid.", "danger")
		return redirect(url_for("ai_workspace"))
	subtasks = payload.get("subtasks")
	if not isinstance(subtasks, list) or not subtasks:
		flash("The AI plan did not contain any subtasks to schedule.", "warning")
		return redirect(url_for("ai_workspace"))

	assignment_title = (request.form.get("assignment_title") or "Assignment Plan").strip()
	module_code = (request.form.get("module_code") or "").strip() or None
	due_date_str = (request.form.get("due_date") or "").strip() or None

	module_id = _lookup_module_id(module_code, current_user.id)
	
	# Look up or create task for the subtasks
	task_id = _lookup_or_create_task(assignment_title, module_id, module_code, current_user.id, due_date_str)
	print(f"[ai-workspace] task_id={task_id} for assignment '{assignment_title}' (module_id={module_id})")
	
	if not task_id:
		print(f"[ai-workspace] WARNING: Could not find or create task for '{assignment_title}'. Subtasks will only be saved to calendar.")
	
	try:
		scheduled, unscheduled = _schedule_ai_subtasks(
			subtasks=subtasks,
			student_id=current_user.id,
			module_id=module_id,
			assignment_title=assignment_title,
			due_date_str=due_date_str,
		)
	except Exception as exc:
		print(f"[ai-workspace] schedule failed user={current_user.id} err={exc}")
		import traceback
		traceback.print_exc()
		flash("Failed to schedule the AI plan. Please try again.", "danger")
		return redirect(url_for("ai_workspace"))

	created_events = 0
	created_subtasks = 0
	
	# Save subtasks to the subtasks table (if task was found/created)
	if task_id:
		print(f"[ai-workspace] Saving {len(subtasks)} subtasks to task_id={task_id}")
		for idx, item in enumerate(subtasks, start=1):
			try:
				sequence = item.get("sequence", idx)
				subtask_title = item.get("title") or f"Subtask {idx}"
				description = item.get("description")
				estimated_hours = item.get("estimated_hours")
				planned_start = item.get("planned_start")
				planned_end = item.get("planned_end")
				
				# Parse planned dates if they're strings
				planned_start_date = None
				planned_end_date = None
				if planned_start:
					try:
						from datetime import datetime
						if isinstance(planned_start, str):
							planned_start_date = datetime.fromisoformat(planned_start.replace('Z', '+00:00')).date()
						else:
							planned_start_date = planned_start.date() if hasattr(planned_start, 'date') else planned_start
					except:
						pass
				if planned_end:
					try:
						from datetime import datetime
						if isinstance(planned_end, str):
							planned_end_date = datetime.fromisoformat(planned_end.replace('Z', '+00:00')).date()
						else:
							planned_end_date = planned_end.date() if hasattr(planned_end, 'date') else planned_end
					except:
						pass
				
				sb_execute(
					"""
					INSERT INTO subtasks (
						task_id, title, description, sequence, estimated_hours,
						planned_start, planned_end
					) VALUES (
						:task_id, :title, :description, :sequence, :estimated_hours,
						:planned_start, :planned_end
					)
					""",
					{
						"task_id": task_id,
						"title": subtask_title[:255],
						"description": description[:1000] if description else None,
						"sequence": int(sequence) if sequence else idx,
						"estimated_hours": float(estimated_hours) if estimated_hours else None,
						"planned_start": planned_start_date,
						"planned_end": planned_end_date,
					}
				)
				created_subtasks += 1
			except Exception as exc:
				print(f"[ai-workspace] subtask insert failed user={current_user.id} subtask_idx={idx} err={exc}")
				import traceback
				traceback.print_exc()
				continue
		
		print(f"[ai-workspace] Successfully saved {created_subtasks}/{len(subtasks)} subtasks to task_id={task_id}")
	else:
		print(f"[ai-workspace] Skipping subtask save - no task_id available")
	
	# Save scheduled subtasks to calendar events
	for item in scheduled:
		start_ts = item.get("start")
		end_ts = item.get("end")
		if not start_ts or not end_ts:
			continue
		title = item.get("title") or assignment_title
		try:
			sb_execute(
				"""
				INSERT INTO events (student_id, module_id, title, start_at, end_at, location)
				VALUES (:student_id, :module_id, :title, :start_at, :end_at, :location)
				""",
				{
					"student_id": current_user.id,
					"module_id": item.get("module_id"),
					"title": title[:255],
					"start_at": start_ts,
					"end_at": end_ts,
					"location": item.get("focus"),
				}
			)
			created_events += 1
		except Exception as exc:
			print(f"[ai-workspace] event insert failed user={current_user.id} err={exc}")
			continue

	# Build success message
	messages = []
	if created_subtasks:
		messages.append(f"Saved {created_subtasks} subtask{'s' if created_subtasks != 1 else ''} to task breakdown.")
	if created_events:
		messages.append(f"Scheduled {created_events} subtask{'s' if created_events != 1 else ''} into your calendar.")
	if unscheduled:
		messages.append(f"{len(unscheduled)} item{'s' if len(unscheduled) != 1 else ''} could not be scheduled due to limited availability.")
	
	if messages:
		flash(" ".join(messages), "success")
	else:
		if unscheduled:
			flash("No suitable time slots were available to schedule those subtasks.", "warning")
		else:
			flash("Unable to save the AI plan. Please try again.", "danger")

	return redirect(url_for("ai_workspace"))

@app.route("/calendar")
@login_required
def calendar_view():
	"""Calendar view of tasks by due date"""
	sql = """
		SELECT t.id, t.title, t.status, t.due_date,
		       t.due_at,
		       t.canvas_assignment_id, m.code AS module_code
		FROM tasks t
		JOIN modules m ON m.id = t.module_id
		WHERE t.student_id = :student_id
		AND t.due_date IS NOT NULL
		ORDER BY t.due_date ASC
		LIMIT 500
	"""
	try:
		rows: List[Dict] = sb_fetch_all(sql, {"student_id": current_user.id})
	except Exception:
		rows = []

	# Reference: ChatGPT (OpenAI) - Calendar Event Time Boundary Handling
	# Date: 2025-11-04
	# Prompt: "Canvas sometimes gives a due_at timestamp or just a due_date. Can you show me how to safely 
	# convert either one into a start and end time in ISO format for my calendar view?"
	# ChatGPT provided a helper block to convert Canvas assignment due information into calendar event start 
	# and end timestamps. When due_at (a full datetime) is available, it is used as the event start, with an 
	# end time five minutes later and clamped to the end of the same day. If only a due_date is available, 
	# the code treats it as an all-day event. This ensures every assignment can be plotted on the calendar, 
	# even when Canvas only supplies date-level precision.
	events: List[Dict] = []
	for r in rows:
		due_at = r.get("due_at")
		if due_at:
			try:
				start_iso = due_at.isoformat()
			except Exception:
				start_iso = str(due_at)
			try:
				from datetime import timedelta
				start_dt = r.get("due_at")
				end_dt = start_dt + timedelta(minutes=5)
				end_of_day = start_dt.replace(hour=23, minute=59, second=59, microsecond=0)
				if end_dt > end_of_day:
					end_dt = end_of_day
				end_iso = end_dt.isoformat()
			except Exception:
				end_iso = start_iso
		else:
			due_date = r.get("due_date")
			start_iso = due_date.strftime("%Y-%m-%d") if hasattr(due_date, "strftime") else str(due_date)
			# Reference: ChatGPT (OpenAI) - FullCalendar All-Day End Date (Exclusive)
			# Date: 2026-02-11
			# Prompt: "FullCalendar treats all-day event end dates as exclusive. If I set
			# end == start for an all-day event, it can render oddly. How should I set
			# the end date for an all-day deadline so it displays on the correct day?"
			# ChatGPT recommended using end = start + 1 day for all-day events.
			try:
				end_date = due_date + timedelta(days=1)
				end_iso = end_date.strftime("%Y-%m-%d")
			except Exception:
				end_iso = start_iso

		status = r.get("status", "pending")
		if status == "completed":
			color = "#16a34a"
		elif status == "in_progress":
			color = "#f59e0b"
		else:
			color = "#2563eb"

		is_canvas = r.get("canvas_assignment_id") is not None
		prefix = "📚 " if is_canvas else ""

		event = {
			"id": r.get("id"),
			"title": f"{prefix}{r.get('title')} [{r.get('module_code')}]",
			"start": start_iso,
			"end": end_iso,
			"allDay": False if due_at else True,
			"color": color,
			"extendedProps": {
				"status": status,
				"module": r.get("module_code")
			},
			"classNames": ["assignment"]
		}
		events.append(event)

	try:
		rows_ev: List[Dict] = sb_fetch_all(
			"""
			SELECT e.id, e.title, e.start_at, e.end_at, e.location, e.canvas_course_id, e.canvas_event_id, m.code AS module_code
			FROM events e
			LEFT JOIN modules m ON m.id = e.module_id
			WHERE e.student_id = :student_id
			AND e.start_at IS NOT NULL
			AND e.end_at IS NOT NULL
			ORDER BY e.start_at ASC
			LIMIT 1000
			""",
			{"student_id": current_user.id}
		)
	except Exception:
		rows_ev = []

	for ev in rows_ev:
		start_val = ev.get("start_at")
		end_val = ev.get("end_at")
		# Reference: FullCalendar.js Documentation - Event Object
		# https://fullcalendar.io/docs/event-object
		# FullCalendar requires ISO-8601 date strings for timed events
		try:
			start_iso = start_val.isoformat()
		except Exception:
			start_iso = str(start_val)
		try:
			end_iso = end_val.isoformat()
		except Exception:
			end_iso = str(end_val)
		mod = ev.get("module_code")
		title = ev.get("title")
		canvas_course_id = ev.get("canvas_course_id")
		# Manual events (no Canvas course ID) get a different color
		if canvas_course_id is None:
			# Manual event (work, training, etc.) - purple color
			color = "#9333ea"
			prefix = "📅 "
		else:
			# Canvas-synced event (lectures, etc.) - cyan color
			color = "#0ea5e9"
			prefix = "📚 "
		events.append({
			"id": f"event-{ev.get('id')}",
			"title": f"{prefix}{title} [{mod}]" if mod else f"{prefix}{title}",
			"start": start_iso,
			"end": end_iso,
			"allDay": False,
			"color": color,
			"extendedProps": {
				"module": mod,
				"location": ev.get("location")
			}
		})

	try:
		print(f"[calendar] user={current_user.id} events_total={len(events)}")
	except Exception:
		pass

	try:
		user_events = sb_fetch_all(
			"""
			SELECT id, title, start_at, end_at, is_recurring
			FROM events
			WHERE student_id = :student_id
			ORDER BY start_at ASC
			LIMIT 200
			""",
			{"student_id": current_user.id}
		)
	except Exception:
		user_events = []

	# Detect recurring patterns by grouping events with same title + day + time
	# This works for both explicitly marked recurring events AND Canvas lectures
	pattern_groups = {}
	
	for event in user_events:
		start_at = event.get('start_at')
		end_at = event.get('end_at')
		title = event.get('title', '')
		
		# Parse start datetime
		if isinstance(start_at, str):
			try:
				start_dt = datetime.fromisoformat(start_at.replace('Z', '+00:00'))
			except:
				start_dt = datetime.now()
		elif start_at:
			start_dt = start_at
		else:
			continue  # Skip events without start time
		
		# Parse end datetime
		if isinstance(end_at, str):
			try:
				end_dt = datetime.fromisoformat(end_at.replace('Z', '+00:00'))
			except:
				end_dt = start_dt
		elif end_at:
			end_dt = end_at
		else:
			end_dt = start_dt
		
		day_of_week = start_dt.strftime('%A')  # e.g., "Monday"
		start_time = start_dt.strftime('%H:%M')  # e.g., "09:00"
		end_time = end_dt.strftime('%H:%M')  # e.g., "17:00"
		
		# Create a unique key for this pattern (title + day + time)
		pattern_key = f"{title}|{day_of_week}|{start_time}|{end_time}"
		
		if pattern_key not in pattern_groups:
			pattern_groups[pattern_key] = {
				'id': event['id'],
				'title': title,
				'day_of_week': day_of_week,
				'start_time': start_time,
				'end_time': end_time,
				'count': 1,
				'events': [event],
				'is_recurring': event.get('is_recurring', False)
			}
		else:
			pattern_groups[pattern_key]['count'] += 1
			pattern_groups[pattern_key]['events'].append(event)
	
	# Separate into recurring (2+ occurrences) and single events
	recurring_events = []
	single_events = []
	
	for pattern_key, group in pattern_groups.items():
		if group['count'] >= 2:
			# This is a recurring pattern (same title, day, time - 2 or more times)
			recurring_events.append({
				'id': group['id'],
				'title': group['title'],
				'day_of_week': group['day_of_week'],
				'start_time': group['start_time'],
				'end_time': group['end_time'],
				'count': group['count'],
				'is_recurring': True
			})
		else:
			# Single occurrence - show as one-time event
			single_events.extend(group['events'])

	return render_template("calendar.html", events=events, user_events=single_events, recurring_events=recurring_events)


@app.route("/calendar/add", methods=["POST"])
@login_required
def add_calendar_event():
	"""Add a manual calendar event (e.g., part-time work, training)"""
	title = request.form.get("title", "").strip()
	start_date = request.form.get("start_date", "").strip()
	start_time = request.form.get("start_time", "").strip()
	end_date = request.form.get("end_date", "").strip()
	end_time = request.form.get("end_time", "").strip()
	location = request.form.get("location", "").strip() or None
	is_recurring = request.form.get("is_recurring") == "on"
	recurrence_end_date_str = request.form.get("recurrence_end_date", "").strip()
	
	if not title:
		flash("Title is required.", "error")
		return redirect(url_for("calendar_view"))
	
	if not start_date or not start_time:
		flash("Start date and time are required.", "error")
		return redirect(url_for("calendar_view"))
	
	if not end_date or not end_time:
		flash("End date and time are required.", "error")
		return redirect(url_for("calendar_view"))
	
	try:
		from datetime import datetime, timedelta, date
		# Combine date and time strings
		start_str = f"{start_date} {start_time}"
		end_str = f"{end_date} {end_time}"
		
		# Parse to datetime (assuming local timezone, convert to UTC if needed)
		try:
			start_dt = datetime.strptime(start_str, "%Y-%m-%d %H:%M")
			end_dt = datetime.strptime(end_str, "%Y-%m-%d %H:%M")
		except ValueError:
			flash("Invalid date or time format. Use YYYY-MM-DD and HH:MM.", "error")
			return redirect(url_for("calendar_view"))
		
		# Ensure timezone aware
		from datetime import timezone
		if start_dt.tzinfo is None:
			# Convert local time to UTC
			start_dt = start_dt.replace(tzinfo=timezone.utc)
		if end_dt.tzinfo is None:
			end_dt = end_dt.replace(tzinfo=timezone.utc)
		
		if end_dt <= start_dt:
			flash("End time must be after start time.", "error")
			return redirect(url_for("calendar_view"))
		
		# Calculate duration for recurring events
		duration = end_dt - start_dt
		
		# Handle recurring events
		if is_recurring:
			if not recurrence_end_date_str:
				# Default to 3 months from start date if no end date specified
				recurrence_end_date = (start_dt.date() + timedelta(days=90))
			else:
				try:
					recurrence_end_date = datetime.strptime(recurrence_end_date_str, "%Y-%m-%d").date()
				except ValueError:
					flash("Invalid recurrence end date format. Use YYYY-MM-DD.", "error")
					return redirect(url_for("calendar_view"))
			
			# Generate weekly recurring instances
			current_start = start_dt
			current_date = start_dt.date()
			created_count = 0
			
			while current_date <= recurrence_end_date:
				current_end = current_start + duration
				
				# Insert each recurring instance
				sb_execute(
					"""
					INSERT INTO events (student_id, title, start_at, end_at, location, is_recurring, recurrence_end_date)
					VALUES (:student_id, :title, :start_at, :end_at, :location, :is_recurring, :recurrence_end_date)
					""",
					{
						"student_id": current_user.id,
						"title": title[:255],
						"start_at": current_start,
						"end_at": current_end,
						"location": location,
						"is_recurring": True,
						"recurrence_end_date": recurrence_end_date,
					}
				)
				created_count += 1
				
				# Move to next week (same day, same time)
				current_start = current_start + timedelta(weeks=1)
				current_date = current_start.date()
			
			flash(f"Recurring event '{title}' added - {created_count} instances created until {recurrence_end_date.strftime('%d/%m/%Y')}.", "success")
		else:
			# Single event (non-recurring)
			sb_execute(
				"""
				INSERT INTO events (student_id, title, start_at, end_at, location, is_recurring)
				VALUES (:student_id, :title, :start_at, :end_at, :location, :is_recurring)
				""",
				{
					"student_id": current_user.id,
					"title": title[:255],
					"start_at": start_dt,
					"end_at": end_dt,
					"location": location,
					"is_recurring": False,
				}
			)
			flash(f"Event '{title}' added to your calendar.", "success")
	except Exception as exc:
		print(f"[calendar] add event failed user={current_user.id} err={exc}")
		flash(f"Failed to add event: {str(exc)}", "danger")
	
	return redirect(url_for("calendar_view"))


@app.route("/events/<int:event_id>/delete", methods=["POST"])
@login_required
def delete_event(event_id: int):
	try:
		result = sb_execute(
			"DELETE FROM events WHERE id = :id AND student_id = :student_id",
			{"id": event_id, "student_id": current_user.id}
		)
		if result == 0:
			flash("Event not found or already deleted.", "warning")
		else:
			flash("Event removed from your calendar.", "success")
	except Exception as exc:
		print(f"[events] delete failed user={current_user.id} event={event_id} err={exc}")
		flash("Failed to delete the event. Please try again.", "danger")
	return redirect(url_for("calendar_view"))


@app.route("/events/<int:event_id>/delete-recurring", methods=["POST"])
@login_required
def delete_recurring_event(event_id: int):
	"""Delete all occurrences of a recurring event by matching title and time pattern"""
	try:
		# First get the event details
		event = sb_fetch_one(
			"SELECT title, start_at, end_at FROM events WHERE id = :id AND student_id = :student_id",
			{"id": event_id, "student_id": current_user.id}
		)
		if not event:
			flash("Event not found.", "warning")
			return redirect(url_for("calendar_view"))
		
		title = event.get('title')
		start_at = event.get('start_at')
		end_at = event.get('end_at')
		
		# Parse to get time
		if isinstance(start_at, str):
			try:
				start_dt = datetime.fromisoformat(start_at.replace('Z', '+00:00'))
			except:
				start_dt = None
		else:
			start_dt = start_at
		
		# Parse end datetime (used for tighter pattern matching)
		if isinstance(end_at, str):
			try:
				end_dt = datetime.fromisoformat(end_at.replace('Z', '+00:00'))
			except Exception:
				end_dt = None
		else:
			end_dt = end_at
		
		if start_dt:
			# Extract time components for matching
			# IMPORTANT: PostgreSQL EXTRACT(DOW) uses 0=Sunday..6=Saturday,
			# while Python weekday() uses 0=Monday..6=Sunday.
			# Convert to Postgres-compatible DOW to avoid off-by-one mismatches.
			postgres_dow = (start_dt.weekday() + 1) % 7
			start_time = start_dt.strftime('%H:%M')
			end_time = end_dt.strftime('%H:%M') if end_dt else None
			
			# Delete all events with same title that occur on the same day/time pattern.
			if end_time:
				result = sb_execute(
					"""
					DELETE FROM events 
					WHERE title = :title 
					AND student_id = :student_id
					AND EXTRACT(DOW FROM start_at) = :day_of_week
					AND TO_CHAR(start_at, 'HH24:MI') = :start_time
					AND TO_CHAR(end_at, 'HH24:MI') = :end_time
					""",
					{
						"title": title,
						"student_id": current_user.id,
						"day_of_week": postgres_dow,
						"start_time": start_time,
						"end_time": end_time,
					}
				)
			else:
				result = sb_execute(
					"""
					DELETE FROM events 
					WHERE title = :title 
					AND student_id = :student_id
					AND EXTRACT(DOW FROM start_at) = :day_of_week
					AND TO_CHAR(start_at, 'HH24:MI') = :start_time
					""",
					{
						"title": title,
						"student_id": current_user.id,
						"day_of_week": postgres_dow,
						"start_time": start_time,
					}
				)
		else:
			# Fallback: just delete by title
			result = sb_execute(
				"DELETE FROM events WHERE title = :title AND student_id = :student_id",
				{"title": title, "student_id": current_user.id}
			)
		
		if result == 0:
			flash("No recurring events found to delete.", "warning")
		else:
			flash(f"Deleted {result} occurrences of '{title}' from your calendar.", "success")
	except Exception as exc:
		print(f"[events] delete recurring failed user={current_user.id} event={event_id} err={exc}")
		flash("Failed to delete the recurring events. Please try again.", "danger")
	return redirect(url_for("calendar_view"))


@app.route("/debug/calendar")
@login_required
def debug_calendar():
	"""Return counts of tasks and events for current user for troubleshooting"""
	counts = {"tasks": 0, "events": 0}
	try:
		row = sb_fetch_one("SELECT COUNT(*) AS c FROM tasks WHERE student_id = :sid", {"sid": current_user.id})
		counts["tasks"] = row["c"] if row else 0
		rowe = sb_fetch_one("SELECT COUNT(*) AS c FROM events WHERE student_id = :sid", {"sid": current_user.id})
		counts["events"] = rowe["c"] if rowe else 0
	except Exception as exc:
		return jsonify({"ok": False, "error": str(exc)}), 500
	return jsonify({"ok": True, "counts": counts}), 200


@app.route("/debug/events")
@login_required
def debug_events():
	"""Return next 50 timed events for troubleshooting"""
	try:
		rows = sb_fetch_all(
			"""
			SELECT id, title, start_at, end_at, canvas_course_id
			FROM events
			WHERE student_id = :sid
			AND start_at >= NOW() - INTERVAL '14 days'
			ORDER BY start_at ASC
			LIMIT 50
			""",
			{"sid": current_user.id}
		)
		def to_iso(x):
			try:
				return x.isoformat()
			except Exception:
				return str(x)
		items = [
			{
				"id": r.get("id"),
				"title": r.get("title"),
				"start_at": to_iso(r.get("start_at")),
				"end_at": to_iso(r.get("end_at")),
				"canvas_course_id": r.get("canvas_course_id"),
			}
			for r in rows
		]
		return jsonify({"ok": True, "events": items}), 200
	except Exception as exc:
		return jsonify({"ok": False, "error": str(exc)}), 500


# Reference: ChatGPT (OpenAI) - Weighted Health Score Calculation
# Date: 2026-02-11
# Prompt: "I need a student 'health score' that combines completion rate,
# on-time completion percentage, and lecture attendance. Some components may be
# missing. Can you provide a weighted scoring pattern that renormalizes weights?"
# ChatGPT provided the weighted-average + fallback normalization pattern below.
def _calculate_health_score(
	*,
	task_completion_rate: Optional[float],
	on_time_rate: Optional[float],
	lecture_attendance_rate: Optional[float]
) -> Dict[str, Any]:
	parts: List[Tuple[float, float]] = []
	if task_completion_rate is not None:
		parts.append((float(task_completion_rate), 0.40))
	if on_time_rate is not None:
		parts.append((float(on_time_rate), 0.30))
	if lecture_attendance_rate is not None:
		parts.append((float(lecture_attendance_rate), 0.30))
	if not parts:
		return {"score": 0.0, "label": "no_data"}
	total_weight = sum(weight for _, weight in parts)
	score = round(sum(value * weight for value, weight in parts) / total_weight, 1)
	if score >= 80:
		label = "on_track"
	elif score >= 60:
		label = "watch"
	else:
		label = "at_risk"
	return {"score": score, "label": label}


@app.route("/events/<int:event_id>/attendance-toggle", methods=["POST"])
@login_required
def toggle_lecture_attendance(event_id: int):
	# Reference: ChatGPT (OpenAI) - Attendance Toggle Upsert Pattern
	# Date: 2026-02-11
	# Prompt: "I need a Flask route that toggles lecture attendance for a given
	# calendar event (event belongs to current user). It should insert if missing
	# and flip attended true/false if existing."
	# ChatGPT provided the ownership-check + insert/update toggle flow below.
	next_url = request.form.get("next") or request.referrer or url_for("analytics")
	event = sb_fetch_one(
		"""
		SELECT id
		FROM events
		WHERE id = :event_id AND student_id = :student_id
		""",
		{"event_id": event_id, "student_id": current_user.id},
	)
	if not event:
		flash("Lecture event not found.", "warning")
		return redirect(next_url)
	try:
		row = sb_fetch_one(
			"""
			SELECT id, attended
			FROM lecture_attendance
			WHERE event_id = :event_id AND student_id = :student_id
			""",
			{"event_id": event_id, "student_id": current_user.id},
		)
		if row:
			new_attended = not bool(row.get("attended"))
			sb_execute(
				"""
				UPDATE lecture_attendance
				SET attended = :attended,
				    attended_at = CASE WHEN :attended THEN NOW() ELSE NULL END
				WHERE id = :id
				""",
				{"attended": new_attended, "id": row.get("id")},
			)
		else:
			sb_execute(
				"""
				INSERT INTO lecture_attendance (student_id, event_id, attended, attended_at)
				VALUES (:student_id, :event_id, TRUE, NOW())
				""",
				{"student_id": current_user.id, "event_id": event_id},
			)
	except Exception as exc:
		print(f"[attendance] toggle failed user={current_user.id} event={event_id} err={exc}")
		flash("Failed to update lecture attendance.", "error")
	return redirect(next_url)


@app.route("/analytics")
@login_required
def analytics():
	charts_dir = os.path.join(app.static_folder or "static", "charts")
	os.makedirs(charts_dir, exist_ok=True)

	cutoff_date = datetime.now().date()

	try:
		status_overview = sb_fetch_all("""
			SELECT status, COUNT(id) as count
			FROM tasks
			WHERE student_id = :student_id
			  AND (due_date IS NULL OR due_date >= :cutoff_date)
			GROUP BY status
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		# Reference: ChatGPT (OpenAI) - Data Extraction with Generator Expressions
		# Date: 2025-10-18
		# Prompt: "I have a list of dictionaries with 'status' and 'count' keys. I need to extract 
		# counts for specific statuses (completed, in_progress, pending) and default to 0 if not found. 
		# Can you show me how to use next() with generator expressions and default values?"
		# ChatGPT provided the pattern using next() with generator expressions to extract values 
		# from a list of dictionaries, with default values if the status is not found.
		total_tasks = sum(s['count'] for s in status_overview)
		completed_tasks = next((s['count'] for s in status_overview if s['status'] == 'completed'), 0)
		in_progress_tasks = next((s['count'] for s in status_overview if s['status'] == 'in_progress'), 0)
		pending_tasks = next((s['count'] for s in status_overview if s['status'] == 'pending'), 0)
		
		# Reference: PostgreSQL Documentation - Date/Time Functions
		# https://www.postgresql.org/docs/current/functions-datetime.html#FUNCTIONS-DATETIME-TRUNC
		# DATE_TRUNC groups dates by week for weekly analytics
		weekly_data = sb_fetch_all("""
			SELECT 
				DATE_TRUNC('week', completed_at) as week,
				COUNT(*) as completions
			FROM tasks
			WHERE student_id = :student_id 
			  AND status = 'completed' 
			  AND completed_at IS NOT NULL
			  AND (due_date IS NULL OR due_date >= :cutoff_date)
			GROUP BY DATE_TRUNC('week', completed_at)
			ORDER BY week
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		max_weekly_completions = max((w['completions'] for w in weekly_data), default=1)
		
		# Reference: ChatGPT (OpenAI) - Analytics SQL with Conditional Aggregation
		# Date: 2025-10-18
		# Prompt: "I need SQL queries for analytics in PostgreSQL. I need to count completed tasks 
		# that were on-time (completed_at <= due_date) vs late. Can you provide the SQL with proper 
		# CASE WHEN statements for conditional counting?"
		# ChatGPT provided the SQL query using CASE WHEN for conditional aggregation to distinguish 
		# between on-time and late completions.
		completion_stats = sb_fetch_one("""
			SELECT 
				COUNT(*) as total_completed,
				SUM(CASE WHEN completed_at <= due_date THEN 1 ELSE 0 END) as on_time,
				SUM(CASE WHEN completed_at > due_date THEN 1 ELSE 0 END) as late
			FROM tasks 
			WHERE student_id = :student_id 
			  AND status = 'completed' 
			  AND completed_at IS NOT NULL
			  AND (due_date IS NULL OR due_date >= :cutoff_date)
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		if completion_stats and completion_stats['total_completed'] > 0:
			on_time_percentage = round((completion_stats['on_time'] / completion_stats['total_completed']) * 100, 1)
			late_percentage = round((completion_stats['late'] / completion_stats['total_completed']) * 100, 1)
			completion_stats['on_time_percentage'] = on_time_percentage
			completion_stats['late_percentage'] = late_percentage
		else:
			on_time_percentage = None

		# Reference: ChatGPT (OpenAI) - Lecture Attendance Analytics Queries
		# Date: 2026-02-11
		# Prompt: "I need SQL for lecture attendance analytics: overall attendance %,
		# module-level attendance %, and a recent lecture session list with attended
		# flags for clickable toggles. Can you provide PostgreSQL queries?"
		# ChatGPT provided the query patterns below.
		lecture_overview = sb_fetch_one(
			"""
			SELECT
				COUNT(e.id) AS total_lectures,
				SUM(CASE WHEN COALESCE(la.attended, FALSE) THEN 1 ELSE 0 END) AS attended_lectures
			FROM events e
			LEFT JOIN lecture_attendance la
			       ON la.event_id = e.id
			      AND la.student_id = :student_id
			WHERE e.student_id = :student_id
			  AND e.canvas_event_id IS NOT NULL
			  AND e.start_at <= NOW()
			""",
			{"student_id": current_user.id},
		) or {"total_lectures": 0, "attended_lectures": 0}
		total_lectures = int(lecture_overview.get("total_lectures") or 0)
		attended_lectures = int(lecture_overview.get("attended_lectures") or 0)
		lecture_attendance_rate = round((attended_lectures * 100.0 / total_lectures), 1) if total_lectures else None

		lecture_sessions_recent = sb_fetch_all(
			"""
			SELECT e.id, e.title, e.start_at, e.end_at, e.module_id, m.code AS module_code,
			       COALESCE(la.attended, FALSE) AS attended
			FROM events e
			LEFT JOIN lecture_attendance la
			       ON la.event_id = e.id
			      AND la.student_id = :student_id
			LEFT JOIN modules m ON m.id = e.module_id
			WHERE e.student_id = :student_id
			  AND e.canvas_event_id IS NOT NULL
			  AND e.start_at >= NOW() - INTERVAL '8 weeks'
			  AND e.start_at <= NOW() + INTERVAL '2 weeks'
			ORDER BY e.start_at DESC
			LIMIT 80
			""",
			{"student_id": current_user.id},
		)

		module_attendance = sb_fetch_all(
			"""
			SELECT m.id AS module_id, m.code AS module_code,
			       COUNT(e.id) AS total_lectures,
			       SUM(CASE WHEN COALESCE(la.attended, FALSE) THEN 1 ELSE 0 END) AS attended_lectures
			FROM modules m
			LEFT JOIN events e
			       ON e.module_id = m.id
			      AND e.student_id = :student_id
			      AND e.canvas_event_id IS NOT NULL
			      AND e.start_at <= NOW()
			LEFT JOIN lecture_attendance la
			       ON la.event_id = e.id
			      AND la.student_id = :student_id
			GROUP BY m.id, m.code
			HAVING COUNT(e.id) > 0
			ORDER BY m.code
			""",
			{"student_id": current_user.id},
		)
		for row in module_attendance:
			total_for_module = int(row.get("total_lectures") or 0)
			attended_for_module = int(row.get("attended_lectures") or 0)
			row["attendance_rate"] = round((attended_for_module * 100.0 / total_for_module), 1) if total_for_module else 0.0

		task_completion_rate = round((completed_tasks * 100.0 / total_tasks), 1) if total_tasks else None
		health_score = _calculate_health_score(
			task_completion_rate=task_completion_rate,
			on_time_rate=on_time_percentage,
			lecture_attendance_rate=lecture_attendance_rate,
		)
		
		# Reference: ChatGPT (OpenAI) - Module Performance Analytics with NULLIF
		# Date: 2025-10-18
		# Prompt: "I need to calculate completion rates per module using LEFT JOIN in PostgreSQL, 
		# handling division by zero with NULLIF. The query should group by module and only show 
		# modules that have tasks. Can you provide the SQL with proper NULLIF for safe division?"
		# ChatGPT provided the SQL query with LEFT JOIN, NULLIF to prevent division by zero, 
		# and HAVING clause to filter modules with tasks.
		module_stats = sb_fetch_all("""
			SELECT 
				m.code as module_code,
				COUNT(t.id) as total_tasks,
				SUM(CASE WHEN t.status = 'completed' THEN 1 ELSE 0 END) as completed,
				ROUND(
					(SUM(CASE WHEN t.status = 'completed' THEN 1 ELSE 0 END)::numeric / NULLIF(COUNT(t.id), 0)) * 100, 
					1
				) as completion_rate
			FROM modules m
			LEFT JOIN tasks t ON m.id = t.module_id 
				AND t.student_id = :student_id
				AND (t.due_date IS NULL OR t.due_date >= :cutoff_date)
			GROUP BY m.id, m.code
			HAVING COUNT(t.id) > 0
			ORDER BY completion_rate DESC
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		# 1. Grade Performance Analytics (using Canvas scores)
		# Note: If canvas_possible is 0 or NULL, treat canvas_score as the percentage directly

		grade_performance = sb_fetch_all("""
			SELECT 
				m.code as module_code,
				COUNT(t.id) as graded_count,
				AVG(CASE 
					WHEN t.canvas_possible > 0 
					THEN (t.canvas_score::numeric / t.canvas_possible::numeric) * 100 
					ELSE t.canvas_score::numeric
				END) as avg_percentage,
				AVG(t.canvas_score) as avg_score,
				AVG(CASE WHEN t.canvas_possible > 0 THEN t.canvas_possible ELSE 100 END) as avg_possible
			FROM modules m
			LEFT JOIN tasks t ON m.id = t.module_id 
				AND t.student_id = :student_id
				AND t.canvas_score IS NOT NULL
			GROUP BY m.id, m.code
			HAVING COUNT(t.id) > 0
			ORDER BY avg_percentage DESC NULLS LAST
		""", {"student_id": current_user.id})
		
		# Overall grade stats
		# Note: If canvas_possible is 0 or NULL, treat canvas_score as the percentage directly
		overall_grade_stats = sb_fetch_one("""
			SELECT 
				COUNT(*) as total_graded,
				AVG(CASE 
					WHEN canvas_possible > 0 
					THEN (canvas_score::numeric / canvas_possible::numeric) * 100 
					ELSE canvas_score::numeric
				END) as overall_avg_percentage
			FROM tasks 
			WHERE student_id = :student_id 
			  AND canvas_score IS NOT NULL
		""", {"student_id": current_user.id})
		
		# Predicted overall grade (weighted by assignment weights)

		predicted_grade = sb_fetch_one("""
			SELECT 
				SUM(CASE 
					WHEN weight_percentage IS NOT NULL AND weight_percentage > 0 AND canvas_score IS NOT NULL
					THEN (
						CASE 
							WHEN canvas_possible > 0 
							THEN (canvas_score::numeric / canvas_possible::numeric) * 100
							ELSE canvas_score::numeric
						END
					) * weight_percentage::numeric
					ELSE NULL
				END) as weighted_grade_sum,
				SUM(CASE 
					WHEN weight_percentage IS NOT NULL AND weight_percentage > 0 AND canvas_score IS NOT NULL
					THEN weight_percentage
					ELSE NULL
				END) as total_weight
			FROM tasks 
			WHERE student_id = :student_id 
			  AND canvas_score IS NOT NULL
			  AND weight_percentage IS NOT NULL
			  AND weight_percentage > 0
		""", {"student_id": current_user.id})
		
		# Individual graded assignments with scores
		# Note: If canvas_possible is 0 or NULL, treat canvas_score as the percentage directly
		individual_grades = sb_fetch_all("""
			SELECT 
				t.id,
				t.title,
				m.code as module_code,
				t.canvas_score,
				CASE WHEN t.canvas_possible > 0 THEN t.canvas_possible ELSE 100 END as canvas_possible,
				CASE 
					WHEN t.canvas_possible > 0 
					THEN ROUND((t.canvas_score::numeric / t.canvas_possible::numeric) * 100, 1)
					ELSE ROUND(t.canvas_score::numeric, 1)
				END as percentage,
				t.canvas_graded_at,
				t.weight_percentage,
				t.due_date
			FROM tasks t
			LEFT JOIN modules m ON t.module_id = m.id
			WHERE t.student_id = :student_id 
			  AND t.canvas_score IS NOT NULL
			ORDER BY t.canvas_graded_at DESC NULLS LAST, t.due_date DESC NULLS LAST
		""", {"student_id": current_user.id})
		
		# 3. Priority & Weighting Analytics
		# High priority tasks (high weight + due soon)
		high_priority_tasks = sb_fetch_all("""
			SELECT 
				t.id, t.title, t.status, t.due_date, t.due_at,
				t.weight_percentage, m.code as module_code,
				CASE 
					WHEN t.due_at IS NOT NULL 
					THEN EXTRACT(EPOCH FROM (t.due_at - NOW())) / 3600
					WHEN t.due_date IS NOT NULL
					THEN EXTRACT(EPOCH FROM (t.due_date::timestamp - NOW())) / 3600
					ELSE NULL
				END as hours_remaining
			FROM tasks t
			JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			  AND t.status != 'completed'
			  AND (t.due_date IS NULL OR t.due_date >= :cutoff_date)
			ORDER BY 
				t.weight_percentage DESC NULLS LAST,
				t.due_at ASC NULLS LAST,
				t.due_date ASC NULLS LAST
			LIMIT 10
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		# Calculate priority scores for high priority tasks
		from services.analytics import calculate_priority, normalise_due_datetime
		from datetime import timezone
		now = datetime.now(timezone.utc)
		high_priority_with_scores = []
		for task in high_priority_tasks:
			due_at = normalise_due_datetime(task, now)
			if due_at:
				hours_remaining = (due_at - now).total_seconds() / 3600
				# Convert weight_percentage to float if it's a Decimal
				weight = task.get('weight_percentage')
				if weight is not None:
					try:
						weight = float(weight)
					except (TypeError, ValueError):
						weight = None
				priority = calculate_priority(
					weight, 
					max(0, hours_remaining)
				)
				task['priority'] = priority
				high_priority_with_scores.append(task)
		
		high_priority_with_scores.sort(key=lambda x: x.get('priority', 0), reverse=True)
		
		# Weight distribution stats
		weight_stats = sb_fetch_one("""
			SELECT 
				COUNT(*) as total_with_weight,
				AVG(weight_percentage) as avg_weight,
				MAX(weight_percentage) as max_weight,
				MIN(weight_percentage) as min_weight
			FROM tasks
			WHERE student_id = :student_id
			  AND weight_percentage IS NOT NULL
			  AND (due_date IS NULL OR due_date >= :cutoff_date)
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		# 7. Actionable Insights
		# Focus module (lowest completion rate)
		focus_module = next((m for m in module_stats if m['completion_rate'] < 100), None)
		
		# At-risk assignments (high weight + due within 7 days)
		at_risk_assignments = sb_fetch_all("""
			SELECT 
				t.id, t.title, t.status, t.due_date, t.due_at,
				t.weight_percentage, m.code as module_code,
				EXTRACT(EPOCH FROM (t.due_date::timestamp - NOW())) / 86400 as days_remaining
			FROM tasks t
			JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			  AND t.status != 'completed'
			  AND t.weight_percentage IS NOT NULL
			  AND t.weight_percentage >= 10
			  AND t.due_date >= :cutoff_date
			  AND t.due_date <= :cutoff_date + INTERVAL '7 days'
			ORDER BY t.weight_percentage DESC, t.due_date ASC
			LIMIT 5
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		# Overdue count by module
		overdue_by_module = sb_fetch_all("""
			SELECT 
				m.code as module_code,
				COUNT(t.id) as overdue_count
			FROM modules m
			LEFT JOIN tasks t ON m.id = t.module_id 
				AND t.student_id = :student_id
				AND t.status != 'completed'
				AND t.due_date < :cutoff_date
				AND (t.due_date IS NULL OR t.due_date >= :cutoff_date - INTERVAL '90 days')
			GROUP BY m.id, m.code
			HAVING COUNT(t.id) > 0
			ORDER BY overdue_count DESC
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})
		
		# Calculate predicted grade percentage
		predicted_grade_pct = None
		if predicted_grade and predicted_grade.get('weighted_grade_sum') and predicted_grade.get('total_weight'):
			if predicted_grade['total_weight'] > 0:
				predicted_grade_pct = round((predicted_grade['weighted_grade_sum'] / predicted_grade['total_weight']), 1)

		# ── Additional Visualisation Data (Iteration 5 – US 18) ──────────
		# Reference: ChatGPT (OpenAI) - Analytics Visualisation Queries
		# Date: 2026-02-10
		# Prompt: "I need extra analytics queries for new Chart.js visuals:
		# 1) daily task due counts for the next 30 days (workload heatmap),
		# 2) cumulative completed tasks over time (progress line),
		# 3) per-module completion percentages for a radar chart.
		# Can you provide compact PostgreSQL queries?"
		# ChatGPT provided the queries below.

		# Daily task density (next 30 days)
		daily_task_density_rows = sb_fetch_all("""
			SELECT due_date::date AS day, COUNT(*) AS task_count
			FROM tasks
			WHERE student_id = :student_id
			  AND due_date >= :cutoff_date
			  AND due_date <= :cutoff_date + INTERVAL '30 days'
			GROUP BY due_date::date
			ORDER BY day
		""", {"student_id": current_user.id, "cutoff_date": cutoff_date})

		# Cumulative completed tasks over time
		cumulative_rows = sb_fetch_all("""
			SELECT completed_at::date AS day,
			       COUNT(*) AS daily_completed,
			       SUM(COUNT(*)) OVER (ORDER BY completed_at::date) AS cumulative
			FROM tasks
			WHERE student_id = :student_id
			  AND status = 'completed'
			  AND completed_at IS NOT NULL
			GROUP BY completed_at::date
			ORDER BY day
		""", {"student_id": current_user.id})

		# Reference: ChatGPT (OpenAI) - Dense Time-Series Chart Preparation
		# Date: 2026-02-10
		# Prompt: "My workload chart only shows dates that have tasks. I want to
		# render a stable 30-day timeline including zero-value days so the chart
		# is easier to read. Can you show a Python pattern to densify SQL rows?"
		# ChatGPT provided the dictionary-lookup + fixed-range expansion pattern.
		density_map = {}
		for row in daily_task_density_rows:
			day_value = row.get("day")
			day_key = day_value.isoformat() if hasattr(day_value, "isoformat") else str(day_value)
			density_map[day_key] = int(row.get("task_count") or 0)
		daily_task_density = []
		for offset in range(31):
			day = cutoff_date + timedelta(days=offset)
			day_key = day.isoformat()
			daily_task_density.append({
				"day": day_key,
				"task_count": density_map.get(day_key, 0),
			})

		cumulative_completions = []
		for row in cumulative_rows:
			day_value = row.get("day")
			cumulative_completions.append({
				"day": day_value.isoformat() if hasattr(day_value, "isoformat") else str(day_value),
				"daily_completed": int(row.get("daily_completed") or 0),
				"cumulative": int(row.get("cumulative") or 0),
			})

		analytics_data = {
			'total_tasks': total_tasks,
			'completed_tasks': completed_tasks,
			'in_progress_tasks': in_progress_tasks,
			'pending_tasks': pending_tasks,
			'weekly_data': weekly_data,
			'max_weekly_completions': max_weekly_completions,
			'completion_stats': completion_stats,
			'module_stats': module_stats,
			# Grade Performance
			'grade_performance': grade_performance,
			'overall_grade_stats': overall_grade_stats,
			'predicted_grade_pct': predicted_grade_pct,
			'individual_grades': individual_grades,
			# Priority & Weighting
			'high_priority_tasks': high_priority_with_scores[:5],
			'weight_stats': weight_stats,
			# Actionable Insights
			'focus_module': focus_module,
			'at_risk_assignments': at_risk_assignments,
			'overdue_by_module': overdue_by_module,
			# Extra Visuals (US 18)
			'daily_task_density': daily_task_density,
			'cumulative_completions': cumulative_completions,
			'lecture_attendance': {
				'total_lectures': total_lectures,
				'attended_lectures': attended_lectures,
				'attendance_rate': lecture_attendance_rate,
			},
			'lecture_sessions_recent': lecture_sessions_recent,
			'module_attendance': module_attendance,
			'task_completion_rate': task_completion_rate,
			'health_score': health_score,
		}
		
	except Exception as e:
		print(f"Error fetching analytics data: {e}")
		import traceback
		traceback.print_exc()
		analytics_data = {
			'total_tasks': 0,
			'completed_tasks': 0,
			'in_progress_tasks': 0,
			'pending_tasks': 0,
			'weekly_data': [],
			'max_weekly_completions': 1,
			'completion_stats': None,
			'module_stats': [],
			'grade_performance': [],
			'overall_grade_stats': None,
			'predicted_grade_pct': None,
			'individual_grades': [],
			'high_priority_tasks': [],
			'weight_stats': None,
			'focus_module': None,
			'at_risk_assignments': [],
			'overdue_by_module': [],
			'daily_task_density': [],
			'cumulative_completions': [],
			'lecture_attendance': {
				'total_lectures': 0,
				'attended_lectures': 0,
				'attendance_rate': None,
			},
			'lecture_sessions_recent': [],
			'module_attendance': [],
			'task_completion_rate': None,
			'health_score': {'score': 0.0, 'label': 'no_data'},
		}

	return render_template("analytics.html", analytics=analytics_data)


# ── Per-Module Dashboard (Iteration 5 – US 17) ─────────────────────
# Reference: ChatGPT (OpenAI) - Per-Module Dashboard Routes
# Date: 2026-02-10
# Prompt: "I need a modules overview page showing all modules with task count,
# completion %, and average grade, plus a detail page for a single module
# showing all tasks, grade breakdown, and progress chart. Can you provide
# the Flask routes and SQL queries?"
# ChatGPT provided the route structure and SQL queries below.
@app.route("/modules")
@login_required
def modules_overview():
	"""List all modules with summary stats."""
	# Reference: ChatGPT (OpenAI) - Module Fallback Matching from Event Title
	# Date: 2026-02-11
	# Prompt: "Some Canvas lecture events may not have module_id populated, but the
	# module code appears in the event title (e.g., IS4408 ...). I need SQL fallback
	# logic so module attendance still aggregates correctly. Can you provide a safe
	# PostgreSQL pattern?"
	# ChatGPT provided the regex extraction fallback used in attendance subqueries.
	try:
		modules = sb_fetch_all("""
			SELECT m.id, m.code,
			       COUNT(t.id) AS total_tasks,
			       SUM(CASE WHEN t.status = 'completed' THEN 1 ELSE 0 END) AS completed,
			       ROUND(
			           SUM(CASE WHEN t.status = 'completed' THEN 1 ELSE 0 END) * 100.0
			           / NULLIF(COUNT(t.id), 0), 1
			       ) AS completion_rate,
			       ROUND(AVG(
			           CASE
			               WHEN t.canvas_score IS NOT NULL AND t.canvas_possible > 0
			               THEN (t.canvas_score / t.canvas_possible) * 100
			               WHEN t.canvas_score IS NOT NULL AND (t.canvas_possible IS NULL OR t.canvas_possible = 0)
			               THEN t.canvas_score
			           END
			       ), 1) AS avg_grade
			       ,
			       COALESCE((
			           SELECT COUNT(*)
			           FROM events e
			           WHERE (
			               e.module_id = m.id
			               OR (
			                   e.module_id IS NULL
			                   AND (
			                       UPPER(e.title) LIKE '%' || UPPER(m.code) || '%'
			                       OR UPPER(e.title) LIKE '%' || UPPER(REGEXP_REPLACE(m.code, '^[0-9]{4}-', '')) || '%'
			                   )
			               )
			           )
			             AND e.student_id = :sid
			             AND e.canvas_event_id IS NOT NULL
			             AND e.start_at <= NOW()
			       ), 0) AS lecture_sessions,
			       COALESCE((
			           SELECT COUNT(*)
			           FROM lecture_attendance la
			           JOIN events e2 ON e2.id = la.event_id
			           WHERE la.student_id = :sid
			             AND la.attended = TRUE
			             AND (
			                 e2.module_id = m.id
			                 OR (
			                     e2.module_id IS NULL
			                     AND (
			                         UPPER(e2.title) LIKE '%' || UPPER(m.code) || '%'
			                         OR UPPER(e2.title) LIKE '%' || UPPER(REGEXP_REPLACE(m.code, '^[0-9]{4}-', '')) || '%'
			                     )
			                 )
			             )
			             AND e2.student_id = :sid
			             AND e2.canvas_event_id IS NOT NULL
			             AND e2.start_at <= NOW()
			       ), 0) AS attended_lectures
			FROM modules m
			LEFT JOIN tasks t ON t.module_id = m.id AND t.student_id = :sid
			GROUP BY m.id, m.code
			ORDER BY m.code
		""", {"sid": current_user.id})
	except Exception as exc:
		print(f"[modules] overview error: {exc}")
		try:
			# Fallback if lecture_attendance table is not migrated yet.
			modules = sb_fetch_all("""
				SELECT m.id, m.code,
				       COUNT(t.id) AS total_tasks,
				       SUM(CASE WHEN t.status = 'completed' THEN 1 ELSE 0 END) AS completed,
				       ROUND(
				           SUM(CASE WHEN t.status = 'completed' THEN 1 ELSE 0 END) * 100.0
				           / NULLIF(COUNT(t.id), 0), 1
				       ) AS completion_rate,
				       ROUND(AVG(
				           CASE
				               WHEN t.canvas_score IS NOT NULL AND t.canvas_possible > 0
				               THEN (t.canvas_score / t.canvas_possible) * 100
				               WHEN t.canvas_score IS NOT NULL AND (t.canvas_possible IS NULL OR t.canvas_possible = 0)
				               THEN t.canvas_score
				           END
				       ), 1) AS avg_grade,
				       0 AS lecture_sessions,
				       0 AS attended_lectures
				FROM modules m
				LEFT JOIN tasks t ON t.module_id = m.id AND t.student_id = :sid
				GROUP BY m.id, m.code
				ORDER BY m.code
			""", {"sid": current_user.id})
		except Exception:
			modules = []
	for item in modules:
		total_lectures = int(item.get("lecture_sessions") or 0)
		attended_lectures = int(item.get("attended_lectures") or 0)
		item["attendance_rate"] = round((attended_lectures * 100.0 / total_lectures), 1) if total_lectures else None
	return render_template("modules_overview.html", modules=modules)


@app.route("/modules/<int:module_id>")
@login_required
def module_detail(module_id):
	"""Detailed view for a single module."""
	try:
		module = sb_fetch_one("SELECT id, code FROM modules WHERE id = :mid", {"mid": module_id})
	except Exception:
		module = None
	if not module:
		flash("Module not found.", "error")
		return redirect(url_for("modules_overview"))

	try:
		tasks = sb_fetch_all("""
			SELECT id, title, status, due_date, due_at, weight_percentage,
			       canvas_score, canvas_possible, canvas_graded_at
			FROM tasks
			WHERE student_id = :sid AND module_id = :mid
			ORDER BY due_date ASC NULLS LAST
		""", {"sid": current_user.id, "mid": module_id})
	except Exception:
		tasks = []

	total = len(tasks)
	completed = sum(1 for t in tasks if t.get("status") == "completed")
	in_progress = sum(1 for t in tasks if t.get("status") == "in_progress")
	pending = sum(1 for t in tasks if t.get("status") == "pending")
	completion_rate = round(completed * 100.0 / total, 1) if total else 0

	# Grade stats
	graded = [t for t in tasks if t.get("canvas_score") is not None]
	avg_grade = None
	if graded:
		pcts = []
		for t in graded:
			cp = t.get("canvas_possible") or 0
			cs = t.get("canvas_score") or 0
			pcts.append((cs / cp * 100) if cp > 0 else cs)
		avg_grade = round(sum(pcts) / len(pcts), 1) if pcts else None

	# Weekly completions for this module
	try:
		weekly = sb_fetch_all("""
			SELECT DATE_TRUNC('week', completed_at) AS week, COUNT(*) AS completions
			FROM tasks
			WHERE student_id = :sid AND module_id = :mid
			  AND status = 'completed' AND completed_at IS NOT NULL
			GROUP BY DATE_TRUNC('week', completed_at)
			ORDER BY week
		""", {"sid": current_user.id, "mid": module_id})
	except Exception:
		weekly = []

	module_code_value = (module.get("code") or "").strip()
	module_code_short = re.sub(r"^[0-9]{4}-", "", module_code_value).strip() or module_code_value

	try:
		module_lecture_summary = sb_fetch_one(
			"""
			SELECT
				COUNT(e.id) AS total_lectures,
				SUM(CASE WHEN COALESCE(la.attended, FALSE) THEN 1 ELSE 0 END) AS attended_lectures
			FROM events e
			LEFT JOIN lecture_attendance la
			       ON la.event_id = e.id
			      AND la.student_id = :sid
			WHERE e.student_id = :sid
			  AND (
			      e.module_id = :mid
			      OR (
			          e.module_id IS NULL
			          AND (
			              UPPER(e.title) LIKE '%' || UPPER(:module_code) || '%'
			              OR UPPER(e.title) LIKE '%' || UPPER(:module_code_short) || '%'
			          )
			      )
			  )
			  AND e.canvas_event_id IS NOT NULL
			  AND e.start_at <= NOW()
			""",
			{
				"sid": current_user.id,
				"mid": module_id,
				"module_code": module_code_value,
				"module_code_short": module_code_short,
			},
		) or {"total_lectures": 0, "attended_lectures": 0}
	except Exception:
		module_lecture_summary = {"total_lectures": 0, "attended_lectures": 0}

	try:
		module_lecture_sessions = sb_fetch_all(
			"""
			SELECT e.id, e.title, e.start_at,
			       COALESCE(la.attended, FALSE) AS attended
			FROM events e
			LEFT JOIN lecture_attendance la
			       ON la.event_id = e.id
			      AND la.student_id = :sid
			WHERE e.student_id = :sid
			  AND (
			      e.module_id = :mid
			      OR (
			          e.module_id IS NULL
			          AND (
			              UPPER(e.title) LIKE '%' || UPPER(:module_code) || '%'
			              OR UPPER(e.title) LIKE '%' || UPPER(:module_code_short) || '%'
			          )
			      )
			  )
			  AND e.canvas_event_id IS NOT NULL
			  AND e.start_at >= NOW() - INTERVAL '8 weeks'
			  AND e.start_at <= NOW() + INTERVAL '2 weeks'
			ORDER BY e.start_at DESC
			LIMIT 40
			""",
			{
				"sid": current_user.id,
				"mid": module_id,
				"module_code": module_code_value,
				"module_code_short": module_code_short,
			},
		)
	except Exception:
		module_lecture_sessions = []

	module_total_lectures = int(module_lecture_summary.get("total_lectures") or 0)
	module_attended_lectures = int(module_lecture_summary.get("attended_lectures") or 0)
	module_attendance_rate = round((module_attended_lectures * 100.0 / module_total_lectures), 1) if module_total_lectures else None

	return render_template("module_detail.html",
		module=module, tasks=tasks,
		total=total, completed=completed, in_progress=in_progress, pending=pending,
		completion_rate=completion_rate, graded_count=len(graded),
		avg_grade=avg_grade, weekly=weekly,
		module_total_lectures=module_total_lectures,
		module_attended_lectures=module_attended_lectures,
		module_attendance_rate=module_attendance_rate,
		module_lecture_sessions=module_lecture_sessions,
	)


@app.route("/study-planner", methods=["GET", "POST"])
@login_required
def study_planner():
	"""AI Study Planner - Generate personalized study recommendations"""
	# Reference: ChatGPT (OpenAI) - Study Planner Prompt Context
	# Date: 2026-01-22
	# Prompt: "I need to summarize tasks and progress into a compact prompt for study
	# recommendations. Can you help structure the summary text?"
	# ChatGPT provided the prompt context structure.
	recommendations: Optional[str] = None
	error: Optional[str] = None
	
	# Fetch tasks data
	try:
		task_rows = sb_fetch_all(
			"""
			SELECT t.id, t.title, t.status, t.due_date, t.due_at, t.weight_percentage,
			       m.code AS module_code
			FROM tasks t
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			ORDER BY t.due_at NULLS LAST, t.due_date NULLS LAST
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[study-planner] failed to load tasks user={current_user.id} error={exc}")
		task_rows = []
	
	# Build progress summary using assess_progress
	progress = assess_progress(task_rows)
	
	# Build tasks summary
	tasks_summary_lines = []
	if task_rows:
		pending = [t for t in task_rows if t.get("status") == "pending"]
		in_progress = [t for t in task_rows if t.get("status") == "in_progress"]
		completed = [t for t in task_rows if t.get("status") == "completed"]
		
		tasks_summary_lines.append(f"Total tasks: {len(task_rows)}")
		tasks_summary_lines.append(f"Pending: {len(pending)}, In Progress: {len(in_progress)}, Completed: {len(completed)}")
		
		if pending or in_progress:
			tasks_summary_lines.append("\nUpcoming/Active tasks:")
			for task in (pending + in_progress)[:10]:  # Limit to 10 for prompt size
				module = task.get("module_code") or "No module"
				due = task.get("due_at") or task.get("due_date") or "No due date"
				weight = task.get("weight_percentage")
				weight_str = f" ({weight}% weight)" if weight else ""
				tasks_summary_lines.append(f"- {task.get('title')} ({module}){weight_str} - Due: {due}")
	else:
		tasks_summary_lines.append("No tasks found.")
	
	tasks_summary = "\n".join(tasks_summary_lines)
	
	# Build progress summary
	progress_summary_lines = [
		f"Status: {progress.get('status', 'unknown').replace('_', ' ').title()}",
		f"Overdue tasks: {progress.get('overdue', 0)}",
		f"Due in 48 hours: {progress.get('nearly_due', 0)}",
		f"Completed this week: {progress.get('completed_this_week', 0)}"
	]
	progress_summary = "\n".join(progress_summary_lines)
	
	# Generate recommendations if POST request
	if request.method == "POST":
		try:
			chatgpt = get_chatgpt_service()
			recommendations = chatgpt.get_study_recommendations(
				tasks_summary=tasks_summary,
				progress_summary=progress_summary
			)
		except ChatGPTClientError as exc:
			error = f"Failed to generate recommendations: {str(exc)}"
			print(f"[study-planner] ChatGPT error: {exc}")
		except Exception as exc:
			error = f"Unexpected error: {str(exc)}"
			print(f"[study-planner] Unexpected error: {exc}")
			import traceback
			traceback.print_exc()
	
	return render_template(
		"study_planner.html",
		recommendations=recommendations,
		error=error,
		tasks_summary=tasks_summary,
		progress_summary=progress_summary
	)


@app.route("/course-bot", methods=["GET", "POST"])
@login_required
def course_bot():
	"""AI Course Bot - Ask questions based on uploaded course materials."""
	# Reference: ChatGPT (OpenAI) - Course Bot Upload + Q&A Flow
	# Date: 2026-01-22
	# Prompt: "I want a course bot that stores uploaded materials and answers questions using
	# only those sources. Can you outline the flow and storage steps?"
	# ChatGPT provided the upload + Q&A orchestration pattern.
	error: Optional[str] = None
	success_message: Optional[str] = None

	if request.method == "POST":
		action = request.form.get("action", "").strip()
		if action == "upload_course_doc":
			file_storage = request.files.get("course_file")
			if not file_storage or not file_storage.filename:
				error = "Please upload a course file."
			else:
				filename = file_storage.filename
				try:
					payload = file_storage.read()
					if not payload:
						raise ValueError("The uploaded file was empty.")
					if len(payload) > _AI_MAX_UPLOAD_BYTES:
						raise ValueError("The file is larger than 4 MB. Upload a smaller file.")

					extracted_text = _extract_text_from_brief(filename, payload)
					if not extracted_text or not extracted_text.strip():
						raise ValueError("Could not extract text from the uploaded file.")

					uploads_dir = os.path.join(app.root_path, "uploads", "course_packs")
					os.makedirs(uploads_dir, exist_ok=True)
					timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
					safe_filename = f"{current_user.id}_{timestamp}_{filename}"
					filepath = os.path.join(uploads_dir, safe_filename)
					with open(filepath, "wb") as f:
						f.write(payload)

					sb_execute(
						"""
						INSERT INTO course_documents (
							student_id, filename, filepath, extracted_text, uploaded_at
						) VALUES (
							:student_id, :filename, :filepath, :extracted_text, NOW()
						)
						""",
						{
							"student_id": current_user.id,
							"filename": filename,
							"filepath": filepath,
							"extracted_text": extracted_text,
						}
					)
					success_message = "Course material uploaded successfully."
				except ValueError as exc:
					error = str(exc)
					print(f"[course-bot] upload validation error user={current_user.id} err={exc}")
				except Exception as exc:
					error = f"Failed to process upload: {str(exc)}"
					print(f"[course-bot] upload failed user={current_user.id} err={exc}")
					import traceback
					traceback.print_exc()
		elif action == "clear_course_docs":
			try:
				rows = sb_fetch_all(
					"""
					SELECT id, filepath
					FROM course_documents
					WHERE student_id = :student_id
					""",
					{"student_id": current_user.id}
				)
				sb_execute(
					"""
					DELETE FROM course_documents
					WHERE student_id = :student_id
					""",
					{"student_id": current_user.id}
				)
				for row in rows:
					filepath = row.get("filepath")
					if filepath and os.path.exists(filepath):
						try:
							os.remove(filepath)
						except Exception:
							pass
				success_message = "Uploaded materials cleared."
			except Exception as exc:
				error = f"Failed to clear materials: {str(exc)}"
				print(f"[course-bot] clear docs failed user={current_user.id} err={exc}")
		elif action == "ask_course_question":
			question = request.form.get("question", "").strip()
			if not question:
				error = "Please enter a question."
			else:
				try:
					documents = sb_fetch_all(
						"""
						SELECT id, filename, extracted_text, uploaded_at
						FROM course_documents
						WHERE student_id = :student_id
						ORDER BY uploaded_at DESC
						""",
						{"student_id": current_user.id}
					)
					if not documents:
						error = "Upload course materials before asking a question."
					else:
						sources = _select_course_materials(documents, question)
						if not sources:
							error = "No relevant course materials found for this question."
						else:
							service = get_chatgpt_service()
							response = service.answer_course_question(
								question=question,
								sources=sources
							)
							citations_payload = [
								{"source": c.source, "quote": c.quote} for c in response.citations
							]
							sb_execute(
								"""
								INSERT INTO course_bot_history (
									student_id, question, answer, citations, asked_at
								) VALUES (
									:student_id, :question, :answer, :citations, NOW()
								)
								""",
								{
									"student_id": current_user.id,
									"question": question,
									"answer": response.answer,
									"citations": json.dumps(citations_payload),
								}
							)
							success_message = "Answer generated. See the history below."
				except ChatGPTClientError as exc:
					error = f"Failed to answer question: {str(exc)}"
					print(f"[course-bot] ChatGPT error user={current_user.id} err={exc}")
				except Exception as exc:
					error = f"Unexpected error: {str(exc)}"
					print(f"[course-bot] Unexpected error user={current_user.id} err={exc}")
					import traceback
					traceback.print_exc()
		elif action == "clear_course_history":
			try:
				sb_execute(
					"""
					DELETE FROM course_bot_history
					WHERE student_id = :student_id
					""",
					{"student_id": current_user.id}
				)
				success_message = "Q&A history cleared."
			except Exception as exc:
				error = f"Failed to clear history: {str(exc)}"
				print(f"[course-bot] clear history failed user={current_user.id} err={exc}")
		else:
			error = "Invalid request."

	docs = []
	history = []
	try:
		docs = sb_fetch_all(
			"""
			SELECT id, filename, uploaded_at
			FROM course_documents
			WHERE student_id = :student_id
			ORDER BY uploaded_at DESC
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[course-bot] failed to load docs user={current_user.id} error={exc}")

	try:
		history = sb_fetch_all(
			"""
			SELECT id, question, answer, citations, asked_at
			FROM course_bot_history
			WHERE student_id = :student_id
			ORDER BY asked_at DESC
			LIMIT 20
			""",
			{"student_id": current_user.id}
		)
		for item in history:
			answer_text = item.get("answer")
			item["qa_pairs"] = []
			if isinstance(answer_text, str) and answer_text.strip().startswith("["):
				try:
					parsed = json.loads(answer_text)
					if isinstance(parsed, list) and all(isinstance(entry, str) for entry in parsed):
						item["answer"] = "\n".join(parsed)
				except Exception:
					pass
			elif isinstance(answer_text, str) and "['" in answer_text and "]" in answer_text:
				try:
					import ast
					parsed = ast.literal_eval(answer_text)
					if isinstance(parsed, list) and all(isinstance(entry, str) for entry in parsed):
						item["answer"] = "\n".join(parsed)
				except Exception:
					pass
			answer_text = item.get("answer")
			if isinstance(answer_text, str) and "| Answer:" in answer_text:
				lines = [line.strip() for line in answer_text.splitlines() if line.strip()]
				for line in lines:
					if "| Answer:" not in line:
						continue
					parts = line.split("| Answer:", 1)
					question_part = parts[0].replace("Q:", "").strip()
					answer_part = parts[1]
					expl = ""
					if "| Explanation:" in answer_part:
						answer_value, expl = answer_part.split("| Explanation:", 1)
						answer_value = answer_value.strip()
						expl = expl.strip()
					else:
						answer_value = answer_part.strip()
					item["qa_pairs"].append({
						"question": question_part,
						"answer": answer_value,
						"explanation": expl,
					})
			raw = item.get("citations")
			try:
				item["citations_list"] = json.loads(raw) if raw else []
			except Exception:
				item["citations_list"] = []
	except Exception as exc:
		print(f"[course-bot] failed to load history user={current_user.id} error={exc}")

	return render_template(
		"course_bot.html",
		docs=docs,
		history=history,
		error=error,
		success_message=success_message
	)


@app.route("/assignment-review", methods=["GET", "POST"])
@login_required
def assignment_review():
	"""AI Assignment Review & Grading - Standalone page"""
	error: Optional[str] = None
	success_message: Optional[str] = None
	
	# Fetch user's tasks for optional linking
	tasks = []
	try:
		tasks = sb_fetch_all(
			"""
			SELECT t.id, t.title, m.code AS module_code
			FROM tasks t
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			ORDER BY t.due_at NULLS LAST, t.due_date NULLS LAST, t.title
			LIMIT 50
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[assignment-review] failed to load tasks user={current_user.id} error={exc}")
	
	# Fetch all reviews for the user
	reviews = []
	try:
		reviews = sb_fetch_all(
			"""
			SELECT ar.id, ar.task_id, ar.filename, ar.ai_feedback, ar.ai_score_estimate, 
			       ar.ai_possible_score, ar.reviewed_at,
			       t.title AS task_title, m.code AS module_code
			FROM assignment_reviews ar
			LEFT JOIN tasks t ON t.id = ar.task_id
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE ar.student_id = :student_id
			ORDER BY ar.reviewed_at DESC
			LIMIT 20
			""",
			{"student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[assignment-review] failed to load reviews user={current_user.id} error={exc}")
	
	if request.method == "POST":
		# Get assignment text from file or text input
		assignment_text: Optional[str] = None
		filename: Optional[str] = None
		filepath: Optional[str] = None
		task_id: Optional[int] = None
		
		# Get optional task_id
		task_id_str = request.form.get("task_id", "").strip()
		if task_id_str:
			try:
				task_id = int(task_id_str)
			except ValueError:
				pass
		
		# Check for file upload
		file_storage = request.files.get("assignment_file")
		if file_storage and file_storage.filename:
			filename = file_storage.filename
			try:
				payload = file_storage.read()
				if not payload:
					raise ValueError("The uploaded file was empty.")
				if len(payload) > _AI_MAX_UPLOAD_BYTES:
					raise ValueError("The file is larger than 4 MB. Upload a smaller file.")
				
				assignment_text = _extract_text_from_brief(filename, payload)
				if not assignment_text:
					raise ValueError("Could not extract text from the uploaded file.")
				
				# Store file on filesystem
				uploads_dir = os.path.join(app.root_path, "uploads", "assignments")
				os.makedirs(uploads_dir, exist_ok=True)
				timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
				safe_filename = f"{current_user.id}_{timestamp}_{filename}"
				filepath = os.path.join(uploads_dir, safe_filename)
				with open(filepath, "wb") as f:
					f.write(payload)
			except ValueError as exc:
				error = str(exc)
				print(f"[assignment-review] file validation error user={current_user.id} err={exc}")
			except Exception as exc:
				error = f"Failed to process uploaded file: {str(exc)}"
				print(f"[assignment-review] file processing failed user={current_user.id} err={exc}")
				import traceback
				traceback.print_exc()
		
		# Check for text input
		if not assignment_text:
			assignment_text = request.form.get("assignment_text", "").strip()
		
		if not assignment_text:
			error = "Please provide assignment content by uploading a file or pasting text."
		
		# Get assignment title and brief (optional)
		assignment_title = "Assignment"
		assignment_brief: Optional[str] = None
		
		# Get brief/rubric from form input (takes priority)
		assignment_brief = request.form.get("assignment_brief", "").strip() or None
		
		if task_id:
			task = sb_fetch_one(
				"""
				SELECT t.id, t.title, t.description, t.student_id
				FROM tasks t
				WHERE t.id = :task_id AND t.student_id = :student_id
				""",
				{"task_id": task_id, "student_id": current_user.id}
			)
			if task:
				assignment_title = task.get("title", "Assignment")
				# Use task description as brief if no form input provided
				if not assignment_brief:
					assignment_brief = task.get("description")
		
		if not error:
			try:
				service = get_chatgpt_service()
				review_result = service.review_and_grade_assignment(
					assignment_text=assignment_text,
					task_title=assignment_title,
					assignment_brief=assignment_brief
				)
				
				# Format comprehensive feedback
				feedback_parts = [review_result.feedback]
				if review_result.strengths:
					feedback_parts.append("\n\n**Strengths:**\n" + "\n".join(f"• {s}" for s in review_result.strengths))
				if review_result.weaknesses:
					feedback_parts.append("\n\n**Areas for Improvement:**\n" + "\n".join(f"• {w}" for w in review_result.weaknesses))
				if review_result.suggestions:
					feedback_parts.append("\n\n**Suggestions:**\n" + "\n".join(f"• {s}" for s in review_result.suggestions))
				comprehensive_feedback = "\n".join(feedback_parts)
				
				# Save review to database
				sb_execute(
					"""
					INSERT INTO assignment_reviews (
						task_id, student_id, filename, filepath, original_text,
						ai_feedback, ai_score_estimate, ai_possible_score, reviewed_at
					) VALUES (
						:task_id, :student_id, :filename, :filepath, :original_text,
						:ai_feedback, :ai_score_estimate, :ai_possible_score, NOW()
					)
					""",
					{
						"task_id": task_id,
						"student_id": current_user.id,
						"filename": filename,
						"filepath": filepath,
						"original_text": assignment_text[:5000] if assignment_text else None,
						"ai_feedback": comprehensive_feedback,
						"ai_score_estimate": review_result.score_estimate,
						"ai_possible_score": review_result.possible_score,
					}
				)
				success_message = "Assignment reviewed successfully!"
				# Reload reviews
				try:
					reviews = sb_fetch_all(
						"""
						SELECT ar.id, ar.task_id, ar.filename, ar.ai_feedback, ar.ai_score_estimate, 
						       ar.ai_possible_score, ar.reviewed_at,
						       t.title AS task_title, m.code AS module_code
						FROM assignment_reviews ar
						LEFT JOIN tasks t ON t.id = ar.task_id
						LEFT JOIN modules m ON m.id = t.module_id
						WHERE ar.student_id = :student_id
						ORDER BY ar.reviewed_at DESC
						LIMIT 20
						""",
						{"student_id": current_user.id}
					)
				except Exception:
					pass
			except ChatGPTClientError as exc:
				error = f"AI review failed: {str(exc)}"
				print(f"[assignment-review] ChatGPT error user={current_user.id} err={exc}")
			except Exception as exc:
				error = f"Unexpected error: {str(exc)}"
				print(f"[assignment-review] Unexpected error user={current_user.id} err={exc}")
				import traceback
				traceback.print_exc()
	
	return render_template(
		"assignment_review.html",
		tasks=tasks,
		reviews=reviews,
		error=error,
		success_message=success_message
	)


@app.route("/reviews")
@login_required
def reviews_history():
	page_str = request.args.get("page", "1").strip()
	try:
		page = max(1, int(page_str))
	except ValueError:
		page = 1
	per_page = 20
	offset = (page - 1) * per_page

	count_row = sb_fetch_one(
		"SELECT COUNT(*) as count FROM assignment_reviews WHERE student_id = :student_id",
		{"student_id": current_user.id}
	)
	total = count_row["count"] if count_row else 0
	total_pages = max(1, (total + per_page - 1) // per_page)

	reviews = sb_fetch_all(
		"""
		SELECT ar.id, ar.filename, ar.ai_feedback, ar.ai_score_estimate, ar.ai_possible_score, ar.reviewed_at,
		       t.title AS task_title, m.code AS module_code
		FROM assignment_reviews ar
		LEFT JOIN tasks t ON t.id = ar.task_id
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE ar.student_id = :student_id
		ORDER BY ar.reviewed_at DESC
		LIMIT :limit OFFSET :offset
		""",
		{"student_id": current_user.id, "limit": per_page, "offset": offset}
	)

	return render_template(
		"reviews_history.html",
		reviews=reviews,
		page=page,
		total_pages=total_pages
	)


@app.route("/activity")
@login_required
def activity_history():
	"""Activity history view with pagination."""
	# Reference: ChatGPT (OpenAI) - Activity History Pagination
	# Date: 2026-02-04
	# Prompt: "I need a paginated activity log page. It should count total rows,
	# fetch a page of recent activity, and display friendly titles."
	# ChatGPT provided the pagination and title-mapping pattern.
	page_str = request.args.get("page", "1").strip()
	try:
		page = max(1, int(page_str))
	except ValueError:
		page = 1
	per_page = 25
	offset = (page - 1) * per_page

	try:
		count_row = sb_fetch_one(
			"SELECT COUNT(*) as count FROM activity_logs WHERE student_id = :student_id",
			{"student_id": current_user.id}
		)
		total = count_row["count"] if count_row else 0
		total_pages = max(1, (total + per_page - 1) // per_page)

		activity_logs = sb_fetch_all(
			"""
			SELECT id, action_type, method, path, endpoint, status_code, duration_ms, created_at
			FROM activity_logs
			WHERE student_id = :student_id
			ORDER BY created_at DESC
			LIMIT :limit OFFSET :offset
			""",
			{"student_id": current_user.id, "limit": per_page, "offset": offset}
		)
		for item in activity_logs:
			item["title"] = _format_activity_title(item)
	except Exception:
		total = 0
		total_pages = 1
		activity_logs = []

	return render_template(
		"activity_history.html",
		activity_logs=activity_logs,
		page=page,
		total_pages=total_pages
	)


@app.route("/subtasks")
@login_required
def subtasks_history():
	page_str = request.args.get("page", "1").strip()
	try:
		page = max(1, int(page_str))
	except ValueError:
		page = 1
	per_page = 20
	offset = (page - 1) * per_page

	count_row = sb_fetch_one(
		"""
		SELECT COUNT(*) as count
		FROM subtasks s
		JOIN tasks t ON t.id = s.task_id
		WHERE t.student_id = :student_id
		""",
		{"student_id": current_user.id}
	)
	total = count_row["count"] if count_row else 0
	total_pages = max(1, (total + per_page - 1) // per_page)

	subtasks = sb_fetch_all(
		"""
		SELECT s.id, s.title, s.description, s.sequence, s.estimated_hours,
		       s.planned_start, s.planned_end, s.is_completed, s.completed_at,
		       t.title AS task_title, m.code AS module_code
		FROM subtasks s
		JOIN tasks t ON t.id = s.task_id
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE t.student_id = :student_id
		ORDER BY s.created_at DESC
		LIMIT :limit OFFSET :offset
		""",
		{"student_id": current_user.id, "limit": per_page, "offset": offset}
	)

	return render_template(
		"subtasks_history.html",
		subtasks=subtasks,
		page=page,
		total_pages=total_pages
	)


@app.route("/lecturer-messages")
@login_required
def lecturer_messages_history():
	page_str = request.args.get("page", "1").strip()
	try:
		page = max(1, int(page_str))
	except ValueError:
		page = 1
	per_page = 20
	offset = (page - 1) * per_page

	count_row = sb_fetch_one(
		"""
		SELECT COUNT(*) as count
		FROM lecturer_messages
		WHERE student_id = :student_id
		""",
		{"student_id": current_user.id}
	)
	total = count_row["count"] if count_row else 0
	total_pages = max(1, (total + per_page - 1) // per_page)

	messages = sb_fetch_all(
		"""
		SELECT lm.id, lm.subject, lm.message, lm.sent_at, lm.email_status, lm.email_error,
		       l.name AS lecturer_name, l.module_code
		FROM lecturer_messages lm
		LEFT JOIN lecturers l ON l.id = lm.lecturer_id
		WHERE lm.student_id = :student_id
		ORDER BY lm.sent_at DESC
		LIMIT :limit OFFSET :offset
		""",
		{"student_id": current_user.id, "limit": per_page, "offset": offset}
	)

	return render_template(
		"lecturer_messages_history.html",
		messages=messages,
		page=page,
		total_pages=total_pages
	)


@app.route("/lecturer-replies")
@login_required
def lecturer_replies():
	"""View lecturer replies received via email"""
	page_str = request.args.get("page", "1").strip()
	try:
		page = max(1, int(page_str))
	except ValueError:
		page = 1
	per_page = 20
	offset = (page - 1) * per_page
	
	# Count total replies
	count_row = sb_fetch_one(
		"SELECT COUNT(*) as count FROM lecturer_replies WHERE student_id = :student_id",
		{"student_id": current_user.id}
	)
	total = count_row["count"] if count_row else 0
	total_pages = max(1, (total + per_page - 1) // per_page)
	
	# Count unread
	unread_row = sb_fetch_one(
		"SELECT COUNT(*) as count FROM lecturer_replies WHERE student_id = :student_id AND is_read = FALSE",
		{"student_id": current_user.id}
	)
	unread_count = unread_row["count"] if unread_row else 0
	
	# Get replies
	replies = sb_fetch_all(
		"""
		SELECT lr.id, lr.from_email, lr.from_name, lr.subject, lr.body,
		       lr.received_at, lr.is_read, lr.lecturer_id,
		       l.name AS lecturer_name
		FROM lecturer_replies lr
		LEFT JOIN lecturers l ON l.id = lr.lecturer_id
		WHERE lr.student_id = :student_id
		ORDER BY lr.received_at DESC
		LIMIT :limit OFFSET :offset
		""",
		{"student_id": current_user.id, "limit": per_page, "offset": offset}
	)
	
	# Check if IMAP is configured
	imap_configured = get_imap_config() is not None
	
	return render_template(
		"lecturer_replies.html",
		replies=replies,
		page=page,
		total_pages=total_pages,
		total=total,
		unread_count=unread_count,
		imap_configured=imap_configured
	)


@app.route("/lecturer-replies/refresh", methods=["POST"])
@login_required
def refresh_lecturer_replies():
	"""Manually check for new lecturer replies via IMAP"""
	result = _fetch_lecturer_replies(current_user.id)
	
	if result["success"]:
		if result["new_count"] > 0:
			flash(f"Found {result['new_count']} new reply(ies) from lecturers!", "success")
		else:
			flash("No new replies found.", "info")
	else:
		flash(f"Failed to check for replies: {result['error']}", "error")
	
	return redirect(url_for("lecturer_replies"))


@app.route("/lecturer-replies/<int:reply_id>/read", methods=["POST"])
@login_required
def mark_reply_read(reply_id: int):
	"""Mark a lecturer reply as read"""
	try:
		sb_execute(
			"UPDATE lecturer_replies SET is_read = TRUE WHERE id = :id AND student_id = :student_id",
			{"id": reply_id, "student_id": current_user.id}
		)
	except Exception as exc:
		print(f"[replies] mark read failed: {exc}")
	
	return redirect(url_for("lecturer_replies"))


@app.route("/lecturer-replies/mark-all-read", methods=["POST"])
@login_required
def mark_all_replies_read():
	"""Mark all lecturer replies as read"""
	try:
		sb_execute(
			"UPDATE lecturer_replies SET is_read = TRUE WHERE student_id = :student_id",
			{"student_id": current_user.id}
		)
		flash("All replies marked as read.", "success")
	except Exception as exc:
		print(f"[replies] mark all read failed: {exc}")
		flash("Failed to mark replies as read.", "error")
	
	return redirect(url_for("lecturer_replies"))


@app.route("/assignment-review/<int:review_id>/delete", methods=["POST"])
@login_required
def delete_assignment_review(review_id: int):
	"""Delete a single assignment review (and stored file) for the current user."""
	try:
		review = sb_fetch_one(
			"""
			SELECT id, filepath
			FROM assignment_reviews
			WHERE id = :review_id AND student_id = :student_id
			""",
			{"review_id": review_id, "student_id": current_user.id}
		)
		if not review:
			flash("Review not found or permission denied.", "error")
			return redirect(url_for("assignment_review"))
		filepath = review.get("filepath")
		sb_execute(
			"""
			DELETE FROM assignment_reviews
			WHERE id = :review_id AND student_id = :student_id
			""",
			{"review_id": review_id, "student_id": current_user.id}
		)
		if filepath and os.path.exists(filepath):
			try:
				os.remove(filepath)
			except Exception:
				pass
		flash("Assignment review deleted.", "success")
	except Exception as exc:
		print(f"[assignment-review] delete failed user={current_user.id} review={review_id} err={exc}")
		flash("Failed to delete assignment review.", "error")
	return redirect(url_for("assignment_review"))


@app.route("/contact-lecturer", methods=["POST"])
@login_required
def contact_lecturer():
	"""Send a Quick Connect message to a lecturer via SMTP."""
	# Reference: ChatGPT (OpenAI) - Quick Connect Lecturer Email Flow
	# Date: 2026-01-23
	# Prompt: "I need a simple flow that lets students pick a lecturer, enter a subject
	# and message, and send it via SMTP with validation. Can you outline the pattern?"
	# ChatGPT provided the validation and send flow used here.
	action = request.form.get("action", "send").strip()
	lecturer_id = request.form.get("lecturer_id", "").strip()
	subject = request.form.get("subject", "").strip()
	request_text = request.form.get("request_text", "").strip()
	message = request.form.get("message", "").strip()
	if not lecturer_id or not subject:
		flash("Lecturer and subject are required.", "error")
		return redirect(url_for("index"))
	try:
		lecturer = sb_fetch_one(
			"SELECT name, email FROM lecturers WHERE id = :id",
			{"id": int(lecturer_id)}
		)
	except Exception:
		lecturer = None
	if not lecturer:
		flash("Lecturer not found.", "error")
		return redirect(url_for("index"))

	if action == "generate":
		if not request_text:
			flash("Please describe what you want from the lecturer.", "error")
			return redirect(url_for("index"))
		try:
			service = get_chatgpt_service()
			draft = service.draft_lecturer_email(
				student_name=current_user.name,
				student_id=current_user.student_number or current_user.id,
				lecturer_name=lecturer.get("name"),
				subject=subject,
				request_text=request_text,
			)
		except ChatGPTClientError as exc:
			flash(f"Failed to generate draft: {exc}", "error")
			return redirect(url_for("index"))
		except Exception as exc:
			print(f"[lecturers] draft failed user={current_user.id} err={exc}")
			flash("Failed to generate draft. Try again.", "error")
			return redirect(url_for("index"))
		session["lecturer_draft"] = draft
		session["lecturer_subject"] = subject
		session["lecturer_request"] = request_text
		session["lecturer_id"] = lecturer_id
		return redirect(url_for("index"))

	if not message:
		flash("Message is required to send.", "error")
		return redirect(url_for("index"))

	student_id_value = current_user.student_number or current_user.id
	signature = f"\n\nRegards,\n{current_user.name} (Student ID: {student_id_value})"
	lower_message = message.lower()
	if "regards," not in lower_message and "student id" not in lower_message:
		message = message.rstrip() + signature

	student_id_value = current_user.student_number or current_user.id
	full_body = (
		f"From: {current_user.name} ({current_user.email})\n"
		f"Student ID: {student_id_value}\n\n"
		f"{message}"
	)
	error = _send_reminder_email(
		to_email=lecturer.get("email"),
		subject=subject,
		body=full_body
	)
	try:
		sb_execute(
			"""
			INSERT INTO lecturer_messages (
				student_id, lecturer_id, subject, message, sent_at, email_status, email_error
			) VALUES (
				:student_id, :lecturer_id, :subject, :message, NOW(), :status, :error
			)
			""",
			{
				"student_id": current_user.id,
				"lecturer_id": int(lecturer_id),
				"subject": subject,
				"message": message,
				"status": "sent" if error is None else "failed",
				"error": error,
			}
		)
	except Exception as exc:
		print(f"[lecturers] message log failed user={current_user.id} err={exc}")
	if error:
		flash(f"Failed to send message: {error}", "error")
	else:
		flash(f"Message sent to {lecturer.get('name')}.", "success")
	return redirect(url_for("index"))


@app.route("/sync-canvas", methods=["GET", "POST"])
@login_required
def sync_canvas():
	"""Sync assignments from Canvas LMS"""
	if request.method == "POST":
		if not current_user.canvas_api_token:
			flash("Please add your Canvas API token in your profile first.", "error")
			return redirect(url_for("sync_canvas"))
		
		try:
			from canvas_sync import sync_canvas_assignments, sync_canvas_calendar_events
			import signal
			
			CANVAS_URL = "https://ucc.instructure.com"
			
			print(f"[sync] Starting assignment sync for user {current_user.id}...")
			stats = sync_canvas_assignments(
				canvas_url=CANVAS_URL,
				api_token=current_user.canvas_api_token,
				student_id=current_user.id,
				db_execute=sb_execute,
				db_fetch_all=sb_fetch_all,
				db_fetch_one=sb_fetch_one
			)
			print(f"[sync] Assignment sync complete: {stats}")

			print(f"[sync] Starting calendar events sync...")
			cal_stats = {"events_new": 0, "events_updated": 0, "errors": 0}
			try:
				api_token = current_user.canvas_api_token
				student_id = current_user.id
				
				# Reference: ChatGPT (OpenAI) - Background Threading with Timeout
				# Date: 2025-11-04
				# Prompt: [Threading pattern for Canvas sync - prompt from documentation]
				# ChatGPT suggested a pattern using threading and queue to run the Canvas sync in a background 
				# worker with a 15 second timeout. The worker thread calls the sync function and reports either 
				# success or an error back through a queue. If the sync exceeds the timeout, the system aborts 
				# the operation, records an error, and displays a warning to the user instead of blocking the request.
				# Reference: Python Documentation - threading Module
				# https://docs.python.org/3/library/threading.html
				# Reference: Python Documentation - queue Module
				# https://docs.python.org/3/library/queue.html
				import threading
				import queue
				
				result_queue = queue.Queue()
				
				def sync_worker():
					try:
						result = sync_canvas_calendar_events(
							canvas_url=CANVAS_URL,
							api_token=api_token,
							student_id=student_id,
							db_execute=sb_execute,
							db_fetch_all=sb_fetch_all,
							db_fetch_one=sb_fetch_one
						)
						result_queue.put(("success", result))
					except Exception as e:
						result_queue.put(("error", str(e)))
				
				thread = threading.Thread(target=sync_worker, daemon=True)
				thread.start()
				thread.join(timeout=45)  # 45 second timeout (increased for large calendars)
				
				if thread.is_alive():
					print(f"[sync] Calendar events sync timed out after 45 seconds")
					cal_stats = {"events_new": 0, "events_updated": 0, "errors": 1}
					flash("Calendar events sync timed out (skipped)", "warning")
				else:
					if not result_queue.empty():
						status, result = result_queue.get()
						if status == "success":
							cal_stats = result
							print(f"[sync] Calendar events sync complete: {cal_stats}")
						else:
							print(f"[sync] Calendar events sync failed: {result}")
							cal_stats = {"events_new": 0, "events_updated": 0, "errors": 1}
							flash(f"Calendar events sync skipped: {result}", "warning")
					else:
						print(f"[sync] Calendar events sync returned no result")
						cal_stats = {"events_new": 0, "events_updated": 0, "errors": 1}
			except Exception as cal_error:
				print(f"[sync] Calendar events sync exception: {cal_error}")
				cal_stats = {"events_new": 0, "events_updated": 0, "errors": 1}
				flash(f"Calendar events sync skipped: {str(cal_error)}", "warning")
			
			if stats['errors']:
				for error in stats['errors']:
					flash(error, "warning")
			
			success_msg = f"✅ Canvas Sync Complete! "
			success_msg += f"Courses: {stats['courses_found']}, "
			success_msg += f"New: {stats['assignments_new']}, "
			success_msg += f"Updated: {stats['assignments_updated']}, "
			success_msg += f"Skipped: {stats['assignments_skipped']}"
			success_msg += f" | Events +{cal_stats['events_new']}/~{cal_stats['events_updated']}"
			
			if stats['modules_created'] > 0:
				success_msg += f", Modules Created: {stats['modules_created']}"
			
			flash(success_msg, "success")
			return redirect(url_for("sync_canvas"))
			
		except Exception as e:
			flash(f"Canvas sync failed: {str(e)}", "error")
			return redirect(url_for("sync_canvas"))
	
	has_token = current_user.canvas_api_token is not None and current_user.canvas_api_token != ""
	
	canvas_tasks_count = sb_fetch_one(
		"SELECT COUNT(*) as count FROM tasks WHERE student_id = :student_id AND canvas_assignment_id IS NOT NULL",
		{"student_id": current_user.id}
	)
	canvas_tasks = canvas_tasks_count['count'] if canvas_tasks_count else 0
	
	total_tasks = sb_fetch_one(
		"SELECT COUNT(*) as count FROM tasks WHERE student_id = :student_id",
		{"student_id": current_user.id}
	)
	total = total_tasks['count'] if total_tasks else 0
	
	return render_template("sync_canvas.html", 
	                      has_token=has_token, 
	                      canvas_tasks=canvas_tasks,
	                      total_tasks=total)


@app.route("/add-data")
@login_required
def add_data_form():
	"""Display forms for adding modules and tasks"""
	modules = sb_fetch_all("SELECT id, code FROM modules ORDER BY code")
	return render_template("add_data.html", modules=modules)


# Reference: ChatGPT (OpenAI) - Voice Transcript Parsing for Task Forms
# Date: 2026-02-10
# Prompt: "I need a Flask endpoint to parse spoken task text into structured fields
# (title, module code, due date, due time, weight, status) for auto-filling a form.
# Can you provide a robust parsing pattern with fallback defaults?"
# ChatGPT provided the helper + endpoint structure below, including date/time parsing
# and module matching fallbacks for imperfect transcripts.
_VOICE_WEEKDAY_MAP = {
	"monday": 0,
	"tuesday": 1,
	"wednesday": 2,
	"thursday": 3,
	"friday": 4,
	"saturday": 5,
	"sunday": 6,
}


def _voice_parse_due_date(transcript: str, base_date) -> Optional[str]:
	text = transcript.lower()
	iso_match = re.search(r"\b(\d{4})-(\d{2})-(\d{2})\b", text)
	if iso_match:
		try:
			return datetime(int(iso_match.group(1)), int(iso_match.group(2)), int(iso_match.group(3))).date().isoformat()
		except ValueError:
			pass

	slash_match = re.search(r"\b(\d{1,2})/(\d{1,2})/(\d{2,4})\b", text)
	if slash_match:
		day = int(slash_match.group(1))
		month = int(slash_match.group(2))
		year = int(slash_match.group(3))
		if year < 100:
			year += 2000
		try:
			return datetime(year, month, day).date().isoformat()
		except ValueError:
			pass

	if "today" in text:
		return base_date.isoformat()
	if "tomorrow" in text:
		return (base_date + timedelta(days=1)).isoformat()

	for weekday_name, weekday_idx in _VOICE_WEEKDAY_MAP.items():
		if f"next {weekday_name}" in text or f"on {weekday_name}" in text or f"{weekday_name}" in text:
			delta = (weekday_idx - base_date.weekday()) % 7
			if delta == 0 or f"next {weekday_name}" in text:
				delta += 7
			return (base_date + timedelta(days=delta)).isoformat()
	return None


def _voice_parse_due_time(transcript: str) -> Optional[str]:
	text = transcript.lower()
	hhmm_match = re.search(r"\b(?:at\s*)?([01]?\d|2[0-3]):([0-5]\d)\b", text)
	if hhmm_match:
		hour = int(hhmm_match.group(1))
		minute = int(hhmm_match.group(2))
		return f"{hour:02d}:{minute:02d}"

	ampm_match = re.search(r"\b(?:at\s*)?(\d{1,2})(?::([0-5]\d))?\s*(am|pm)\b", text)
	if ampm_match:
		hour = int(ampm_match.group(1))
		minute = int(ampm_match.group(2) or "0")
		period = ampm_match.group(3)
		if period == "pm" and hour < 12:
			hour += 12
		if period == "am" and hour == 12:
			hour = 0
		if 0 <= hour <= 23:
			return f"{hour:02d}:{minute:02d}"
	return None


# Reference: ChatGPT (OpenAI) - Voice Command Intent Router + Action Handlers
# Date: 2026-02-11
# Prompt: "I already have voice-to-task autofill. I now need one endpoint that detects
# voice intents for (1) dashboard query ('what is due this week/any overdue'),
# (2) recurring calendar event creation, and (3) microtask generation from a spoken
# command. Can you provide a safe Flask pattern with helper functions, regex parsing,
# DB inserts, and structured JSON responses?"
# ChatGPT provided the intent-router + helper structure below, including command
# detection, recurring event parsing, microtask creation flow, and response payload
# patterns adapted to this codebase.
def _voice_parse_time_token(value: str) -> Optional[str]:
	"""Parse a single time token (e.g., 6pm, 18:30) into HH:MM."""
	text = (value or "").strip().lower()
	if not text:
		return None
	hhmm_match = re.fullmatch(r"([01]?\d|2[0-3]):([0-5]\d)", text)
	if hhmm_match:
		return f"{int(hhmm_match.group(1)):02d}:{int(hhmm_match.group(2)):02d}"
	ampm_match = re.fullmatch(r"(\d{1,2})(?::([0-5]\d))?\s*(am|pm)", text)
	if ampm_match:
		hour = int(ampm_match.group(1))
		minute = int(ampm_match.group(2) or "0")
		period = ampm_match.group(3)
		if period == "pm" and hour < 12:
			hour += 12
		if period == "am" and hour == 12:
			hour = 0
		if 0 <= hour <= 23:
			return f"{hour:02d}:{minute:02d}"
	return None


def _voice_detect_intent(transcript: str) -> str:
	text = transcript.lower()
	if any(phrase in text for phrase in ["what is due this week", "what's due this week", "any overdue", "overdue tasks", "due this week"]):
		return "dashboard_query"
	if "break" in text and "into" in text and any(word in text for word in ["steps", "microtasks", "subtasks"]):
		return "create_microtasks"
	if any(word in text for word in ["mark", "set", "update"]) and any(word in text for word in ["done", "completed", "in progress", "pending"]):
		return "update_task_status"
	if any(word in text for word in ["move", "reschedule"]) and any(word in text for word in ["today", "tomorrow", "next ", " on ", " at ", ":"]):
		return "reschedule_task"
	if any(word in text for word in ["weight", "percent", "priority"]) and any(word in text for word in ["set", "update", "change", "make"]):
		return "set_task_weight"
	if any(word in text for word in ["add", "create", "schedule"]) and "every" in text:
		return "create_recurring_event"
	return "unknown"


def _voice_match_module(modules: List[Dict[str, Any]], transcript: str) -> Tuple[Optional[int], Optional[str]]:
	upper_text = transcript.upper()
	for module in modules:
		code = (module.get("code") or "").upper().strip()
		if code and re.search(rf"\b{re.escape(code)}\b", upper_text):
			return module.get("id"), code
	return None, None


# Reference: ChatGPT (OpenAI) - Voice Task Lifecycle Command Handlers
# Date: 2026-02-18
# Prompt: "Extend voice commands to handle task lifecycle actions in Flask:
# mark task status (done/in-progress/pending), reschedule due date/time, and
# update weighting/priority. I need task title extraction + fuzzy matching and
# safe DB update patterns. Can you provide helper functions and intent hooks?"
# ChatGPT provided the normalization, task matching, and update handler pattern
# adapted below for this codebase.
def _voice_normalize_title(value: str) -> str:
	return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def _voice_extract_task_phrase(transcript: str) -> Optional[str]:
	text = (transcript or "").strip()
	patterns = [
		r"(?:mark|set|update|change|move|reschedule|make)\s+(.+?)\s+(?:as\s+)?(?:done|completed|in progress|pending|todo)\b",
		r"(?:move|reschedule|set|change)\s+(.+?)\s+(?:to|for|on)\b",
		r"(?:set|update|change|make)\s+(.+?)\s+(?:weight|priority)\b",
	]
	for pattern in patterns:
		match = re.search(pattern, text, flags=re.IGNORECASE)
		if match:
			candidate = (match.group(1) or "").strip(" ,.-")
			if candidate:
				return candidate
	return None


def _voice_find_task_for_command(transcript: str, student_id: int) -> Optional[Dict[str, Any]]:
	modules = sb_fetch_all("SELECT id, code FROM modules ORDER BY code")
	module_id, _ = _voice_match_module(modules, transcript)
	candidate = _voice_extract_task_phrase(transcript) or transcript
	candidate_norm = _voice_normalize_title(candidate)
	if not candidate_norm:
		return None
	rows = sb_fetch_all(
		"""
		SELECT t.id, t.title, t.status, t.due_date, t.due_at, t.weight_percentage, m.code AS module_code
		FROM tasks t
		LEFT JOIN modules m ON m.id = t.module_id
		WHERE t.student_id = :student_id
		  AND (:module_id IS NULL OR t.module_id = :module_id)
		ORDER BY COALESCE(t.due_at, t.due_date::timestamp) ASC NULLS LAST, t.id DESC
		LIMIT 300
		""",
		{"student_id": student_id, "module_id": module_id},
	)
	best_row = None
	best_score = 0
	candidate_tokens = {t for t in candidate_norm.split() if len(t) > 2}
	for row in rows:
		title_norm = _voice_normalize_title(row.get("title") or "")
		if not title_norm:
			continue
		score = 0
		if candidate_norm == title_norm:
			score += 100
		if candidate_norm in title_norm:
			score += 60
		elif title_norm in candidate_norm:
			score += 45
		title_tokens = {t for t in title_norm.split() if len(t) > 2}
		overlap = len(candidate_tokens & title_tokens)
		score += overlap * 8
		if score > best_score:
			best_score = score
			best_row = row
	return best_row if best_score >= 16 else None


def _voice_update_task_status(transcript: str, student_id: int) -> Dict[str, Any]:
	text = transcript.lower()
	new_status = "pending"
	if "in progress" in text:
		new_status = "in_progress"
	elif "done" in text or "completed" in text or "finish" in text:
		new_status = "completed"
	elif "pending" in text or "todo" in text or "to do" in text:
		new_status = "pending"
	task = _voice_find_task_for_command(transcript, student_id)
	if not task:
		raise ValueError("Could not match a task title in that command.")
	sb_execute(
		"""
		UPDATE tasks
		SET status = :status,
		    completed_at = CASE
		        WHEN :status = 'completed' THEN COALESCE(completed_at, NOW())
		        ELSE NULL
		    END
		WHERE id = :task_id
		  AND student_id = :student_id
		""",
		{"status": new_status, "task_id": task.get("id"), "student_id": student_id},
	)
	return {
		"task_id": task.get("id"),
		"title": task.get("title"),
		"status": new_status,
	}


def _voice_reschedule_task(transcript: str, student_id: int) -> Dict[str, Any]:
	base_date = datetime.now().date()
	due_date_str = _voice_parse_due_date(transcript, base_date)
	due_time_str = _voice_parse_due_time(transcript)
	if not due_date_str:
		raise ValueError("Please include a new due date, e.g. 'move database lab to next Friday'.")
	task = _voice_find_task_for_command(transcript, student_id)
	if not task:
		raise ValueError("Could not match a task title in that command.")
	due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()
	due_at = None
	if due_time_str:
		due_at = datetime.fromisoformat(f"{due_date_str}T{due_time_str}")
	elif task.get("due_at"):
		existing_dt = task.get("due_at")
		if hasattr(existing_dt, "hour") and hasattr(existing_dt, "minute"):
			due_at = datetime.combine(due_date, time(existing_dt.hour, existing_dt.minute))
	sb_execute(
		"""
		UPDATE tasks
		SET due_date = :due_date,
		    due_at = :due_at
		WHERE id = :task_id
		  AND student_id = :student_id
		""",
		{
			"due_date": due_date,
			"due_at": due_at,
			"task_id": task.get("id"),
			"student_id": student_id,
		},
	)
	return {
		"task_id": task.get("id"),
		"title": task.get("title"),
		"due_date": due_date_str,
		"due_time": due_time_str,
	}


def _voice_set_task_weight(transcript: str, student_id: int) -> Dict[str, Any]:
	text = transcript.lower()
	weight = None
	percent_match = re.search(r"(\d{1,3}(?:\.\d+)?)\s*(?:%|percent)\b", text)
	if percent_match:
		weight = float(percent_match.group(1))
	else:
		if "high" in text:
			weight = 25.0
		elif "medium" in text:
			weight = 15.0
		elif "low" in text:
			weight = 5.0
	if weight is None:
		raise ValueError("Please include a weight like '20 percent' or priority high/medium/low.")
	weight = max(0.0, min(100.0, weight))
	task = _voice_find_task_for_command(transcript, student_id)
	if not task:
		raise ValueError("Could not match a task title in that command.")
	sb_execute(
		"""
		UPDATE tasks
		SET weight_percentage = :weight
		WHERE id = :task_id
		  AND student_id = :student_id
		""",
		{"weight": weight, "task_id": task.get("id"), "student_id": student_id},
	)
	return {
		"task_id": task.get("id"),
		"title": task.get("title"),
		"weight_percentage": round(weight, 1),
	}


def _voice_dashboard_query(student_id: int) -> Dict[str, Any]:
	today = datetime.now().date()
	week_end = today + timedelta(days=7)
	due_this_week = []
	overdue = []
	try:
		due_this_week = sb_fetch_all(
			"""
			SELECT t.id, t.title, t.status, t.due_date, t.due_at, m.code AS module_code
			FROM tasks t
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			  AND t.status != 'completed'
			  AND COALESCE(t.due_at::date, t.due_date) >= :today
			  AND COALESCE(t.due_at::date, t.due_date) <= :week_end
			ORDER BY COALESCE(t.due_at, t.due_date::timestamp) ASC
			LIMIT 10
			""",
			{"student_id": student_id, "today": today, "week_end": week_end},
		)
	except Exception:
		due_this_week = []
	try:
		overdue = sb_fetch_all(
			"""
			SELECT t.id, t.title, t.status, t.due_date, t.due_at, m.code AS module_code
			FROM tasks t
			LEFT JOIN modules m ON m.id = t.module_id
			WHERE t.student_id = :student_id
			  AND t.status != 'completed'
			  AND COALESCE(t.due_at::date, t.due_date) < :today
			ORDER BY COALESCE(t.due_at, t.due_date::timestamp) ASC
			LIMIT 10
			""",
			{"student_id": student_id, "today": today},
		)
	except Exception:
		overdue = []
	return {
		"due_this_week_count": len(due_this_week),
		"overdue_count": len(overdue),
		"due_this_week": due_this_week,
		"overdue": overdue,
	}


def _voice_parse_until_date(transcript: str, base_date) -> Optional[datetime.date]:
	text = transcript.lower()
	# Try explicit date first (YYYY-MM-DD / DD/MM/YYYY / weekday phrases)
	parsed = _voice_parse_due_date(transcript, base_date)
	if parsed:
		try:
			return datetime.strptime(parsed, "%Y-%m-%d").date()
		except ValueError:
			pass
	month_match = re.search(r"\buntil\s+([a-z]+)(?:\s+(\d{4}))?\b", text)
	if not month_match:
		return None
	month_name = month_match.group(1).strip().title()
	year = int(month_match.group(2)) if month_match.group(2) else base_date.year
	try:
		month_num = datetime.strptime(month_name, "%B").month
	except ValueError:
		try:
			month_num = datetime.strptime(month_name[:3], "%b").month
		except ValueError:
			return None
	# If month already passed this year and no year was provided, roll to next year.
	if not month_match.group(2) and month_num < base_date.month:
		year += 1
	# End date = last day of month
	if month_num == 12:
		next_month = datetime(year + 1, 1, 1).date()
	else:
		next_month = datetime(year, month_num + 1, 1).date()
	return next_month - timedelta(days=1)


def _voice_create_recurring_event(transcript: str, student_id: int) -> Dict[str, Any]:
	text = transcript.lower()
	weekday_name = None
	weekday_idx = None
	for name, idx in _VOICE_WEEKDAY_MAP.items():
		if re.search(rf"\bevery\s+{name}\b", text):
			weekday_name = name.title()
			weekday_idx = idx
			break
	if weekday_idx is None:
		raise ValueError("Please include a weekday, e.g. 'every Tuesday'.")

	time_range_match = re.search(
		r"\b(\d{1,2}(?::\d{2})?\s*(?:am|pm)?|\d{1,2}:\d{2})\s*(?:to|-|until)\s*(\d{1,2}(?::\d{2})?\s*(?:am|pm)?|\d{1,2}:\d{2})\b",
		text,
	)
	if not time_range_match:
		raise ValueError("Please include a time range, e.g. '6pm to 8pm'.")
	start_time_text = time_range_match.group(1)
	end_time_text = time_range_match.group(2)
	start_time_hhmm = _voice_parse_time_token(start_time_text)
	end_time_hhmm = _voice_parse_time_token(end_time_text)
	if not start_time_hhmm or not end_time_hhmm:
		raise ValueError("Could not parse event start/end time.")

	title_match = re.search(r"\b(?:add|create|schedule)\s+(.+?)\s+every\b", text)
	event_title = title_match.group(1).strip(" ,.-").title() if title_match else "Voice Event"
	if not event_title:
		event_title = "Voice Event"

	today = datetime.now().date()
	recur_until = _voice_parse_until_date(transcript, today) or (today + timedelta(days=56))
	delta = (weekday_idx - today.weekday()) % 7
	first_date = today + timedelta(days=delta)
	start_hour, start_minute = [int(x) for x in start_time_hhmm.split(":")]
	end_hour, end_minute = [int(x) for x in end_time_hhmm.split(":")]

	created = 0
	current_date = first_date
	while current_date <= recur_until:
		start_dt = datetime.combine(current_date, time(start_hour, start_minute))
		end_dt = datetime.combine(current_date, time(end_hour, end_minute))
		if end_dt <= start_dt:
			end_dt = end_dt + timedelta(days=1)
		sb_execute(
			"""
			INSERT INTO events (student_id, title, start_at, end_at, location, is_recurring, recurrence_end_date)
			VALUES (:student_id, :title, :start_at, :end_at, :location, :is_recurring, :recurrence_end_date)
			""",
			{
				"student_id": student_id,
				"title": event_title[:255],
				"start_at": start_dt,
				"end_at": end_dt,
				"location": None,
				"is_recurring": True,
				"recurrence_end_date": recur_until,
			},
		)
		created += 1
		current_date = current_date + timedelta(days=7)

	return {
		"title": event_title,
		"weekday": weekday_name,
		"start_time": start_time_hhmm,
		"end_time": end_time_hhmm,
		"until": recur_until.isoformat(),
		"created_count": created,
	}


def _voice_create_microtasks(transcript: str, student_id: int) -> Dict[str, Any]:
	text = transcript.lower()
	modules = sb_fetch_all("SELECT id, code FROM modules ORDER BY code")
	module_id, module_code = _voice_match_module(modules, transcript)

	step_match = re.search(r"\binto\s+(\d{1,2})\s+(?:steps|microtasks|subtasks)\b", text)
	if not step_match:
		raise ValueError("Please include number of steps, e.g. 'into 5 steps'.")
	step_count = max(2, min(12, int(step_match.group(1))))

	title_match = re.search(
		r"\bbreak(?:\s+down)?\s+(.+?)\s+into\s+\d{1,2}\s+(?:steps|microtasks|subtasks)\b",
		transcript,
		flags=re.IGNORECASE,
	)
	assignment_title = (title_match.group(1).strip(" ,.-") if title_match else "").strip()
	if not assignment_title:
		raise ValueError("Please include assignment name, e.g. 'Break IS4416 database report into 5 steps'.")
	if module_code:
		assignment_title = re.sub(rf"\b{re.escape(module_code)}\b", "", assignment_title, flags=re.IGNORECASE).strip(" -")

	due_date_str = _voice_parse_due_date(transcript, datetime.now().date())
	task_id = _lookup_or_create_task(
		assignment_title=assignment_title[:255],
		module_id=module_id,
		module_code=module_code,
		student_id=student_id,
		due_date_str=due_date_str,
	)
	if not task_id:
		raise ValueError("Could not create or find a task for this assignment.")

	row = sb_fetch_one("SELECT COALESCE(MAX(sequence), 0) AS max_seq FROM subtasks WHERE task_id = :task_id", {"task_id": task_id})
	base_sequence = int(row.get("max_seq") or 0) if row else 0

	template_titles = [
		"Clarify assignment brief",
		"Research and collect sources",
		"Build outline / structure",
		"Draft main content",
		"Add references and polish",
		"Final proofread and submit",
	]
	due_date_obj = None
	if due_date_str:
		try:
			due_date_obj = datetime.strptime(due_date_str, "%Y-%m-%d").date()
		except ValueError:
			due_date_obj = None
	start_date = datetime.now().date()
	span_days = max((due_date_obj - start_date).days, 0) if due_date_obj else 0

	created_titles = []
	for i in range(step_count):
		if i < len(template_titles):
			subtask_title = template_titles[i]
		else:
			subtask_title = f"Step {i + 1}"
		planned_day = None
		if due_date_obj:
			offset = int((span_days * (i / max(step_count - 1, 1)))) if step_count > 1 else 0
			planned_day = start_date + timedelta(days=offset)
		sb_execute(
			"""
			INSERT INTO subtasks (
				task_id, title, description, sequence, estimated_hours,
				planned_start, planned_end
			) VALUES (
				:task_id, :title, :description, :sequence, :estimated_hours,
				:planned_start, :planned_end
			)
			""",
			{
				"task_id": task_id,
				"title": subtask_title[:255],
				"description": None,
				"sequence": base_sequence + i + 1,
				"estimated_hours": None,
				"planned_start": planned_day,
				"planned_end": planned_day,
			},
		)
		created_titles.append(subtask_title)

	return {
		"task_id": task_id,
		"task_title": assignment_title,
		"module_code": module_code,
		"due_date": due_date_str,
		"created_count": step_count,
		"subtasks": created_titles,
	}


@app.route("/voice-parse", methods=["POST"])
@login_required
def voice_parse():
	payload = request.get_json(silent=True) or {}
	transcript = (payload.get("transcript") or "").strip()
	if not transcript:
		return jsonify({"ok": False, "error": "Transcript is required."}), 400

	try:
		modules = sb_fetch_all("SELECT id, code FROM modules ORDER BY code")
	except Exception:
		modules = []

	now_date = datetime.now().date()
	upper_text = transcript.upper()
	module_id = None
	module_code = None
	for module in modules:
		code = (module.get("code") or "").upper().strip()
		if code and re.search(rf"\b{re.escape(code)}\b", upper_text):
			module_id = module.get("id")
			module_code = code
			break

	weight_percentage = None
	weight_match = re.search(r"\b(\d{1,3}(?:\.\d+)?)\s*(?:%|percent)\b", transcript.lower())
	if weight_match:
		try:
			weight_percentage = max(0.0, min(100.0, float(weight_match.group(1))))
		except ValueError:
			weight_percentage = None

	status = "pending"
	lt = transcript.lower()
	if "in progress" in lt or "ongoing" in lt:
		status = "in_progress"
	elif "completed" in lt or "done" in lt or "finished" in lt:
		status = "completed"

	clean_title = re.sub(
		r"^\s*(add|create|new)\s+(a\s+)?(task|assignment)\s*(called|named)?\s*",
		"",
		transcript,
		flags=re.IGNORECASE,
	).strip()
	clean_title = re.split(r"\b(?:due|by|on|at)\b", clean_title, maxsplit=1, flags=re.IGNORECASE)[0].strip(" ,.-")
	title = clean_title or transcript[:120]

	due_date = _voice_parse_due_date(transcript, now_date)
	due_time = _voice_parse_due_time(transcript)

	warnings = []
	if module_id is None:
		warnings.append("No module code detected in voice note. Please choose one manually.")
	if due_date is None:
		warnings.append("No due date detected. Please set a due date.")

	return jsonify({
		"ok": True,
		"parsed": {
			"title": title,
			"module_id": module_id,
			"module_code": module_code,
			"due_date": due_date,
			"due_time": due_time,
			"weight_percentage": weight_percentage,
			"status": status,
		},
		"warnings": warnings,
	}), 200


# Reference: ChatGPT (OpenAI) - Voice Command API Response Contract
# Date: 2026-02-11
# Prompt: "For multiple voice actions, I need a single Flask endpoint that returns
# consistent JSON (`ok`, `intent`, `message`, `data`) and gracefully handles unknown
# intents and user-facing validation errors. Can you provide a clean route pattern?"
# ChatGPT provided the standardized endpoint response contract and intent dispatch
# pattern used in this route.
@app.route("/voice-command", methods=["POST"])
@login_required
def voice_command():
	payload = request.get_json(silent=True) or {}
	transcript = (payload.get("transcript") or "").strip()
	if not transcript:
		return jsonify({"ok": False, "error": "Transcript is required."}), 400

	intent = _voice_detect_intent(transcript)
	try:
		if intent == "dashboard_query":
			data = _voice_dashboard_query(current_user.id)
			return jsonify({
				"ok": True,
				"intent": intent,
				"message": f"You have {data['overdue_count']} overdue and {data['due_this_week_count']} due in the next 7 days.",
				"data": data,
			}), 200

		if intent == "create_recurring_event":
			data = _voice_create_recurring_event(transcript, current_user.id)
			return jsonify({
				"ok": True,
				"intent": intent,
				"message": f"Created {data['created_count']} recurring calendar events for '{data['title']}'.",
				"data": data,
			}), 200

		if intent == "create_microtasks":
			data = _voice_create_microtasks(transcript, current_user.id)
			return jsonify({
				"ok": True,
				"intent": intent,
				"message": f"Created {data['created_count']} microtasks for '{data['task_title']}'.",
				"data": data,
			}), 200

		if intent == "update_task_status":
			data = _voice_update_task_status(transcript, current_user.id)
			return jsonify({
				"ok": True,
				"intent": intent,
				"message": f"Updated '{data['title']}' to {data['status'].replace('_', ' ')}.",
				"data": data,
			}), 200

		if intent == "reschedule_task":
			data = _voice_reschedule_task(transcript, current_user.id)
			time_part = f" at {data['due_time']}" if data.get("due_time") else ""
			return jsonify({
				"ok": True,
				"intent": intent,
				"message": f"Rescheduled '{data['title']}' to {data['due_date']}{time_part}.",
				"data": data,
			}), 200

		if intent == "set_task_weight":
			data = _voice_set_task_weight(transcript, current_user.id)
			return jsonify({
				"ok": True,
				"intent": intent,
				"message": f"Set weighting for '{data['title']}' to {data['weight_percentage']}%.",
				"data": data,
			}), 200

		return jsonify({
			"ok": False,
			"intent": "unknown",
			"error": "I could not detect a supported command. Try: 'Mark database report as done', 'Move IS4416 lab to next Friday 3pm', 'Set strategy essay to 20 percent', 'What is due this week?', or 'Break IS4416 database report into 5 steps due next Thursday'.",
		}), 400
	except Exception as exc:
		print(f"[voice-command] failed user={current_user.id} intent={intent} err={exc}")
		return jsonify({"ok": False, "intent": intent, "error": str(exc)}), 400


@app.route("/add-module", methods=["POST"])
@login_required
def add_module():
	"""Add a new module to the database"""
	try:
		code = request.form.get("code", "").strip().upper()
		if not code:
			flash("Module code is required", "error")
			return redirect(url_for("add_data_form"))
		
		sb_execute(
			"INSERT INTO modules (code) VALUES (:code)",
			{"code": code}
		)
		flash(f"Module '{code}' added successfully!", "success")
	except Exception as e:
		flash(f"Error adding module: {str(e)}", "error")
	
	return redirect(url_for("add_data_form"))


@app.route("/add-task", methods=["POST"])
@login_required
def add_task():
	"""Add a new task/assignment for the current user"""
	try:
		title = request.form.get("title", "").strip()
		module_id = request.form.get("module_id")
		due_date = request.form.get("due_date")
		due_time = request.form.get("due_time")
		description = request.form.get("description", "").strip() or None
		weight = request.form.get("weight_percentage")
		status = request.form.get("status", "pending")
		
		if not title:
			flash("Task title is required", "error")
			return redirect(url_for("add_data_form"))
		if not module_id:
			flash("Please select a module", "error")
			return redirect(url_for("add_data_form"))
		if not due_date:
			flash("Due date is required", "error")
			return redirect(url_for("add_data_form"))
		
		due_at_value = None
		if due_time:
			try:
				dt_combined = datetime.fromisoformat(f"{due_date}T{due_time}")
				due_at_value = dt_combined
			except ValueError:
				flash("Invalid due time provided; saving as all-day deadline.", "warning")
				due_at_value = None
		
		weight_value = None
		if weight:
			try:
				weight_value = float(weight)
			except ValueError:
				flash("Weighting must be a number.", "warning")
				weight_value = None
		
		sb_execute(
			"""INSERT INTO tasks (title, student_id, module_id, due_date, due_at, description, status, weight_percentage) 
			   VALUES (:title, :student_id, :module_id, :due_date, :due_at, :description, :status, :weight_percentage)""",
			{
				"title": title,
				"student_id": current_user.id,
				"module_id": int(module_id),
				"due_date": due_date,
				"due_at": due_at_value,
				"description": description,
				"status": status,
				"weight_percentage": weight_value
			}
		)
		flash(f"Task '{title}' added successfully!", "success")
	except Exception as e:
		flash(f"Error adding task: {str(e)}", "error")
	
	return redirect(url_for("add_data_form"))


@app.route("/add-lecturer", methods=["POST"])
@login_required
def add_lecturer():
	"""Add a lecturer contact for Quick Connect."""
	try:
		name = request.form.get("name", "").strip()
		email = request.form.get("email", "").strip()
		module_code = request.form.get("module_code", "").strip().upper() or None
		if not name or not email:
			flash("Lecturer name and email are required.", "error")
			return redirect(url_for("add_data_form"))
		sb_execute(
			"""
			INSERT INTO lecturers (name, email, module_code)
			VALUES (:name, :email, :module_code)
			""",
			{"name": name, "email": email, "module_code": module_code}
		)
		flash("Lecturer added successfully.", "success")
	except Exception as exc:
		print(f"[lecturers] add failed user={current_user.id} err={exc}")
		flash("Failed to add lecturer.", "error")
	return redirect(url_for("add_data_form"))


@app.route("/update-task-status/<int:task_id>", methods=["POST"])
@login_required
def update_task_status(task_id):
	"""Update the status of a task"""
	try:
		status = request.form.get("status")
		if status not in ["pending", "in_progress", "completed"]:
			flash("Invalid status", "error")
			return redirect(url_for("tasks"))
		
		task = sb_fetch_one("SELECT id FROM tasks WHERE id = :id AND student_id = :student_id", 
		                    {"id": task_id, "student_id": current_user.id})
		if not task:
			flash("Task not found or you don't have permission to update it", "error")
			return redirect(url_for("tasks"))
		
		if status == "completed":
			sb_execute(
				"""UPDATE tasks 
				   SET status = :status, completed_at = NOW() 
				   WHERE id = :task_id AND student_id = :student_id""",
				{"status": status, "task_id": task_id, "student_id": current_user.id}
			)
		else:
			sb_execute(
				"""UPDATE tasks 
				   SET status = :status, completed_at = NULL 
				   WHERE id = :task_id AND student_id = :student_id""",
				{"status": status, "task_id": task_id, "student_id": current_user.id}
			)
		
		flash("Task status updated successfully!", "success")
	except Exception as e:
		flash(f"Error updating task: {str(e)}", "error")
	
	return redirect(url_for("tasks"))


# ── ICS Calendar Export ──────────────────────────────────────────────
# Reference: ChatGPT (OpenAI) - ICS Calendar Export with icalendar Library
# Date: 2026-02-10
# Prompt: "I need a Flask route that generates an ICS feed from my tasks
# (due dates) and calendar events tables so students can subscribe in
# Google Calendar or Outlook. Can you show me how to build VEVENT entries
# with the icalendar library and return the file as a download?"
# ChatGPT provided the route pattern, VEVENT construction, and Content-Type headers.
@app.route("/export/calendar.ics")
@login_required
def export_ics():
	"""Generate an ICS file for the student's tasks and events."""
	from icalendar import Calendar as ICalendar, Event as IEvent, vText
	import pytz

	include_tasks = request.args.get("include_tasks", "true") == "true"
	include_events = request.args.get("include_events", "true") == "true"

	cal = ICalendar()
	cal.add("prodid", "-//SmartStudentPlanner//EN")
	cal.add("version", "2.0")
	cal.add("calscale", "GREGORIAN")
	cal.add("x-wr-calname", f"{current_user.name} – Academic Schedule")

	utc = pytz.UTC

	# ── Tasks (assignment deadlines) ──
	if include_tasks:
		try:
			tasks = sb_fetch_all(
				"""
				SELECT t.id, t.title, t.due_date, t.due_at, t.status,
				       t.description, m.code AS module_code
				FROM tasks t
				LEFT JOIN modules m ON m.id = t.module_id
				WHERE t.student_id = :sid AND t.due_date IS NOT NULL
				ORDER BY t.due_date
				""",
				{"sid": current_user.id},
			)
		except Exception:
			tasks = []

		for t in tasks:
			ev = IEvent()
			title = t.get("title", "Task")
			mod = t.get("module_code")
			ev.add("summary", f"📋 {title} [{mod}]" if mod else f"📋 {title}")
			ev.add("uid", f"task-{t['id']}@smartstudentplanner")

			due_at = t.get("due_at")
			due_date = t.get("due_date")
			if due_at:
				try:
					dt = due_at if hasattr(due_at, "hour") else datetime.fromisoformat(str(due_at))
					if dt.tzinfo is None:
						dt = utc.localize(dt)
					ev.add("dtstart", dt)
					ev.add("dtend", dt + timedelta(minutes=30))
				except Exception:
					ev.add("dtstart", due_date)
			else:
				ev.add("dtstart", due_date)

			desc_parts = []
			if t.get("status"):
				desc_parts.append(f"Status: {t['status']}")
			if t.get("description"):
				desc_parts.append(t["description"])
			if desc_parts:
				ev.add("description", "\n".join(desc_parts))

			cal.add_component(ev)

	# ── Calendar events (lectures, manual blocks) ──
	if include_events:
		try:
			events = sb_fetch_all(
				"""
				SELECT e.id, e.title, e.start_at, e.end_at, e.location,
				       m.code AS module_code, e.canvas_course_id
				FROM events e
				LEFT JOIN modules m ON m.id = e.module_id
				WHERE e.student_id = :sid
				ORDER BY e.start_at
				""",
				{"sid": current_user.id},
			)
		except Exception:
			events = []

		for e in events:
			ev = IEvent()
			title = e.get("title", "Event")
			mod = e.get("module_code")
			prefix = "📚 " if e.get("canvas_course_id") else "📅 "
			ev.add("summary", f"{prefix}{title} [{mod}]" if mod else f"{prefix}{title}")
			ev.add("uid", f"event-{e['id']}@smartstudentplanner")

			start_at = e.get("start_at")
			end_at = e.get("end_at")
			try:
				st = start_at if hasattr(start_at, "hour") else datetime.fromisoformat(str(start_at))
				en = end_at if hasattr(end_at, "hour") else datetime.fromisoformat(str(end_at))
				if st.tzinfo is None:
					st = utc.localize(st)
				if en.tzinfo is None:
					en = utc.localize(en)
				ev.add("dtstart", st)
				ev.add("dtend", en)
			except Exception:
				continue

			if e.get("location"):
				ev.add("location", vText(e["location"]))

			cal.add_component(ev)

	response = app.response_class(
		cal.to_ical(),
		mimetype="text/calendar",
		headers={"Content-Disposition": "attachment; filename=timetable.ics"},
	)
	return response


if __name__ == "__main__":
	test_database_connection()
	
	cfg = get_flask_config()
	app.run(host=cfg.host, port=cfg.port, debug=cfg.debug)
